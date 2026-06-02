#!/usr/bin/env python3
"""
商場管理 Dashboard（/mall/overview）教學手冊自動生成腳本
=========================================================
使用方式：
  pip install playwright python-docx
  playwright install chromium
  python mall_overview_generate_manual.py

輸出：
  manual_screenshots_mall_overview/
  商場管理Dashboard教學手冊.docx
"""

import asyncio
import sys
from pathlib import Path
from datetime import datetime

try:
    from playwright.async_api import async_playwright
except ImportError:
    print("❌ 請先執行：pip install playwright && playwright install chromium"); sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx"); sys.exit(1)

# ── 設定 ──────────────────────────────────────────────────────────────────────
BASE_URL    = "http://localhost:5173"
TARGET_URL  = f"{BASE_URL}/mall/overview"
OUTPUT_DIR  = Path(__file__).parent / "manual_screenshots_mall_overview"
MANUAL_PATH = Path(__file__).parent / "商場管理Dashboard教學手冊.docx"
USERNAME    = "admin"
PASSWORD    = "Admin@2026"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
GREEN   = RGBColor(0x52, 0xC4, 0x1A)
ORANGE  = RGBColor(0xFA, 0xAD, 0x14)
RED     = RGBColor(0xCF, 0x13, 0x22)
GRAY    = RGBColor(0x59, 0x59, 0x59)

CONTENT_X  = 185
CONTENT_W  = 1390
VIEWPORT_H = 900


# ── 截圖邏輯 ──────────────────────────────────────────────────────────────────
async def capture_screenshots(page, output_dir: Path) -> dict:
    shots = {}

    async def snap(name: str, scroll_y: int = 0, clip=None, wait_ms: int = 800):
        try:
            await page.evaluate(f"window.scrollTo(0, {scroll_y})")
            await page.wait_for_timeout(wait_ms)
            path = output_dir / f"{name}.png"
            await page.screenshot(path=str(path), full_page=False,
                                  **({"clip": clip} if clip else {}))
            shots[name] = path
            print(f"   ✓ {name}")
        except Exception as e:
            print(f"   ✗ {name}：{e}")

    async def click_tab(tab_text: str, wait_ms: int = 2000):
        try:
            tab = page.locator("div.ant-tabs-tab").filter(has_text=tab_text).first
            await tab.click()
            await page.wait_for_timeout(wait_ms)
        except Exception as e:
            print(f"   ⚠ Tab [{tab_text}] 失敗：{e}")

    # ── 登入 ──────────────────────────────────────────────────────────────────
    print("🔐 登入中...")
    await page.goto(f"{BASE_URL}/login")
    await page.wait_for_load_state("domcontentloaded")
    await page.wait_for_timeout(800)
    await page.locator("input[placeholder*='帳號'], input[type='text']").first.fill(USERNAME)
    await page.locator("input[type='password']").first.fill(PASSWORD)
    await page.locator("button[type='submit'], button:has-text('登入')").first.click()
    try:
        await page.wait_for_url(f"{BASE_URL}/dashboard", timeout=15000)
        print("   ✓ 登入成功")
    except Exception:
        print(f"   ✓ 已跳轉至 {page.url}")

    # ── 導航 ──────────────────────────────────────────────────────────────────
    await page.goto(TARGET_URL)
    await page.wait_for_load_state("networkidle")
    try:
        await page.locator(".ant-card").first.wait_for(timeout=12000)
        await page.wait_for_timeout(2500)
        print("   ✓ 頁面載入完成")
    except Exception:
        await page.wait_for_timeout(5000)

    # ════════════════════════════════════════════════════════════
    # Tab A：Dashboard
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：Tab A Dashboard")
    await page.evaluate("window.scrollTo(0, 0)")
    await page.wait_for_timeout(600)

    # 1. 頁頭 + 篩選列（含年月 + 巡檢日期 + PPT 匯出按鈕）
    await snap("01_header_filter",
               clip={"x": CONTENT_X, "y": 55, "width": CONTENT_W, "height": 185})

    # 2. 彙總 KPI 列（5 張小卡）
    await snap("02_kpi_aggregate",
               clip={"x": CONTENT_X, "y": 240, "width": CONTENT_W, "height": 130})

    # 3. 第一排來源卡片（商場例行維護 / 全棟例行維護 / 商場工務巡檢 / 整棟巡檢）
    await snap("03_source_row1",
               clip={"x": CONTENT_X, "y": 385, "width": CONTENT_W, "height": 210})

    # 4. 第二排來源卡片（商場工務報修 / 商場主管交辦 / 商場緊急事件）
    await snap("04_source_row2",
               clip={"x": CONTENT_X, "y": 610, "width": CONTENT_W, "height": 210})

    # 5. 圖表區（兩個 Bar + Pie + Line）
    await snap("05_charts",
               scroll_y=820,
               clip={"x": CONTENT_X, "y": 820, "width": CONTENT_W, "height": 400})

    # 6. 費用摘要（從 luqunData）
    await snap("06_fee_summary",
               scroll_y=1220,
               clip={"x": CONTENT_X, "y": 1220, "width": CONTENT_W, "height": 200})

    # ════════════════════════════════════════════════════════════
    # Tab B：每日累計
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：Tab B 每日累計")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("B. 每日累計", wait_ms=2500)
    await snap("07_tab_b_daily",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab C：每月累計
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：Tab C 每月累計")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("C. 每月累計", wait_ms=2500)
    await snap("08_tab_c_monthly",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab D：每年累計
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：Tab D 每年累計")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("D. 每年累計", wait_ms=2500)
    await snap("09_tab_d_yearly",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab：人員工時%
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：人員工時%")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("人員工時%", wait_ms=2500)
    await snap("10_person_pct",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab：人員排名
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：人員排名")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("人員排名", wait_ms=2000)
    await snap("11_ranking",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    return shots


# ── Word 手冊輔助函式 ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def add_image_safe(doc, path, width_inches: float = 5.8):
    if isinstance(path, Path) and path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph("[ 截圖未能取得，請參考系統實際畫面 ]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = GRAY; p.runs[0].font.italic = True


def add_tip(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run("💡 提示："); r.bold = True; r.font.color.rgb = PRIMARY
    p.add_run(text)


def hdr_table(doc, headers, bg="1B3A5C"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, bg)
    return table


def h1(doc, t):
    h = doc.add_heading(t, 1); h.runs[0].font.color.rgb = PRIMARY; return h

def h2(doc, t):
    h = doc.add_heading(t, 2); h.runs[0].font.color.rgb = ACCENT; return h


# ── Word 手冊建構 ─────────────────────────────────────────────────────────────
def build_word_manual(shots: dict, output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11.69); sec.page_width  = Inches(8.27)
    sec.top_margin  = Inches(0.9);   sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1.1);   sec.right_margin  = Inches(1.1)

    # ── 封面 ──────────────────────────────────────────────────────────────────
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_heading("商場管理 Dashboard", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY; t.runs[0].font.size = Pt(28)
    t2 = doc.add_heading("教學操作手冊", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT; t2.runs[0].font.size = Pt(22)
    doc.add_paragraph()
    sub = doc.add_paragraph(
        f"路由：/mall/overview\n{datetime.now().strftime('%Y 年 %m 月 %d 日')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    doc.add_page_break()

    # ── 一、功能簡介 ──────────────────────────────────────────────────────────
    h1(doc, "一、功能簡介")
    p = doc.add_paragraph()
    p.add_run("商場管理 Dashboard").bold = True
    p.add_run(
        "（路徑：商場管理 ▶ ★商場管理 Dashboard，路由：/mall/overview）"
        "整合商場所有工務作業的即時狀態，讓管理者一頁掌握七大作業來源的健康度、工時與完成率。"
        "與飯店管理 Dashboard 架構對應，但資料來源為商場端。"
    )
    doc.add_paragraph()

    tab_info = [
        ("A. Dashboard（預設）", "7 來源狀態卡片 + 彙總 KPI + 圖表 + 費用摘要"),
        ("B. 每日累計",          "五大工項類別 × 每日工時表（年月篩選）"),
        ("C. 每月累計",          "五大工項類別 × 每月工時矩陣（年份篩選）"),
        ("D. 每年累計",          "年度累計工時走勢（年份篩選）"),
        ("人員工時%",            "五大來源 × Top-15 人員工時佔比表格"),
        ("人員排名",             "人員工時排名 + 堆疊來源 Bar Chart"),
    ]
    table = hdr_table(doc, ["頁籤", "說明"])
    for name, desc in tab_info:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ── 二、Tab A：篩選列與頁頭工具 ──────────────────────────────────────────
    h1(doc, "二、Tab A — 篩選列與頁頭工具")
    add_image_safe(doc, shots.get("01_header_filter"))
    doc.add_paragraph()
    doc.add_paragraph("Tab A 頂部的篩選列控制 7 個來源卡片與圖表資料（各 Tab 有獨立篩選）：")

    filter_items = [
        ("年度",       "選擇工務資料的年份"),
        ("月份",       "選「全年」看全年彙整；選特定月份看當月資料"),
        ("巡檢日期",   "DatePicker——控制「商場工務巡檢」的日別統計口徑；點「今日」快速回到今天"),
        ("全部重新整理","重新載入所有 7 個來源的資料"),
        ("匯出 PowerPoint","紫色漸層按鈕，將當前 KPI 資料匯出為 .pptx（需選定特定月份）"),
    ]
    table = hdr_table(doc, ["控制項", "說明"])
    for name, desc in filter_items:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    add_tip(doc, "年月篩選改變後，7 個來源卡片會各自重新載入資料，可能需要 3–8 秒才能全部完成。")
    doc.add_page_break()

    # ── 三、Tab A：彙總 KPI ────────────────────────────────────────────────────
    h1(doc, "三、Tab A — 彙總 KPI 列")
    add_image_safe(doc, shots.get("02_kpi_aggregate"))
    doc.add_paragraph()
    doc.add_paragraph("7 個來源卡片上方有 5 張彙總 KPI 卡，匯總四大主要來源（商場例行維護、全棟例行維護、商場工務巡檢、商場工務報修）的數字：")

    kpi_rows = [
        ("本期總工項",   "四大來源案件/工項數加總", "深藍"),
        ("已完成工項",   "四大來源已完成件數加總", "綠色"),
        ("本期工時合計", "四大來源工時加總（hr，含問號圖示說明口徑）", "橙色"),
        ("異常/未完成",  "各來源異常/未完成件數加總；=0 顯示「全部正常」", ">0 紅色 / =0 綠色"),
        ("逾期未完成",   "各來源逾期件數加總（主要來自例行維護逾期）", ">0 深紅 / =0 綠色"),
    ]
    table = hdr_table(doc, ["KPI 名稱", "說明", "顏色"])
    for name, desc, color in kpi_rows:
        row = table.add_row()
        for i, v in enumerate([name, desc, color]):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ── 四、Tab A：7 來源狀態卡片 ─────────────────────────────────────────────
    h1(doc, "四、Tab A — 7 來源狀態卡片")

    h2(doc, "4.1  第一排（維護 / 巡檢類，4 張）")
    add_image_safe(doc, shots.get("03_source_row1"))
    doc.add_paragraph()

    row1 = [
        ("商場例行維護", "深藍 #1B3A5C", "/mall/periodic-maintenance",   "商場 PM 完成率"),
        ("全棟例行維護", "藍色 #4BA8E8", "/mall/full-building-maintenance","全棟 PM 完成率"),
        ("商場工務巡檢", "紫色 #722ED1", "/mall-facility-inspection",     "5 張巡檢表場次完成率"),
        ("整棟巡檢",     "綠色 #52C41A", "/full-building-inspection",      "整棟巡檢批次完成率"),
    ]
    table = hdr_table(doc, ["來源名稱", "卡片顏色", "跳轉路徑", "完成率口徑"])
    for name, color, route, desc in row1:
        row = table.add_row()
        for i, v in enumerate([name, color, route, desc]):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    h2(doc, "4.2  第二排（報修 / 交辦 / 緊急事件類，3 張）")
    add_image_safe(doc, shots.get("04_source_row2"))
    doc.add_paragraph()

    row2 = [
        ("商場工務報修", "橙色 #FA8C16", "/luqun-repair/dashboard", "結案率（已結案/總件數）"),
        ("商場主管交辦", "深紅 #C0392B", "/mall/other-tasks",       "本期件數 + 工時（無完成率）"),
        ("商場緊急事件", "紅色 #D4380D", "/mall/other-tasks",       "本期件數 + 工時（無完成率）"),
    ]
    table = hdr_table(doc, ["來源名稱", "卡片顏色", "跳轉路徑", "顯示重點"])
    for name, color, route, desc in row2:
        row = table.add_row()
        for i, v in enumerate([name, color, route, desc]):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    add_tip(doc, "主管交辦和緊急事件卡片沒有完成率（因為這兩類任務通常不做完成率統計），只顯示件數和工時。點擊任一卡片可跳轉至對應模組詳細頁。")
    doc.add_page_break()

    # ── 五、Tab A：圖表 & 費用摘要 ────────────────────────────────────────────
    h1(doc, "五、Tab A — 圖表區 & 費用摘要")
    add_image_safe(doc, shots.get("05_charts"))
    doc.add_paragraph()
    chart_items = [
        ("各來源工項數比較（水平 Bar）", "每列一個來源，深色=總數、淺色=完成數，快速比較各來源工作量"),
        ("各來源完成率（水平 Bar）",     "每列一個來源，顯示完成率（%），可對比維護類與報修類的效率差異"),
        ("工時來源占比（圓餅圖）",       "顯示各來源工時比例，停留顯示 HR 數與占比"),
        ("商場報修 12 個月趨勢（折線）", "商場工務報修的近 12 月報修件數（深色）與結案件數（淺色）趨勢"),
    ]
    for title, desc in chart_items:
        p = doc.add_paragraph()
        p.add_run(f"▸ {title}：").bold = True
        p.add_run(desc)

    doc.add_paragraph()

    h2(doc, "5.1  費用摘要")
    add_image_safe(doc, shots.get("06_fee_summary"))
    doc.add_paragraph()
    doc.add_paragraph(
        "圖表下方顯示商場工務報修（陸群）的費用摘要，"
        "包含委外+維修費用（累計 YTD）、扣款費用、當月金額小計。"
        "資料直接來自 luqunData 的 kpi 欄位，與報修來源卡片共用同一組資料。"
    )
    doc.add_page_break()

    # ── 六、Tab B — 每日累計 ──────────────────────────────────────────────────
    h1(doc, "六、Tab B — 每日累計")
    add_image_safe(doc, shots.get("07_tab_b_daily"))
    doc.add_paragraph()
    doc.add_paragraph(
        "「B. 每日累計」頁籤顯示商場工務報修的每日工時累計，"
        "以五大工項類別（現場報修、上級交辦、緊急事件、例行維護、每日巡檢）為列，"
        "當月每一天為欄，展示各類別每日的案件件數與工時。"
        "\n\n"
        "左上角年月選單控制查詢範圍（此篩選器獨立於 Tab A 的篩選器）。"
        "首次切換到此 Tab 時系統自動載入資料（懶載入）。"
    )
    add_tip(doc, "此 Tab 的工時口徑依「occupied_at」欄位歸屬日期，與 Tab A 商場工務報修卡片的口徑一致。")
    doc.add_page_break()

    # ── 七、Tab C — 每月累計 ──────────────────────────────────────────────────
    h1(doc, "七、Tab C — 每月累計")
    add_image_safe(doc, shots.get("08_tab_c_monthly"))
    doc.add_paragraph()
    doc.add_paragraph(
        "「C. 每月累計」頁籤以年度為橫軸，五大工項類別為縱軸，"
        "顯示全年 12 個月的工時累計矩陣。"
        "未來月份欄位顯示「—」。"
        "\n\n"
        "左上角年份選單控制查詢年度（獨立篩選）。"
        "可點欄位標題排序，找出工時最高的月份或類別。"
    )
    doc.add_page_break()

    # ── 八、Tab D — 每年累計 ──────────────────────────────────────────────────
    h1(doc, "八、Tab D — 每年累計")
    add_image_safe(doc, shots.get("09_tab_d_yearly"))
    doc.add_paragraph()
    doc.add_paragraph(
        "「D. 每年累計」以年度視角顯示各月份工時的累計走勢（Running Total），"
        "適合比較全年工務量的月份分布，找出旺季與淡季規律。"
        "年份選單獨立控制，可切換不同年度比較。"
    )
    doc.add_page_break()

    # ── 九、人員工時% & 人員排名 ─────────────────────────────────────────────
    h1(doc, "九、人員工時% & 人員排名")

    h2(doc, "9.1  人員工時%")
    add_image_safe(doc, shots.get("10_person_pct"))
    doc.add_paragraph()
    doc.add_paragraph(
        "以熱度表格顯示：列 = 五大工項類別（現場報修、上級交辦、緊急事件、例行維護、每日巡檢），"
        "欄 = 工時排名前 15 的人員。"
        "每格顯示該人員在該類別的工時佔個人總工時的百分比。"
        "\n\n"
        "顏色規則：≥30% 紅色（高度集中）、15–30% 橙色、>0% 綠色、—灰色（無工時）。"
    )
    doc.add_paragraph()

    h2(doc, "9.2  人員排名")
    add_image_safe(doc, shots.get("11_ranking"))
    doc.add_paragraph()
    doc.add_paragraph(
        "以降序列出各人員的工時數，搭配堆疊 Bar Chart 直觀呈現各人員的工時來源結構。"
        "Bar Chart 以三大類別（現場報修=橙色、例行維護=深藍、每日巡檢=紫色）堆疊顯示。"
        "\n\n"
        "表格包含：排名、人員姓名、總工時（HR）、工時占比（%），以及各類別的工時分解。"
        "\n\n"
        "人員工時% 與人員排名兩個 Tab 共享同一組資料載入，"
        "切換到其中一個後，另一個也不需要再等待。"
    )
    doc.add_page_break()

    # ── 十、常見操作情境 ──────────────────────────────────────────────────────
    h1(doc, "十、常見操作情境")
    scenarios = [
        ("月初快速確認商場工務整體健康狀態",
         ["進入商場管理 Dashboard（Tab A 預設開啟）",
          "年月選本月，等待 7 個來源卡片全部載入",
          "掃視卡片顏色：全部綠色代表正常；橙色/紅色代表有警示",
          "查看彙總 KPI 的「異常/未完成」與「逾期未完成」是否為 0"]),
        ("查看商場例行維護本月完成進度",
         ["Tab A → 找「商場例行維護」卡片，查看完成率",
          "若完成率低於 70%，點卡片跳轉至商場例行維護詳細清單",
          "在清單頁篩選「未完成」確認待辦項目"]),
        ("分析本月哪個工項佔用最多工時",
         ["Tab A → 查看「工時來源占比」圓餅圖",
          "停留在最大區塊看是哪個來源的工時最多",
          "切換到 Tab B 每日累計查看各工項的每日分布",
          "切換到 Tab C 每月累計查看全年趨勢比較"]),
        ("找出工時最重的工務人員",
         ["切換到「人員排名」Tab",
          "第一名即工時最重的人員",
          "切換到「人員工時%」Tab，找到該人員欄位確認工時集中在哪個類別"]),
        ("匯出月報 PowerPoint",
         ["確認年月選到目標月份（不可選全年）",
          "點頁頭右側「匯出 PowerPoint」紫色按鈕",
          "等待下載完成"]),
    ]
    for i, (title, steps) in enumerate(scenarios, 1):
        p = doc.add_paragraph()
        p.add_run(f"情境 {i}：{title}").bold = True
        p.runs[0].font.color.rgb = PRIMARY
        for step in steps:
            doc.add_paragraph(f"　→ {step}")
        doc.add_paragraph()

    # ── 附錄 ──────────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "附錄：7 來源跳轉路徑對照")
    table = hdr_table(doc, ["來源名稱", "顏色", "跳轉路徑"])
    sources = [
        ("商場例行維護", "深藍 #1B3A5C",  "/mall/periodic-maintenance"),
        ("全棟例行維護", "藍色 #4BA8E8",  "/mall/full-building-maintenance"),
        ("商場工務巡檢", "紫色 #722ED1",  "/mall-facility-inspection/dashboard"),
        ("整棟巡檢",     "綠色 #52C41A",  "/full-building-inspection/dashboard"),
        ("商場工務報修", "橙色 #FA8C16",  "/luqun-repair/dashboard"),
        ("商場主管交辦", "深紅 #C0392B",  "/mall/other-tasks"),
        ("商場緊急事件", "紅色 #D4380D",  "/mall/other-tasks"),
    ]
    for name, color, route in sources:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = color; row.cells[2].text = route
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 手冊已儲存：{output_path}")


# ── 主程式 ────────────────────────────────────────────────────────────────────
async def main():
    OUTPUT_DIR.mkdir(exist_ok=True)
    print("=" * 55)
    print("  商場管理 Dashboard 教學手冊生成器")
    print("=" * 55)
    print(f"目標：{TARGET_URL}")
    print(f"截圖：{OUTPUT_DIR}")
    print(f"手冊：{MANUAL_PATH}\n")

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=["--disable-blink-features=AutomationControlled"],
        )
        context = await browser.new_context(
            viewport={"width": 1600, "height": VIEWPORT_H},
            device_scale_factor=1.5,
            locale="zh-TW",
            timezone_id="Asia/Taipei",
        )
        page = await context.new_page()
        shots = await capture_screenshots(page, OUTPUT_DIR)
        await browser.close()

    print(f"\n截圖完成，共 {len(shots)} 張")
    print("\n📄 生成 Word 手冊...")
    build_word_manual(shots, MANUAL_PATH)
    print(f"\n✅ 執行完成！請開啟：{MANUAL_PATH}")


if __name__ == "__main__":
    asyncio.run(main())
