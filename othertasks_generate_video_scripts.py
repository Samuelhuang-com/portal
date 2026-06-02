#!/usr/bin/env python3
"""
主管交辦／緊急事件（/hotel/other-tasks）影音教學旁白腳本生成器
=============================================================
使用方式：
  python othertasks_generate_video_scripts.py

輸出：主管交辦緊急事件_影音教學腳本.docx
"""

import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx"); sys.exit(1)

OUTPUT_PATH = Path(__file__).parent / "主管交辦緊急事件_影音教學腳本.docx"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
ORANGE  = RGBColor(0xFA, 0xAD, 0x14)
GRAY    = RGBColor(0x59, 0x59, 0x59)


# ── 輔助函式 ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def narration_block(doc, text: str):
    for part in text.split("\n\n"):
        part = part.strip()
        if not part:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(part); r.font.size = Pt(11.5)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "12"); left.set(qn("w:color"), "4BA8E8")
        pBdr.append(left); pPr.append(pBdr)


def label_box(doc, label: str, text: str, fill_hex: str):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    lc = table.rows[0].cells[0]
    set_cell_bg(lc, fill_hex)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label); lr.bold = True; lr.font.size = Pt(10)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = table.rows[0].cells[1]
    set_cell_bg(rc, "F8F9FA")
    rc.paragraphs[0].add_run(text).font.size = Pt(10.5)
    doc.add_paragraph()


def action_block(doc, text: str):  label_box(doc, "🖱 操作", text, "52C41A")
def hint_block(doc, text: str):    label_box(doc, "⏸ 提示", text, "FA8C16")


def timing_row(doc, t_range: str, label: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"[{t_range}]  ")
    r1.bold = True; r1.font.color.rgb = ORANGE; r1.font.size = Pt(10)
    r2 = p.add_run(label)
    r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = PRIMARY


def section_divider(doc):
    p = doc.add_paragraph("─" * 64)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── 腳本資料 ──────────────────────────────────────────────────────────────────
SCRIPTS = [

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第一集",
        "title": "功能簡介 & 篩選列操作",
        "duration": "建議時長：3–4 分鐘",
        "goal": "讓觀眾了解主管交辦與緊急事件的用途與差異，以及七個篩選條件的使用方式。",
        "segments": [
            {
                "time": "00:00–00:30",
                "label": "功能簡介",
                "narration": (
                    "大家好，歡迎收看主管交辦／緊急事件教學系列。"
                    "這個頁面管理兩種特殊的工務任務記錄："
                    "\n\n"
                    "第一種是「上級交辦」——主管臨時指派的非例行工務任務，"
                    "例如某個設備需要特別處理，或者有跨部門的協調工作。"
                    "\n\n"
                    "第二種是「緊急事件」——飯店或商場發生的突發性緊急情況，"
                    "例如漏水、停電、設備突然故障等，需要立刻處理的事項。"
                    "\n\n"
                    "這兩種記錄資料來源都是 Ragic，以建立日期作為月份歸屬。"
                ),
                "actions": ["進入頁面，讓觀眾看到完整畫面"],
                "hints": [],
            },
            {
                "time": "00:30–01:10",
                "label": "進入方式與歸屬自動篩選",
                "narration": (
                    "這個頁面有兩個入口："
                    "從「飯店管理」側邊欄點進來，頁面會自動把「歸屬」篩選設為「飯店」；"
                    "從「商場管理」點進來則預設「商場」。"
                    "\n\n"
                    "所以如果你想同時看飯店和商場的記錄，"
                    "可以把歸屬篩選清空，就會顯示全部的記錄。"
                    "\n\n"
                    "畫面頂部顯示標題「主管交辦／緊急事件」，"
                    "底下是篩選列，再下方是兩個頁籤的清單。"
                ),
                "actions": [
                    "指向歸屬篩選的「飯店」設定",
                    "示範清空歸屬篩選",
                ],
                "hints": [],
            },
            {
                "time": "01:10–02:40",
                "label": "七個篩選條件詳解",
                "narration": (
                    "篩選列有七個條件，我們依序說明。"
                    "\n\n"
                    "第一個「年份」——選擇要查看的年度，可以清空顯示所有年份。"
                    "\n\n"
                    "第二個「月份」——選擇月份，不選則查全年。"
                    "\n\n"
                    "第三個「狀態」——選擇案件狀態，選項是從 Ragic 動態取得的，"
                    "包含：結案、已結案、處理中、候辦、待辦、待排程、取消等。"
                    "\n\n"
                    "第四個「交辦主管」——篩選特定主管指派的任務，"
                    "如果你只負責某位主管的業務，篩選後就只看到那位主管的交辦。"
                    "\n\n"
                    "第五個「工程人員」——篩選特定工程師承接的任務，"
                    "適合主管查某位工程師的工作量。"
                    "\n\n"
                    "第六個「歸屬」——選飯店或商場，剛才已經說過了。"
                    "\n\n"
                    "第七個「關鍵字搜尋」——在問題說明和備註兩個欄位裡搜尋，"
                    "輸入關鍵字後按 Enter 或點搜尋圖示觸發。"
                    "\n\n"
                    "所有篩選條件都是即時生效的，改了之後清單和頁籤數字徽章自動更新，"
                    "不需要按「查詢」按鈕。"
                ),
                "actions": [
                    "逐一指向七個篩選條件",
                    "示範選擇一個狀態篩選，觀察清單即時更新",
                    "示範輸入關鍵字搜尋",
                ],
                "hints": [],
            },
            {
                "time": "02:40–03:15",
                "label": "Tab 件數 Badge",
                "narration": (
                    "頁籤標題旁邊有一個件數徽章——"
                    "「上級交辦」旁邊是深藍色數字，「緊急事件」旁邊是深紅色數字。"
                    "\n\n"
                    "這個數字會隨篩選條件即時更新，"
                    "所以當你選了「待辦」狀態篩選，"
                    "兩個 Tab 的數字就會分別告訴你各類型有幾件待辦任務，"
                    "不用切換頁籤就能掌握整體狀況。"
                    "\n\n"
                    "好，這集篩選列介紹完了。下一集看清單和 Drawer 明細。"
                ),
                "actions": [
                    "指向 Tab Badge 數字",
                    "改變篩選條件，展示數字即時更新",
                ],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第二集",
        "title": "任務清單解讀 & Tab 切換",
        "duration": "建議時長：3 分鐘",
        "goal": "讓觀眾掌握清單各欄位的意義、列底色警示、分頁設定，以及兩個 Tab 的切換方式。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "這集介紹清單的各個欄位和 Tab 切換方式。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–02:00",
                "label": "清單欄位解讀",
                "narration": (
                    "清單欄位從左到右說明："
                    "\n\n"
                    "「歸屬」——藍色 Tag 是飯店、綠色 Tag 是商場。"
                    "\n\n"
                    "「建立日期」——任務或事件的建立時間，也是月份歸屬的依據。"
                    "\n\n"
                    "「交辦主管」——指派此任務的主管名稱。"
                    "\n\n"
                    "「工程人員」——負責執行的工程師名稱。"
                    "\n\n"
                    "「問題說明」——任務的核心描述，如果內容太長會以「...」截斷，"
                    "點擊該列展開 Drawer 才能看到完整說明。"
                    "\n\n"
                    "「備註」——補充說明，同樣超出寬度以「...」截斷。"
                    "\n\n"
                    "「最後更新日期」——最後一次修改記錄的時間，"
                    "可以判斷案件是否最近有更新進度。"
                    "\n\n"
                    "「狀態」——以彩色 Tag 顯示：綠色是結案/完成，藍色是處理中，"
                    "黃色是候辦/待辦/待排程，灰色是取消。"
                    "\n\n"
                    "「維修工時」——工時（小時），大於零時以深藍加粗顯示，零或無資料顯示橫線。"
                    "\n\n"
                    "特別注意——狀態是「候辦」、「待辦」、「待排程」的列，"
                    "整列會有黃色底色，讓主管一眼找出還沒開始的任務。"
                ),
                "actions": [
                    "指向各欄位依序說明",
                    "指向一列黃色底色的列說明",
                    "指向維修工時欄的深藍粗體數字",
                ],
                "hints": [],
            },
            {
                "time": "02:00–02:30",
                "label": "切換 Tab & 分頁操作",
                "narration": (
                    "點「緊急事件」Tab，清單切換為緊急事件的記錄，"
                    "右側「共 N 筆」也同步更新為緊急事件的筆數。"
                    "\n\n"
                    "清單底部有分頁控制——預設每頁 50 筆，"
                    "右側可以選擇改為每頁 20 或 100 筆。"
                    "資料筆數多的時候，可以改為 100 筆減少翻頁次數。"
                    "\n\n"
                    "好，這集清單介紹完了。下一集看任務明細 Drawer 和操作情境。"
                ),
                "actions": [
                    "點「緊急事件」Tab",
                    "指向底部分頁控制和每頁筆數選單",
                ],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第三集",
        "title": "任務明細 Drawer & 常見操作情境",
        "duration": "建議時長：3–4 分鐘",
        "goal": "讓觀眾學會使用明細 Drawer 查看完整任務資訊（含附圖），並示範四個常用操作情境。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場",
                "narration": (
                    "最後一集，介紹點擊任務後展開的明細 Drawer，"
                    "以及幾個最常用的操作情境。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–02:00",
                "label": "明細 Drawer 完整介紹",
                "narration": (
                    "點擊清單任意一列，右側滑出任務明細 Drawer。"
                    "\n\n"
                    "Drawer 標題依格式顯示：類型 Tag + 任務類型 + 識別碼 + 「在 Ragic 查看」連結。"
                    "識別碼優先取「報修編號」，沒有的話取「問題說明」的前 20 個字。"
                    "「在 Ragic 查看」點擊在新分頁開啟 Ragic 原始記錄。"
                    "\n\n"
                    "Drawer 內容分三區："
                    "\n\n"
                    "第一區是「基本資訊」兩欄顯示："
                    "歸屬（飯店/商場 Tag）、屬性（上級交辦/緊急事件 Tag）、"
                    "交辦主管、工程人員、建立日期、最後更新、狀態（彩色 Tag）、維修工時。"
                    "\n\n"
                    "第二區是「明細資訊」單欄："
                    "「問題說明」完整顯示（粗體，保留所有換行）——"
                    "這裡才能看到在清單中被截斷的完整描述。"
                    "「備註」也是完整版本。"
                    "\n\n"
                    "第三區是「附圖」——"
                    "如果任務有上傳圖片，底部會顯示附圖區域，以三欄排列縮圖。"
                    "點擊任一縮圖可展開 Lightbox 全螢幕預覽，支援左右切換。"
                    "如果有附圖，Drawer 寬度會從 480px 自動加寬到 640px。"
                    "\n\n"
                    "圖片是系統從資料庫自動載入的，不需要手動操作，"
                    "載入中會顯示「載入附圖...」的提示。"
                ),
                "actions": [
                    "點清單第一列展開 Drawer",
                    "逐一指向三個資訊區塊",
                    "指向「在 Ragic 查看」連結",
                    "若有附圖，點縮圖示範 Lightbox",
                ],
                "hints": ["建議選有附圖的記錄示範，或提前確認哪幾筆有圖片"],
            },
            {
                "time": "02:00–02:35",
                "label": "情境 1：確認本月待辦的主管交辦件數",
                "narration": (
                    "情境一：月初確認有哪些主管交辦還沒開始執行。"
                    "\n\n"
                    "年份月份選本月，狀態篩選選「候辦」或「待辦」，"
                    "查看「上級交辦」Tab 的清單——黃色底色的列就是待辦任務。"
                    "Tab Badge 的數字告訴你本月共有幾件待辦，"
                    "同時也能看緊急事件 Tab 的待辦件數。"
                ),
                "actions": ["示範篩選本月待辦狀態"],
                "hints": [],
            },
            {
                "time": "02:35–03:10",
                "label": "情境 2：查詢某工程師的本月任務量",
                "narration": (
                    "情境二：主管想了解某位工程師這個月承接了多少任務。"
                    "\n\n"
                    "年月選本月，工程人員篩選選目標工程師，"
                    "兩個 Tab 的 Badge 就分別顯示他的上級交辦件數和緊急事件件數，"
                    "加總就是他的本月任務量。"
                    "\n\n"
                    "在清單中可以進一步確認各件任務的狀態，"
                    "了解哪些已結案、哪些還在處理中。"
                ),
                "actions": ["示範篩選特定工程人員"],
                "hints": [],
            },
            {
                "time": "03:10–03:45",
                "label": "結尾",
                "narration": (
                    "主管交辦／緊急事件的三集教學全部結束了。"
                    "\n\n"
                    "快速回顧：第一集學功能簡介和七個篩選條件操作，"
                    "第二集學清單欄位和 Tab 切換，"
                    "第三集學明細 Drawer 和操作情境。"
                    "\n\n"
                    "這個模組是追蹤非例行工務任務的重要工具，"
                    "善用篩選條件和即時更新的 Badge，"
                    "可以快速掌握待辦任務的整體狀況。謝謝收看！"
                ),
                "actions": ["畫面停在清單頁面做結尾"],
                "hints": [],
            },
        ],
    },
]


# ── Word 文件生成 ─────────────────────────────────────────────────────────────
def build_doc(scripts: list, output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height   = Inches(11.69); sec.page_width    = Inches(8.27)
    sec.top_margin    = Inches(0.85);  sec.bottom_margin = Inches(0.85)
    sec.left_margin   = Inches(1.1);   sec.right_margin  = Inches(1.1)

    # 封面
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_heading("主管交辦／緊急事件", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY; t.runs[0].font.size = Pt(26)
    t2 = doc.add_heading("影音教學旁白腳本", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT; t2.runs[0].font.size = Pt(20)
    doc.add_paragraph()
    sub = doc.add_paragraph(
        f"共三集 · 路由：/hotel/other-tasks（飯店）/ /mall/other-tasks（商場）\n"
        f"{datetime.now().strftime('%Y 年 %m 月 %d 日')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    doc.add_paragraph()

    # 使用說明
    doc.add_heading("腳本使用說明", 1).runs[0].font.color.rgb = PRIMARY
    for tip in [
        "🎬 旁白：錄製時直接念，可依個人習慣調整，不必逐字照念。",
        "🖱 操作：錄製時同步要做的滑鼠動作。",
        "⏸ 提示：錄製前的準備建議，不需要說出來。",
        "[時間軸]：僅供參考，依實際節奏調整。",
        "建議選有候辦/待辦狀態記錄的月份，讓黃色底色列清楚可見。",
        "若需展示附圖功能，提前確認哪幾筆記錄有附圖。",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(tip).font.size = Pt(10.5)

    doc.add_page_break()

    for script in scripts:
        h = doc.add_heading(f"{script['ep']}  {script['title']}", 1)
        h.runs[0].font.color.rgb = PRIMARY; h.runs[0].font.size = Pt(16)

        info = doc.add_paragraph()
        r1 = info.add_run(f"⏱ {script['duration']}　　")
        r1.font.size = Pt(10); r1.font.color.rgb = ORANGE; r1.bold = True

        gp = doc.add_paragraph()
        gp.paragraph_format.space_after = Pt(8)
        gp.add_run("學習目標：").bold = True
        gp.runs[0].font.color.rgb = ACCENT
        gp.add_run(script["goal"])

        section_divider(doc)

        for seg in script["segments"]:
            timing_row(doc, seg["time"], seg["label"])
            narration_block(doc, seg["narration"])
            for act in seg.get("actions", []):
                action_block(doc, act)
            for hint in seg.get("hints", []):
                hint_block(doc, hint)
            doc.add_paragraph()

        doc.add_page_break()

    # 附錄
    doc.add_heading("附錄：狀態顏色速查", 1).runs[0].font.color.rgb = PRIMARY
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["狀態", "Tag 顏色", "代表意義", "列底色"]):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1B3A5C")

    rows_data = [
        ("結案 / 已結案 / 已完成", "🟢 success 綠",   "已完成結案",   "—"),
        ("處理中 / 進行中",        "🔵 processing 藍","執行中",       "—"),
        ("候辦 / 待辦 / 待排程",   "🟡 warning 黃",   "尚未開始",     "🟡 黃色"),
        ("取消",                   "⬜ default 灰",   "已取消",       "—"),
    ]
    for status, color, desc, row_bg in rows_data:
        row = table.add_row()
        row.cells[0].text = status; row.cells[1].text = color
        row.cells[2].text = desc;   row.cells[3].text = row_bg
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 腳本手冊已儲存：{output_path}")


if __name__ == "__main__":
    build_doc(SCRIPTS, OUTPUT_PATH)
    print(f"執行完成！請開啟：{OUTPUT_PATH}")
