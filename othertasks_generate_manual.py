#!/usr/bin/env python3
"""
主管交辦／緊急事件（/hotel/other-tasks）教學手冊自動生成腳本
=============================================================
使用方式：
  pip install playwright python-docx
  playwright install chromium
  python othertasks_generate_manual.py

輸出：
  manual_screenshots_othertasks/
  主管交辦緊急事件教學手冊.docx
"""

import asyncio
import sys
from pathlib import Path
from datetime import datetime

try:
    from playwright.async_api import async_playwright
except ImportError:
    print("❌ 請先執行：pip install playwright && playwright install chromium"); sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx"); sys.exit(1)

# ── 設定 ──────────────────────────────────────────────────────────────────────
BASE_URL    = "http://localhost:5173"
TARGET_URL  = f"{BASE_URL}/hotel/other-tasks"
OUTPUT_DIR  = Path(__file__).parent / "manual_screenshots_othertasks"
MANUAL_PATH = Path(__file__).parent / "主管交辦緊急事件教學手冊.docx"
USERNAME    = "admin"
PASSWORD    = "Admin@2026"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
GREEN   = RGBColor(0x52, 0xC4, 0x1A)
ORANGE  = RGBColor(0xFA, 0xAD, 0x14)
RED     = RGBColor(0xCF, 0x13, 0x22)
GRAY    = RGBColor(0x59, 0x59, 0x59)
CONTENT_X = 185
CONTENT_W = 1390


# ── 截圖邏輯 ──────────────────────────────────────────────────────────────────
async def capture_screenshots(page, output_dir: Path) -> dict:
    shots = {}

    async def snap(name: str, scroll_y: int = 0, clip=None, wait_ms: int = 700):
        try:
            await page.evaluate(f"window.scrollTo(0, {scroll_y})")
            await page.wait_for_timeout(wait_ms)
            path = output_dir / f"{name}.png"
            await page.screenshot(path=str(path), full_page=False,
                                  **({"clip": clip} if clip else {}))
            shots[name] = path
            print(f"   ✓ {name}")
        except Exception as e:
            print(f"   ✗ {name}：{e}")

    async def click_tab(tab_text: str, wait_ms: int = 1200):
        try:
            tab = page.locator("div.ant-tabs-tab").filter(has_text=tab_text).first
            await tab.click()
            await page.wait_for_timeout(wait_ms)
        except Exception as e:
            print(f"   ⚠ Tab [{tab_text}] 失敗：{e}")

    # ── 登入 ──────────────────────────────────────────────────────────────────
    print("🔐 登入中...")
    await page.goto(f"{BASE_URL}/login")
    await page.wait_for_load_state("domcontentloaded")
    await page.wait_for_timeout(800)
    await page.locator("input[placeholder*='帳號'], input[type='text']").first.fill(USERNAME)
    await page.locator("input[type='password']").first.fill(PASSWORD)
    await page.locator("button[type='submit'], button:has-text('登入')").first.click()
    try:
        await page.wait_for_url(f"{BASE_URL}/dashboard", timeout=15000)
        print("   ✓ 登入成功")
    except Exception:
        print(f"   ✓ 已跳轉至 {page.url}")

    # ── 導航 ──────────────────────────────────────────────────────────────────
    await page.goto(TARGET_URL)
    await page.wait_for_load_state("networkidle")
    try:
        await page.locator(".ant-table-tbody tr").first.wait_for(timeout=12000)
        await page.wait_for_timeout(1500)
        print("   ✓ 頁面載入完成")
    except Exception:
        await page.wait_for_timeout(4000)

    # ── 1. 頁面標題 + 篩選列 ──────────────────────────────────────────────────
    print("📸 截圖：頁面標題 + 篩選列")
    await page.evaluate("window.scrollTo(0, 0)")
    await page.wait_for_timeout(500)
    await snap("01_title_filter",
               clip={"x": CONTENT_X, "y": 55, "width": CONTENT_W, "height": 230})

    # ── 2. 上級交辦 Tab 清單（預設）──────────────────────────────────────────
    print("📸 截圖：上級交辦 Tab")
    await snap("02_tab_supervisor",
               clip={"x": CONTENT_X, "y": 240, "width": CONTENT_W, "height": 560})

    # ── 3. 切換到緊急事件 Tab ─────────────────────────────────────────────────
    print("📸 截圖：緊急事件 Tab")
    await click_tab("緊急事件", wait_ms=1500)
    await snap("03_tab_emergency",
               clip={"x": CONTENT_X, "y": 240, "width": CONTENT_W, "height": 560})

    # ── 4. 切回上級交辦，點擊第一列展開 Drawer ────────────────────────────────
    print("📸 截圖：明細 Drawer")
    await click_tab("上級交辦", wait_ms=1200)
    try:
        first_row = page.locator(".ant-table-tbody tr").first
        await first_row.click()
        await page.wait_for_timeout(1500)
        await snap("04_drawer",
                   clip={"x": 850, "y": 55, "width": 750, "height": 720})
        # 關掉 Drawer
        close_btn = page.locator(".ant-drawer-close").first
        await close_btn.click()
        await page.wait_for_timeout(400)
    except Exception as e:
        print(f"   ✗ Drawer：{e}")

    # ── 5. 篩選器操作示意（選一個狀態）──────────────────────────────────────
    print("📸 截圖：篩選器（展示各選單）")
    await page.evaluate("window.scrollTo(0, 0)")
    await page.wait_for_timeout(400)
    await snap("05_filter_overview",
               clip={"x": CONTENT_X, "y": 55, "width": CONTENT_W, "height": 150})

    return shots


# ── Word 手冊輔助函式 ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def add_image_safe(doc, path, width_inches: float = 5.8):
    if isinstance(path, Path) and path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph("[ 截圖未能取得，請參考系統實際畫面 ]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = GRAY; p.runs[0].font.italic = True


def add_tip(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run("💡 提示："); r.bold = True; r.font.color.rgb = PRIMARY
    p.add_run(text)


def hdr_table(doc, headers, bg="1B3A5C"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, bg)
    return table


def h1(doc, t):
    h = doc.add_heading(t, 1); h.runs[0].font.color.rgb = PRIMARY; return h

def h2(doc, t):
    h = doc.add_heading(t, 2); h.runs[0].font.color.rgb = ACCENT; return h


# ── Word 手冊建構 ─────────────────────────────────────────────────────────────
def build_word_manual(shots: dict, output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11.69); sec.page_width  = Inches(8.27)
    sec.top_margin  = Inches(0.9);   sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1.1);   sec.right_margin  = Inches(1.1)

    # ── 封面 ──────────────────────────────────────────────────────────────────
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_heading("主管交辦／緊急事件", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY; t.runs[0].font.size = Pt(28)
    t2 = doc.add_heading("教學操作手冊", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT; t2.runs[0].font.size = Pt(22)
    doc.add_paragraph()
    sub = doc.add_paragraph(
        f"路由：/hotel/other-tasks（飯店）/ /mall/other-tasks（商場）\n"
        f"{datetime.now().strftime('%Y 年 %m 月 %d 日')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    doc.add_page_break()

    # ── 一、功能簡介 ──────────────────────────────────────────────────────────
    h1(doc, "一、功能簡介")
    p = doc.add_paragraph()
    p.add_run("主管交辦／緊急事件").bold = True
    p.add_run(
        "（飯店：/hotel/other-tasks；商場：/mall/other-tasks）"
        "管理主管臨時交辦的工務任務與飯店/商場突發的緊急事件記錄。"
        "資料來源為 Ragic other-tasks 模組，以建立日期作為月份歸屬。"
        "\n\n"
        "頁面進入時依路徑自動帶入歸屬篩選："
        "從「飯店管理」進入預設顯示飯店任務；從「商場管理」進入預設顯示商場任務。"
    )
    doc.add_paragraph()
    doc.add_paragraph("兩個頁籤內容結構完全相同，僅任務類型不同：")
    for tab, color, desc in [
        ("🔔 上級交辦", "深藍", "主管指派的工務任務，通常為非例行性工作"),
        ("🚨 緊急事件", "深紅", "突發性的緊急情況，需即時處理"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"• {tab}（{color}）：").bold = True
        p.add_run(desc)
    doc.add_page_break()

    # ── 二、頁面標題 & 篩選列 ─────────────────────────────────────────────────
    h1(doc, "二、頁面標題 & 篩選列")
    add_image_safe(doc, shots.get("01_title_filter"))
    doc.add_paragraph()

    h2(doc, "2.1  篩選條件（7 個）")
    filter_rows = [
        ("年份",     "篩選特定年份，可清空（顯示所有年份）"),
        ("月份",     "篩選特定月份（1–12），可清空（顯示全年）"),
        ("狀態",     "篩選案件狀態（動態從 Ragic 取得選項）"),
        ("交辦主管", "篩選指派主管姓名（動態選項）"),
        ("工程人員", "篩選負責執行的工程師（動態選項）"),
        ("歸屬",     "篩選飯店 / 商場（從 hotel 路徑進入預設「飯店」）"),
        ("關鍵字",   "在問題說明與備註欄位中搜尋關鍵字"),
    ]
    table = hdr_table(doc, ["篩選條件", "說明"])
    for name, desc in filter_rows:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    add_tip(doc, "所有篩選條件即時生效，更改後清單和 Tab 件數 Badge 自動更新，無需按查詢按鈕。")
    doc.add_paragraph()

    h2(doc, "2.2  Tab 件數 Badge")
    doc.add_paragraph(
        "每個 Tab 標題旁顯示符合目前篩選條件的件數 Badge（藍色為上級交辦，紅色為緊急事件），"
        "讓管理者在切換 Tab 之前就能知道各類型的待辦數量。"
        "篩選條件改變時 Badge 自動更新。"
    )
    doc.add_page_break()

    # ── 三、任務清單 ──────────────────────────────────────────────────────────
    h1(doc, "三、任務清單")

    h2(doc, "3.1  上級交辦 Tab")
    add_image_safe(doc, shots.get("02_tab_supervisor"))
    doc.add_paragraph()

    h2(doc, "3.2  緊急事件 Tab")
    add_image_safe(doc, shots.get("03_tab_emergency"))
    doc.add_paragraph()

    h2(doc, "3.3  表格欄位說明")
    col_rows = [
        ("歸屬",        "飯店（藍色 Tag）/ 商場（綠色 Tag）"),
        ("建立日期",    "任務或事件的建立時間（yyyy/mm/dd HH:mm）"),
        ("交辦主管",    "指派此任務的主管姓名"),
        ("工程人員",    "負責執行的工程師姓名"),
        ("問題說明",    "任務或事件的詳細描述（超出寬度以「...」截斷）"),
        ("備註",        "補充說明或處理備注（超出以「...」截斷）"),
        ("最後更新日期","最後一次修改記錄的時間"),
        ("狀態",        "彩色 Tag 標示目前狀態（見附錄說明）"),
        ("維修工時",    "工時（小時，精確到 2 位小數）；為 0 時顯示「—」"),
    ]
    table = hdr_table(doc, ["欄位", "說明"])
    for name, desc in col_rows:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph(
        "候辦 / 待辦 / 待排程 狀態的列會以黃色底色標示，便於快速識別尚未開始執行的任務。"
    )
    doc.add_paragraph()

    h2(doc, "3.4  分頁設定")
    doc.add_paragraph("• 預設每頁顯示 50 筆，可切換為 20 或 100 筆")
    doc.add_paragraph("• 右上角顯示「共 N 筆」符合篩選條件的總件數")
    doc.add_paragraph("• 支援頁碼點擊翻頁")
    doc.add_page_break()

    # ── 四、任務明細 Drawer ───────────────────────────────────────────────────
    h1(doc, "四、任務明細 Drawer")
    add_image_safe(doc, shots.get("04_drawer"), width_inches=4.5)
    doc.add_paragraph()
    doc.add_paragraph(
        "點擊清單任意一列，右側滑出「任務明細 Drawer」。"
        "Drawer 寬度：有附圖時 640px，無附圖時 480px。"
    )
    doc.add_paragraph()

    h2(doc, "4.1  標題列格式")
    doc.add_paragraph(
        "Drawer 標題依規範格式顯示：\n"
        "[類型 Tag]  [任務類型]：[問題說明前 20 字]  [🔗 在 Ragic 查看]"
    )
    doc.add_paragraph()

    h2(doc, "4.2  基本資訊區（2 欄）")
    basic_rows = [
        ("歸屬",       "飯店（藍色）/ 商場（綠色）Tag"),
        ("屬性",       "上級交辦（深藍）/ 緊急事件（深紅）Tag"),
        ("交辦主管",   "指派主管姓名"),
        ("工程人員",   "執行工程師姓名"),
        ("建立日期",   "任務建立時間"),
        ("最後更新",   "最後修改時間"),
        ("狀態",       "彩色 Tag"),
        ("維修工時",   "hr（若有）"),
    ]
    table = hdr_table(doc, ["欄位", "說明"])
    for name, desc in basic_rows:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    h2(doc, "4.3  明細資訊區（1 欄）")
    doc.add_paragraph("• 問題說明：完整顯示（粗體、保留換行）")
    doc.add_paragraph("• 備註：完整說明文字")
    doc.add_paragraph()

    h2(doc, "4.4  附圖預覽")
    doc.add_paragraph(
        "若任務有上傳圖片，Drawer 底部顯示附圖區域（3 欄排列），"
        "點擊縮圖可展開 Lightbox 全螢幕預覽（Image.PreviewGroup 支援左右切換）。"
        "系統自動從資料庫載入圖片，載入中顯示「載入附圖...」轉圈提示。"
    )
    doc.add_paragraph()
    add_tip(doc, "標題列的「在 Ragic 查看」連結點擊後在新分頁開啟 Ragic 原始記錄，可查看完整欄位與操作記錄。")
    doc.add_page_break()

    # ── 五、常見操作情境 ──────────────────────────────────────────────────────
    h1(doc, "五、常見操作情境")
    scenarios = [
        ("確認本月有哪些未完成的主管交辦",
         ["篩選列選定本月年月",
          "狀態篩選選「候辦」或「待辦」",
          "查看上級交辦 Tab 的清單，黃色底色列即為待辦任務",
          "點擊任務展開 Drawer 確認詳情"]),
        ("查詢特定工程師本月承接的所有任務",
         ["工程人員篩選選目標工程師",
          "月份篩選選本月",
          "查看上級交辦 + 緊急事件兩個 Tab 的清單",
          "Tab Badge 數字即為各類型的任務件數"]),
        ("確認某次緊急事件的處理記錄",
         ["切換到「緊急事件」Tab",
          "關鍵字搜尋輸入事件相關描述關鍵字",
          "找到目標記錄後點擊展開 Drawer",
          "查看問題說明（完整版）、工程人員、工時和狀態",
          "若需查看原始資料，點「在 Ragic 查看」連結"]),
        ("查看商場的緊急事件（從飯店路徑進入）",
         ["歸屬篩選從「飯店」改為「商場」",
          "切換到「緊急事件」Tab",
          "清單自動更新為商場的緊急事件記錄"]),
    ]
    for i, (title, steps) in enumerate(scenarios, 1):
        p = doc.add_paragraph()
        p.add_run(f"情境 {i}：{title}").bold = True
        p.runs[0].font.color.rgb = PRIMARY
        for step in steps:
            doc.add_paragraph(f"　→ {step}")
        doc.add_paragraph()

    # ── 附錄 ──────────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "附錄：狀態顏色對照表")
    table = hdr_table(doc, ["狀態", "Tag 顏色", "代表意義", "列底色"])
    status_rows = [
        ("結案 / 已結案 / 已完成 / 完成", "🟢 success 綠色",   "任務已完成結案",   "—"),
        ("處理中 / 進行中",               "🔵 processing 藍色","執行中",           "—"),
        ("候辦 / 待辦 / 待排程",          "🟡 warning 黃色",   "尚未開始執行",     "🟡 黃色底色"),
        ("取消",                          "⬜ default 灰色",   "任務已取消",       "—"),
    ]
    for status, color, desc, row_bg in status_rows:
        row = table.add_row()
        row.cells[0].text = status; row.cells[1].text = color
        row.cells[2].text = desc;   row.cells[3].text = row_bg
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 手冊已儲存：{output_path}")


# ── 主程式 ────────────────────────────────────────────────────────────────────
async def main():
    OUTPUT_DIR.mkdir(exist_ok=True)
    print("=" * 55)
    print("  主管交辦／緊急事件 教學手冊生成器")
    print("=" * 55)
    print(f"目標：{TARGET_URL}")
    print(f"截圖：{OUTPUT_DIR}")
    print(f"手冊：{MANUAL_PATH}\n")

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=["--disable-blink-features=AutomationControlled"],
        )
        context = await browser.new_context(
            viewport={"width": 1600, "height": 900},
            device_scale_factor=1.5,
            locale="zh-TW",
            timezone_id="Asia/Taipei",
        )
        page = await context.new_page()
        shots = await capture_screenshots(page, OUTPUT_DIR)
        await browser.close()

    print(f"\n截圖完成，共 {len(shots)} 張")
    print("\n📄 生成 Word 手冊...")
    build_word_manual(shots, MANUAL_PATH)
    print(f"\n✅ 執行完成！請開啟：{MANUAL_PATH}")


if __name__ == "__main__":
    asyncio.run(main())
