#!/usr/bin/env python3
"""
全棟例行維護（/mall/full-building-maintenance）影音教學旁白腳本生成器
=================================================================
使用方式：
  python fullbldg_pm_generate_video_scripts.py

輸出：全棟例行維護_影音教學腳本.docx

本腳本以「與商場週期保養表的差異」為核心，
熟悉商場版的觀眾可以只看第一集和第二集。
"""

import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx"); sys.exit(1)

OUTPUT_PATH = Path(__file__).parent / "全棟例行維護_影音教學腳本.docx"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
ORANGE  = RGBColor(0xFA, 0xAD, 0x14)
GRAY    = RGBColor(0x59, 0x59, 0x59)


# ── 輔助函式 ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def narration_block(doc, text: str):
    for part in text.split("\n\n"):
        part = part.strip()
        if not part:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(part); r.font.size = Pt(11.5)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "12"); left.set(qn("w:color"), "4BA8E8")
        pBdr.append(left); pPr.append(pBdr)


def label_box(doc, label: str, text: str, fill_hex: str):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    lc = table.rows[0].cells[0]
    set_cell_bg(lc, fill_hex)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label); lr.bold = True; lr.font.size = Pt(10)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = table.rows[0].cells[1]
    set_cell_bg(rc, "F8F9FA")
    rc.paragraphs[0].add_run(text).font.size = Pt(10.5)
    doc.add_paragraph()


def action_block(doc, text: str):  label_box(doc, "🖱 操作", text, "52C41A")
def hint_block(doc, text: str):    label_box(doc, "⏸ 提示", text, "FA8C16")
def diff_block(doc, text: str):    label_box(doc, "⭐ 全棟特有", text, "1565C0")
def warn_block(doc, text: str):    label_box(doc, "⚠ 注意", text, "C0392B")


def timing_row(doc, t_range: str, label: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"[{t_range}]  ")
    r1.bold = True; r1.font.color.rgb = ORANGE; r1.font.size = Pt(10)
    r2 = p.add_run(label)
    r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = PRIMARY


def section_divider(doc):
    p = doc.add_paragraph("─" * 64)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── 腳本資料 ──────────────────────────────────────────────────────────────────
SCRIPTS = [

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第一集",
        "title": "功能簡介（與商場版差異）& Dashboard KPI 六卡",
        "duration": "建議時長：4 分鐘",
        "goal": "讓觀眾快速掌握全棟例行維護與商場週期保養表的三點差異，並了解多出的「保養時間」KPI 卡的意義。",
        "segments": [
            {
                "time": "00:00–00:45",
                "label": "功能簡介與差異說明",
                "narration": (
                    "大家好，歡迎收看全棟例行維護教學系列。"
                    "這個模組管理整棟大樓的例行維護保養，"
                    "和商場週期保養表的架構幾乎完全相同——"
                    "如果你已經看過商場週期保養表的教學，絕大多數的操作方式都一樣。"
                    "\n\n"
                    "全棟版和商場版有三個差異，我先說清楚。"
                    "\n\n"
                    "第一：Dashboard KPI 是六張，商場版是五張——"
                    "全棟版多了一張「保養時間」，顯示實際完成的保養工時。"
                    "\n\n"
                    "第二：每日巡檢表目前標示「本地同步功能開發中」，"
                    "篩選器也簡化為只有單日 DatePicker，沒有年月選單。"
                    "\n\n"
                    "第三：年度計劃表的格子多了兩種狀態——"
                    "粉紅色的「！」代表應保養但無資料，"
                    "淺灰色的「∅」代表這個月沒有設定保養頻率。"
                    "\n\n"
                    "其他六個 Tab 的功能和商場版完全相同，不再重複說明。"
                    "這集先介紹差異最大的 Dashboard KPI 六卡。"
                ),
                "actions": ["進入頁面，指向八個頁籤"],
                "hints": [],
            },
            {
                "time": "00:45–01:15",
                "label": "年月篩選與本月批次入口",
                "narration": (
                    "和商場版一樣，Dashboard 左上角有年份和月份的篩選選單，"
                    "切換後點「重新整理」更新資料。"
                    "頁面底部若有本月批次，會顯示快速入口直接跳轉。"
                ),
                "actions": ["指向年月篩選"],
                "hints": [],
            },
            {
                "time": "01:15–03:30",
                "label": "KPI 六卡詳解",
                "narration": (
                    "KPI 卡片有六張，前四張和商場版相同，後兩張是關鍵。"
                    "\n\n"
                    "第一張「有效項目」深藍色——當月應執行的保養工項總數。"
                    "\n\n"
                    "第二張「已完成」綠色——完成件數加上括弧裡的完成率百分比。"
                    "\n\n"
                    "第三張「逾期件數」深紅色——過了排定日期還沒完成的件數。"
                    "\n\n"
                    "第四張「異常待追蹤」紫色——有異常記錄需追蹤的件數。"
                    "\n\n"
                    "第五張「預估工時」藍色——依保養計劃加總的計劃工時，"
                    "也就是按照原本排程預估這個月應該要花多少小時。"
                    "\n\n"
                    "第六張「保養時間」綠色——這是全棟版獨有的！"
                    "顯示實際完成保養的工時，對應 Ragic 記錄的實際作業時間。"
                    "\n\n"
                    "比較「預估工時」和「保養時間」這兩張很有意義——"
                    "如果保養時間明顯低於預估工時，可能代表有些項目還沒做完；"
                    "如果超過，代表實際工作量比預期重，值得關注原因。"
                    "\n\n"
                    "下方一樣有完成率進度條，格式是「百分比（完成/總件數）」。"
                ),
                "actions": [
                    "放大顯示六張卡片，依序指向",
                    "特別強調第六張「保養時間」為全棟版獨有",
                    "指向第五和第六張說明比較意義",
                ],
                "hints": [],
            },
            {
                "time": "03:30–04:00",
                "label": "圖表與預警區",
                "narration": (
                    "圖表區和預警區和商場版完全相同——"
                    "水平堆疊長條圖、狀態環狀圖、逾期 Top 10、待執行項目。"
                    "詳細說明請參考商場週期保養表教學的第一集。"
                    "\n\n"
                    "下一集介紹全棟版每日巡檢表的開發中狀態。"
                ),
                "actions": ["快速掃過圖表和預警區"],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第二集",
        "title": "每日巡檢表（開發中）& 年度計劃表（多兩種狀態）",
        "duration": "建議時長：3–4 分鐘",
        "goal": "讓觀眾了解全棟版每日巡檢表的開發中狀態與限制，以及年度計劃表新增的兩種格子狀態。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "這集介紹全棟版的兩個差異點——"
                    "每日巡檢表的開發中狀態，以及年度計劃表新增的格子狀態。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:40",
                "label": "每日巡檢表（開發中）",
                "narration": (
                    "切換到「每日巡檢表」Tab。"
                    "\n\n"
                    "頁面最上方有一個藍色的提示橫幅，說明目前狀態：「整棟巡檢本地同步功能開發中。"
                    "目前每日巡檢表尚未接通本地 DB，模板欄位已備妥。"
                    "請至 Ragic 系統填寫各樓層巡檢表單，接通後資料將自動顯示於此。」"
                    "\n\n"
                    "這個提示的意思是：系統的欄位框架已經建好了，"
                    "但目前還無法從本地資料庫讀取整棟巡檢的資料。"
                    "正式的巡檢記錄請繼續在 Ragic 系統填寫，等功能完成後這裡會自動顯示。"
                    "\n\n"
                    "篩選器和商場版不同——全棟版只有一個單日 DatePicker，"
                    "沒有年份和月份選單。選擇日期後點「重新整理」，"
                    "系統會嘗試取得當日的巡檢資料。"
                    "\n\n"
                    "表格欄位和商場版相同：樓層、項目、檢查內容、運轉狀況、"
                    "實際巡檢人員、異常說明、時間。"
                    "目前若無接通資料，表格會顯示「尚無巡檢資料」。"
                ),
                "actions": [
                    "切換到「每日巡檢表」Tab",
                    "指向藍色提示橫幅說明",
                    "指向單日 DatePicker（比較商場版的年月選單）",
                ],
                "hints": ["如果目前沒有資料，表格為空，告訴觀眾這是正常的開發中狀態"],
            },
            {
                "time": "01:40–03:00",
                "label": "年度計劃表（多兩種格子狀態）",
                "narration": (
                    "切換到「年度計劃表」Tab。"
                    "\n\n"
                    "全棟版的年度計劃表和商場版大致相同，"
                    "但多了兩種格子狀態，我來說明。"
                    "\n\n"
                    "第一種新狀態：粉紅底色的「！」圖示，狀態碼是 no_data。"
                    "代表這個保養項目在這個月份「應該保養，但系統找不到任何資料」。"
                    "可能的原因是排程沒有建立，或者 Ragic 的記錄還沒同步過來。"
                    "看到「！」的格子，點擊後 Drawer 會提示你去「排程管理」Tab 建立排程，"
                    "或者前往 Ragic 確認記錄狀態。"
                    "\n\n"
                    "第二種新狀態：淺灰底色的「∅」圖示，狀態碼是 no_frequency。"
                    "代表這個保養項目在系統設定裡「沒有設定這個月份的執行頻率」。"
                    "也就是說，這個月本來就不需要保養，是正常的。"
                    "\n\n"
                    "加上商場版原有的六種狀態，全棟版的年度計劃表共有八種格子狀態。"
                    "掌握「！」和「∅」的差異，就能正確判斷格子代表的情況。"
                    "\n\n"
                    "其他的格子操作——點格子展開 Drawer 明細——和商場版完全相同。"
                ),
                "actions": [
                    "切換到「年度計劃表」Tab",
                    "指向「！」狀態的格子（若有）說明意義",
                    "指向「∅」狀態的格子說明",
                    "點一個有資料的格子展示 Drawer",
                ],
                "hints": ["若沒有「！」或「∅」的格子，只要指向圖例說明即可"],
            },
            {
                "time": "03:00–03:30",
                "label": "其他 Tab 簡介",
                "narration": (
                    "剩下的 Tab——每月維護、每季維護、每年維護、排程管理、批次清單——"
                    "和商場週期保養表完全相同，這邊就不再重複說明了。"
                    "\n\n"
                    "記住月初要執行「排程管理」Tab 的「產生排程」，"
                    "為當月應執行的全棟保養項目建立排程記錄。"
                    "\n\n"
                    "全棟例行維護的兩集教學到這邊全部結束。"
                    "核心差異記住三點：多一張保養時間 KPI 卡、"
                    "每日巡檢表開發中、年度計劃表多兩種狀態。謝謝收看！"
                ),
                "actions": ["快速切換幾個 Tab 讓觀眾看到是同樣的界面"],
                "hints": [],
            },
        ],
    },
]


# ── Word 文件生成 ─────────────────────────────────────────────────────────────
def build_doc(scripts: list, output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height   = Inches(11.69); sec.page_width    = Inches(8.27)
    sec.top_margin    = Inches(0.85);  sec.bottom_margin = Inches(0.85)
    sec.left_margin   = Inches(1.1);   sec.right_margin  = Inches(1.1)

    # 封面
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_heading("全棟例行維護", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY; t.runs[0].font.size = Pt(26)
    t2 = doc.add_heading("影音教學旁白腳本", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT; t2.runs[0].font.size = Pt(20)
    doc.add_paragraph()
    sub = doc.add_paragraph(
        f"共二集 · 路由：/mall/full-building-maintenance\n{datetime.now().strftime('%Y 年 %m 月 %d 日')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    doc.add_paragraph()

    # 使用說明
    doc.add_heading("腳本使用說明", 1).runs[0].font.color.rgb = PRIMARY
    for tip in [
        "🎬 旁白：錄製時直接念，可依個人習慣調整，不必逐字照念。",
        "🖱 操作：錄製時同步要做的滑鼠動作。",
        "⭐ 全棟特有：藍色標示框代表全棟版與商場週期保養表的差異。",
        "⚠ 注意：紅色標示框代表需要特別提醒的事項。",
        "此腳本刻意簡短（只有兩集），假設觀眾已看過商場週期保養表教學。",
        "若觀眾未看過商場版，建議先引導至商場週期保養表教學系列（五集）。",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(tip).font.size = Pt(10.5)

    doc.add_page_break()

    for script in scripts:
        h = doc.add_heading(f"{script['ep']}  {script['title']}", 1)
        h.runs[0].font.color.rgb = PRIMARY; h.runs[0].font.size = Pt(16)

        info = doc.add_paragraph()
        r1 = info.add_run(f"⏱ {script['duration']}　　")
        r1.font.size = Pt(10); r1.font.color.rgb = ORANGE; r1.bold = True

        gp = doc.add_paragraph()
        gp.paragraph_format.space_after = Pt(8)
        gp.add_run("學習目標：").bold = True
        gp.runs[0].font.color.rgb = ACCENT
        gp.add_run(script["goal"])

        section_divider(doc)

        for seg in script["segments"]:
            timing_row(doc, seg["time"], seg["label"])
            narration_block(doc, seg["narration"])
            for act in seg.get("actions", []):
                action_block(doc, act)
            for hint in seg.get("hints", []):
                hint_block(doc, hint)
            doc.add_paragraph()

        doc.add_page_break()

    # 附錄：差異速查
    doc.add_heading("附錄：全棟版 vs 商場版差異速查", 1).runs[0].font.color.rgb = PRIMARY
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["差異點", "商場週期保養表", "全棟例行維護"]):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1B3A5C")
    diff_rows = [
        ("路由",               "/mall/periodic-maintenance", "/mall/full-building-maintenance"),
        ("Dashboard KPI 數量", "5 張",                       "6 張（多「保養時間」）"),
        ("每日巡檢表篩選",     "年月 + 單日 DatePicker",     "只有單日 DatePicker"),
        ("每日巡檢表狀態",     "已接通，有完整資料",          "開發中，藍色 Alert 提示"),
        ("年度計劃表格子狀態", "6 種",                        "8 種（+！no_data，+∅ no_frequency）"),
    ]
    for diff, mall, full in diff_rows:
        row = table.add_row()
        row.cells[0].text = diff; row.cells[1].text = mall; row.cells[2].text = full
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    doc.add_heading("年度計劃表格子狀態（全棟版 8 種）", 2).runs[0].font.color.rgb = ACCENT
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Table Grid"
    for i, h in enumerate(["圖示", "說明", "格子顏色", "全棟獨有？"]):
        c = table2.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1B3A5C")
    annual_rows = [
        ("✅", "已完成",       "綠底 #f6ffed",  "—"),
        ("🔴", "逾期",         "紅底 #fff1f0",  "—"),
        ("🔵", "進行中",       "藍底 #e6f4ff",  "—"),
        ("⭕", "待執行",       "橙底 #fff7e6",  "—"),
        ("?",  "未排定",       "黃底 #fffbe6",  "—"),
        ("─",  "非本月",       "淺灰 #fafafa",  "—"),
        ("！", "應保養無資料", "粉紅 #fff0f6",  "⭐ 全棟版獨有"),
        ("∅",  "無此月頻率",   "灰白 #f5f5f5",  "⭐ 全棟版獨有"),
    ]
    for icon, desc, color, note in annual_rows:
        row = table2.add_row()
        row.cells[0].text = icon; row.cells[1].text = desc
        row.cells[2].text = color; row.cells[3].text = note
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 腳本手冊已儲存：{output_path}")


if __name__ == "__main__":
    build_doc(SCRIPTS, OUTPUT_PATH)
    print(f"執行完成！請開啟：{OUTPUT_PATH}")
