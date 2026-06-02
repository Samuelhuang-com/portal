#!/usr/bin/env python3
"""
集團決策 Dashboard 教學手冊自動生成腳本
=====================================
使用方式：
  pip install playwright python-docx Pillow
  playwright install chromium
  python generate_dashboard_manual.py

輸出：集團決策Dashboard教學手冊.docx（存在腳本同層目錄）
"""

import asyncio
import sys
from pathlib import Path
from datetime import datetime

# ── 相依性檢查 ────────────────────────────────────────────────────────────────
try:
    from playwright.async_api import async_playwright
except ImportError:
    print("❌ 缺少 playwright。請執行：")
    print("   pip install playwright && playwright install chromium")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 缺少 python-docx。請執行：")
    print("   pip install python-docx")
    sys.exit(1)

# ── 設定 ──────────────────────────────────────────────────────────────────────
BASE_URL    = "http://localhost:5173"
OUTPUT_DIR  = Path(__file__).parent / "manual_screenshots"
MANUAL_PATH = Path(__file__).parent / "集團決策Dashboard教學手冊.docx"

USERNAME = "admin"
PASSWORD = "Admin@2026"


# ── 截圖邏輯 ──────────────────────────────────────────────────────────────────
async def capture_screenshots(page, output_dir: Path) -> dict:
    """
    導航到各頁面並截圖，回傳 {name: Path} 字典。
    若某張截圖失敗則跳過（手冊仍會生成）。
    """
    shots = {}

    async def snap(name: str, scroll_y: int = 0, clip=None, wait_ms: int = 800):
        try:
            if scroll_y > 0:
                await page.evaluate(f"window.scrollTo(0, {scroll_y})")
                await page.wait_for_timeout(wait_ms)
            path = output_dir / f"{name}.png"
            kwargs = {}
            if clip:
                kwargs["clip"] = clip
            await page.screenshot(path=str(path), full_page=False, **kwargs)
            shots[name] = path
            print(f"   ✓ {name}")
        except Exception as e:
            print(f"   ✗ {name}：{e}")

    # ── 登入 ──────────────────────────────────────────────────────────────────
    print("🔐 登入中...")
    await page.goto(f"{BASE_URL}/login")
    await page.wait_for_load_state("domcontentloaded")
    await page.wait_for_timeout(800)

    # 填入帳號
    await page.locator("input[placeholder*='帳號'], input[placeholder*='Email'], input[type='text']").first.fill(USERNAME)
    # 填入密碼
    await page.locator("input[type='password']").first.fill(PASSWORD)
    # 點登入按鈕
    await page.locator("button[type='submit'], button:has-text('登入')").first.click()

    # 等待跳轉到 dashboard（最多 15 秒）
    try:
        await page.wait_for_url(f"{BASE_URL}/dashboard", timeout=15000)
        print("   ✓ 登入成功")
    except Exception:
        # 若 URL 不完全符合，只要離開 /login 就算成功
        current = page.url
        if "/login" not in current:
            print(f"   ✓ 登入成功（跳轉至 {current}）")
        else:
            print("   ⚠ 登入可能失敗，繼續嘗試...")

    # ── 導航到 exec-work-dashboard ────────────────────────────────────────────
    await page.goto(f"{BASE_URL}/exec-work-dashboard")
    await page.wait_for_load_state("networkidle")
    # 等待 KPI 卡片出現（確認資料已載入）
    try:
        await page.locator(".ant-statistic-content-value").first.wait_for(timeout=12000)
        await page.wait_for_timeout(1500)   # 讓圖表渲染完畢
        print("   ✓ 資料載入完成")
    except Exception:
        await page.wait_for_timeout(5000)   # fallback 等待

    print("📸 截圖：集團工務概覽")

    # 1. 頁頭 + 三列 KPI 卡片（scroll=0，clip 去掉左側 sidebar）
    await page.evaluate("window.scrollTo(0, 0)")
    await page.wait_for_timeout(600)
    await snap("01_kpi_header",
               clip={"x": 185, "y": 60, "width": 1390, "height": 430})

    # 2. 工務報修摘要（年月篩選 + 三欄卡片）
    await page.evaluate("window.scrollTo(0, 430)")
    await page.wait_for_timeout(600)
    await snap("02_repair_cards",
               clip={"x": 185, "y": 430, "width": 1390, "height": 510})

    # 3. 近 12 個月趨勢圖 + 圓餅圖（飯店）
    await page.evaluate("window.scrollTo(0, 940)")
    await page.wait_for_timeout(800)
    await snap("03_charts_hotel",
               clip={"x": 185, "y": 840, "width": 1390, "height": 480})

    # 4. 近 12 個月趨勢圖 + 圓餅圖（商場）
    await page.evaluate("window.scrollTo(0, 1420)")
    await page.wait_for_timeout(600)
    await snap("04_charts_mall",
               clip={"x": 185, "y": 1320, "width": 1390, "height": 480})

    # 5. 展開所有 Collapse，截圖每日累計表
    await page.evaluate("window.scrollTo(0, 1800)")
    await page.wait_for_timeout(500)
    try:
        btn = page.locator("button:has-text('全展開')")
        if await btn.count() > 0:
            await btn.first.click()
            await page.wait_for_timeout(1200)
    except Exception:
        pass
    await page.evaluate("window.scrollTo(0, 1900)")
    await page.wait_for_timeout(600)
    await snap("05_daily_table",
               clip={"x": 185, "y": 1800, "width": 1390, "height": 520})

    # 6. 明細分析：每日/每月工時表（展開後繼續往下滾）
    await page.evaluate("window.scrollTo(0, 2600)")
    await page.wait_for_timeout(600)
    await snap("06_hours_tables",
               clip={"x": 185, "y": 2500, "width": 1390, "height": 500})

    # 7. 人員負荷 + 飯店vs商場比較表
    await page.evaluate("window.scrollTo(0, 3300)")
    await page.wait_for_timeout(600)
    await snap("07_burden_compare",
               clip={"x": 185, "y": 3200, "width": 1390, "height": 500})

    # 8. 工項矩陣 + 異常提醒
    await page.evaluate("window.scrollTo(0, 4000)")
    await page.wait_for_timeout(600)
    await snap("08_matrix_alerts",
               clip={"x": 185, "y": 3900, "width": 1390, "height": 500})

    # ── 工作日誌 TAB ──────────────────────────────────────────────────────────
    print("📸 截圖：工作日誌")
    try:
        await page.evaluate("window.scrollTo(0, 0)")
        tab = page.locator("div.ant-tabs-tab", has_text="工作日誌").first
        await tab.click()
        await page.wait_for_timeout(1500)
        await snap("09_journal_overview",
                   clip={"x": 185, "y": 55, "width": 1390, "height": 680})

        # 示範：切換到區間模式並查詢
        range_btn = page.locator("div.ant-segmented-item", has_text="區間").first
        await range_btn.click()
        await page.wait_for_timeout(400)
        await snap("10_journal_range_mode",
                   clip={"x": 185, "y": 55, "width": 1390, "height": 300})
    except Exception as e:
        print(f"   ✗ 工作日誌截圖失敗：{e}")

    # ── 統計基準說明 TAB ───────────────────────────────────────────────────────
    print("📸 截圖：統計基準說明")
    try:
        await page.evaluate("window.scrollTo(0, 0)")
        tab = page.locator("div.ant-tabs-tab", has_text="統計基準說明").first
        await tab.click()
        await page.wait_for_timeout(2500)   # iframe 需要較長載入
        await snap("11_methodology",
                   clip={"x": 185, "y": 55, "width": 1390, "height": 680})
    except Exception as e:
        print(f"   ✗ 統計基準說明截圖失敗：{e}")

    return shots


# ── Word 手冊建構 ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """為表格儲存格設定背景色"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_image_safe(doc, path, width_inches: float = 5.8):
    """插入圖片，若檔案不存在則插入提示文字"""
    if isinstance(path, Path) and path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph("[ 截圖未能取得，請參考系統實際畫面 ]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p.runs[0].font.italic = True


def add_tip(doc, text: str):
    """插入操作提示段落"""
    p = doc.add_paragraph()
    r = p.add_run("💡 操作提示：")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    p.add_run(text)
    return p


def build_word_manual(shots: dict, output_path: Path):
    doc = Document()

    # ── 頁面設定（A4）──────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width  = Inches(8.27)
    section.top_margin  = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

    PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
    ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
    GRAY    = RGBColor(0x59, 0x59, 0x59)

    # ────────────────────────────────────────────────────────────────────────
    # 封面
    # ────────────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_heading("集團決策 Dashboard", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY
    t.runs[0].font.size = Pt(28)

    t2 = doc.add_heading("教學操作手冊", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT
    t2.runs[0].font.size = Pt(22)

    doc.add_paragraph()
    sub = doc.add_paragraph(f"集團工務決策駕駛艙 · 操作指南\n{datetime.now().strftime('%Y 年 %m 月 %d 日')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────────
    # 目錄（手動）
    # ────────────────────────────────────────────────────────────────────────
    doc.add_heading("目錄", 1)
    toc_items = [
        ("一、功能簡介", "3"),
        ("二、如何進入", "3"),
        ("三、集團工務概覽", "4"),
        ("　3.1 KPI 指標卡片（16 個指標）", "4"),
        ("　3.2 工務報修摘要", "5"),
        ("　3.3 近 12 個月趨勢圖 & 報修類型分布", "6"),
        ("　3.4 飯店 / 商場每日累計案件數", "7"),
        ("　3.5 明細分析區", "7"),
        ("四、工作日誌", "9"),
        ("　4.1 查詢模式", "9"),
        ("　4.2 工作日誌表格欄位", "10"),
        ("　4.3 班表整合與異常標示", "10"),
        ("　4.4 Excel 匯出", "11"),
        ("五、統計基準說明", "11"),
        ("六、常見操作情境", "12"),
        ("附錄：指標顏色對照表", "13"),
    ]
    for label, pg in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(label)
        r.font.size = Pt(11)
        # 分隔點
        tab = p.add_run(" " + "·" * max(1, (52 - len(label))) + f" {pg}")
        tab.font.size = Pt(10)
        tab.font.color.rgb = GRAY

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────────
    # 一、功能簡介
    # ────────────────────────────────────────────────────────────────────────
    doc.add_heading("一、功能簡介", 1)
    p = doc.add_paragraph()
    p.add_run("集團決策 Dashboard").bold = True
    p.add_run(
        "（路由：/exec-work-dashboard）是集團工務管理的核心決策視角，"
        "整合飯店工務部與商場工務報修的即時 KPI 數據，讓管理層一眼掌握全集團工務狀態。"
        "頁面資料在進入時自動從後端 API 平行載入，無需手動觸發同步。"
    )

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("資料來源（自動平行載入）：").bold = True
    sources = [
        "商場工務報修 KPI（year/month 篩選）",
        "飯店工務部 KPI（year/month 篩選）",
        "飯店工項類別月累計（year 篩選）",
        "商場工項類別月累計（year 篩選）",
        "飯店工項類別日累計（year/month 篩選）",
        "商場工項類別日累計（year/month 篩選）",
        "明細分析工時表（year/month，含所有來源）",
    ]
    for s in sources:
        doc.add_paragraph(f"• {s}", style="List Bullet")

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.add_run("本頁面包含三個主要頁籤：").bold = True
    tabs_desc = [
        ("集團工務概覽", "KPI 儀表板、報修摘要、趨勢圖、每日累計表、明細分析、異常提醒"),
        ("工作日誌", "逐筆工作記錄查詢（單日 / 區間 / 整月 / 人員模式）"),
        ("統計基準說明", "各指標計算口徑的官方說明文件"),
    ]
    for name, desc in tabs_desc:
        p = doc.add_paragraph()
        p.add_run(f"① {name}：").bold = True
        p.add_run(desc)

    # ────────────────────────────────────────────────────────────────────────
    # 二、如何進入
    # ────────────────────────────────────────────────────────────────────────
    doc.add_heading("二、如何進入", 1)
    steps = [
        "登入系統後，點選左側側邊欄最上方的「★集團決策 Dashboard」（標星號代表快速入口）",
        "系統自動載入當月資料，等待右上角顯示「更新於 HH:MM:SS」即代表載入完成（約 3–5 秒）",
        "如需重新載入資料，點選右上角「重新整理」按鈕",
        "如需查看不同月份，於「工務報修」區塊調整「報修年月」下拉選單",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────────
    # 三、集團工務概覽
    # ────────────────────────────────────────────────────────────────────────
    doc.add_heading("三、集團工務概覽", 1)

    # 3.1 KPI
    doc.add_heading("3.1  KPI 指標卡片（16 個指標）", 2)
    add_image_safe(doc, shots.get("01_kpi_header"))
    doc.add_paragraph()
    doc.add_paragraph(
        "頁面頂部共有 16 個 KPI 指標卡片，分三列排列。每張卡片以彩色上邊框區分類型，"
        "數值顏色會依狀態自動變化（紅 / 橙 / 綠）。"
    )
    doc.add_paragraph()

    # KPI 說明表格
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["指標名稱", "說明", "顏色警示規則"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], "1B3A5C")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    kpi_rows = [
        # 第一列
        ("本月總案件", "飯店 + 商場當月報修案件總數", "深藍（固定）"),
        ("本月總工時 ℹ", "含工時記錄案件之總工時（HR）\n⚠ 僅含 work_hours > 0 的案件，可能低於各模組工時加總", "藍色（固定）"),
        ("完成件數", "已結案案件總數", "綠色（固定）"),
        ("未完成件數", "尚未結案件數", ">0 → 紅色\n=0 → 綠色"),
        ("完成率", "完成件數 ÷ 總案件 × 100%", "≥80% 綠\n50–79% 橙\n<50% 紅"),
        ("均工時/件", "總工時 ÷ 有工時記錄件數", "橙色（固定）"),
        ("飯店案件占比", "飯店案件 ÷ 總案件 × 100%", "深綠（固定）"),
        ("商場案件占比", "商場案件 ÷ 總案件 × 100%", "深藍（固定）"),
        # 第二列
        ("飯店待辦驗數", "飯店工務待客戶驗收件數", ">0 → 橙色\n=0 → 綠色"),
        ("商場待辦驗數", "商場工務待客戶驗收件數", ">0 → 橙色\n=0 → 綠色"),
        ("飯店上期未結", "飯店上月遺留未結案件數", ">0 → 紅色\n=0 → 綠色"),
        ("商場上期未結", "商場上月遺留未結案件數", ">0 → 紅色\n=0 → 綠色"),
        # 第三列
        ("上級交辦件數 ℹ", "主管交辦任務件數（本月，依建立日期篩選）", "紫色（固定）"),
        ("上級交辦工時", "主管交辦任務總工時（HR）", "淡紫色（固定）"),
        ("緊急事件件數 ℹ", "緊急事件任務件數（本月，依建立日期篩選）", "深紅（固定）"),
        ("緊急事件工時", "緊急事件任務總工時（HR）", "紅色（固定）"),
    ]
    for row_data in kpi_rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    add_tip(doc, "卡片標題右方有 ℹ 圖示的指標，游標停留可查看詳細的資料來源說明。")

    doc.add_paragraph()

    # 3.2 工務報修摘要
    doc.add_heading("3.2  工務報修摘要", 2)
    add_image_safe(doc, shots.get("02_repair_cards"))
    doc.add_paragraph()

    repair_items = [
        ("報修年月篩選器",
         "「報修年月」下拉選單可選擇年份與月份，切換後頁面自動重新載入該月數據。預設為當月。"),
        ("飯店工務部 摘要卡",
         "顯示飯店工務部的報修件數、已結案、待辦驗數、未結案、結案率（含進度條）、"
         "平均結案天數、工時（h），以及當月費用小計（委外＋維修、扣款費用、扣款專櫃、當月小計）。\n"
         "卡片底部有一句話「智慧摘要」，自動判斷並提示最重要的異常（如：最久未結 X 天）。"),
        ("商場工務報修 摘要卡",
         "結構同上，顯示商場工務部數據。"),
        ("飯店／商場工項類別比較",
         "表格並列飯店與商場各工項類別的案件件數與占比。"
         "工項包含：現場報修、上級交辦、緊急事件、例行維護、每日巡檢。"),
    ]
    for title, desc in repair_items:
        p = doc.add_paragraph()
        p.add_run(f"▸ {title}：").bold = True
        p.add_run(desc)
        doc.add_paragraph()

    add_tip(doc, "點選摘要卡右上角「查看詳情 ›」可直接跳轉到對應模組的詳細清單頁面。")

    doc.add_page_break()

    # 3.3 趨勢圖
    doc.add_heading("3.3  近 12 個月趨勢圖 & 報修類型分布", 2)
    add_image_safe(doc, shots.get("03_charts_hotel"))
    doc.add_paragraph("▲ 飯店工務部")
    doc.add_paragraph()
    add_image_safe(doc, shots.get("04_charts_mall"))
    doc.add_paragraph("▲ 商場工務報修")
    doc.add_paragraph()
    doc.add_paragraph(
        "此區各有兩組圖表，飯店與商場各一排：\n"
        "• 折線圖（左）：顯示近 12 個月「報修件數（深藍線）」與「完成件數（綠線）」趨勢，"
        "可快速判斷工務量增減與結案效率。\n"
        "• 圓餅圖（右）：顯示當月報修類型分布（衛浴、空調、內裝、建築等），"
        "占比超過 4% 的類型會顯示標籤。"
    )

    doc.add_paragraph()

    # 3.4 每日累計表
    doc.add_heading("3.4  飯店 / 商場每日累計案件數", 2)
    add_image_safe(doc, shots.get("05_daily_table"))
    doc.add_paragraph()
    doc.add_paragraph(
        "以 Collapse（折疊面板）呈現，預設收合，點選標題列展開。"
        "頁面右上角「⊕ 全展開」/ 「⊖ 全收合」可同時控制飯店與商場兩張表。"
    )
    doc.add_paragraph()
    doc.add_paragraph("表格結構說明：")
    doc.add_paragraph("• 列 = 各工項類別（現場報修 / 上級交辦 / 緊急事件 / 例行維護 / 每日巡檢 / TOTAL）",
                      style="List Bullet")
    doc.add_paragraph("• 欄 = 當月每一天（日期 + 星期）", style="List Bullet")
    doc.add_paragraph("• 每格顯示當日案件件數，可點選跳轉至對應模組詳細頁", style="List Bullet")
    doc.add_paragraph("• 上級交辦 / 緊急事件 僅有月合計（標示「月計」），無每日分拆", style="List Bullet")
    doc.add_paragraph()
    add_tip(doc, "點選工項類別名稱（如「現場報修」）可直接跳轉至對應模組清單。")

    doc.add_paragraph()

    # 3.5 明細分析
    doc.add_heading("3.5  明細分析區", 2)
    add_image_safe(doc, shots.get("06_hours_tables"))
    doc.add_paragraph()
    add_image_safe(doc, shots.get("07_burden_compare"))
    doc.add_paragraph()
    add_image_safe(doc, shots.get("08_matrix_alerts"))
    doc.add_paragraph()
    doc.add_paragraph("以下六個分析面板均以 Collapse 折疊呈現，可個別點選展開：")
    doc.add_paragraph()

    analysis_items = [
        ("📅 每日累計工時表",
         "顯示選定月份每天各工項類別的工時（HR）。"
         "工時以顏色標記：≥8HR 紅色、4–8HR 橙色、>0 一般顯示。"
         "注意：需選擇「月份」（非全年）才會顯示此表，全年模式顯示提示文字。"),
        ("📆 每月累計工時表",
         "顯示選定年度全年（1–12 月）各工項類別的工時。"
         "未來月份以「—」顯示。支援依 TOTAL 欄排序。"),
        ("🧮 人員負荷與效率分析",
         "列出每位工務人員的：總工時（HR）、件數、均工時/件、主要工項類別、"
         "負荷判斷（需關注🔴 / 工時偏高🟠 / 正常🟢）。\n"
         "判斷基準：均工時/件 ≥ 3.0HR → 需關注；2.5–3.0HR → 工時偏高；< 2.5HR → 正常。\n"
         "欄位標題有 ℹ 圖示，停留可查看詳細計算說明。"),
        ("🏢 飯店 vs 商場比較表",
         "並排比較飯店工務部與商場工務報修的：案件數、工時（h）、完成件數、"
         "未完成件數、完成率、主要工項類別。最底列顯示集團合計。"),
        ("📊 工項類別 × 單位矩陣",
         "以工項類別為列，顯示飯店件數、商場件數、合計件數、件占比、"
         "飯店工時、商場工時、總工時。最底列顯示各欄合計。"),
        ("⚠️ 異常提醒（預設展開）",
         "系統自動分析並顯示以下警示（若無異常則顯示「✅ 本月無異常警示」）：\n"
         "• 🔴 未完成件數警示（>0 件）\n"
         "• 🔴 完成率偏低（<60%）\n"
         "• 🟠 完成率注意（60–79%）\n"
         "• 🟡 工項類別集中（單類 >60%）\n"
         "• 🟠 人員超載（工時 >80HR）\n"
         "• 🔵 單日工時暴增（>月均 ×2）"),
    ]
    for title, desc in analysis_items:
        p = doc.add_paragraph()
        p.add_run(f"▸ {title}").bold = True
        doc.add_paragraph(desc)
        doc.add_paragraph()

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────────
    # 四、工作日誌
    # ────────────────────────────────────────────────────────────────────────
    doc.add_heading("四、工作日誌", 1)
    add_image_safe(doc, shots.get("09_journal_overview"))
    doc.add_paragraph()
    doc.add_paragraph(
        "工作日誌頁籤整合三個資料來源的逐筆工作記錄，讓管理者可以查詢"
        "任意日期或人員的完整工作明細。"
    )

    # 4.1 查詢模式
    doc.add_heading("4.1  查詢模式", 2)
    add_image_safe(doc, shots.get("10_journal_single_mode"))
    doc.add_paragraph()

    modes = [
        ("單日模式", "選擇單一日期，查看當日全部人員工作記錄。人員間以 Collapse 展開，可個別收合。"),
        ("區間模式", "選擇起迄日期，查看跨日記錄；以日期分層展開，每天再展開各人員。"),
        ("整月模式", "選擇年月，一次載入整月記錄，結構同區間模式。"),
        ("人員模式", "先選人員、再選日期區間，集中顯示該人員在區間內的所有工作記錄。"),
    ]
    for mode, desc in modes:
        p = doc.add_paragraph()
        p.add_run(f"• {mode}：").bold = True
        p.add_run(desc)

    doc.add_paragraph()
    add_tip(doc, "有資料後右側會出現「全部縮合 / 全部展開」按鈕，可快速收合所有人員/日期面板。")

    # 4.2 表格欄位
    doc.add_heading("4.2  工作日誌表格欄位", 2)
    fields = [
        ("項次", "此人員當日工作序號"),
        ("工項類別", "現場報修 / 上級交辦 / 緊急事件 / 例行維護 / 每日巡檢（以 ✓ 標記）"),
        ("工作事項", "工作描述，前方有「飯」/「商」小標籤區分飯店或商場記錄"),
        ("預估耗時(min)", "作業前預估所需時間（分鐘）"),
        ("起", "工作開始時間（HH:MM）"),
        ("迄", "工作結束時間（HH:MM）"),
        ("工時(min)", "實際工時（分鐘），以深藍加粗顯示"),
        ("備註", "補充說明（橙色字體）"),
        ("回報事項", "需向主管回報的內容"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["欄位名稱", "說明"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(table.rows[0].cells[i], "1B3A5C")
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for name, desc in fields:
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = desc
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    add_tip(doc, "點選任一列可展開 Drawer 明細，顯示完整原始欄位（Ragic 來源資料）及維修圖片。")

    # 4.3 班表整合
    doc.add_heading("4.3  班表整合與異常標示", 2)
    doc.add_paragraph(
        "每位人員姓名前方會顯示班別 Tag，幫助管理者快速發現「非上班日卻有工作記錄」的異常情況："
    )
    shift_cases = [
        ("彩色班別代碼（如 A、B）", "正常上班班別，游標停留可查看班別名稱"),
        ("🔴 MinusCircle（—）", "當日為非上班班別（排休），但仍有工作記錄"),
        ("🟡 FileUnknown（?）", "班表查無此人記錄，但仍有工作記錄"),
        ("❌ UserDelete", "人員未指定（工作記錄無歸屬人員）"),
    ]
    for tag, desc in shift_cases:
        p = doc.add_paragraph()
        p.add_run(f"• {tag}：").bold = True
        p.add_run(desc)

    # 4.4 Excel 匯出
    doc.add_heading("4.4  Excel 匯出", 2)
    doc.add_paragraph(
        "查詢完成後，右側會出現「匯出 Excel」按鈕（綠色外框）。"
        "點選後下載 .xlsx 檔案，檔名格式為「工作日誌_起始日期_結束日期.xlsx」，"
        "人員模式會額外加上人員姓名。"
    )

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────────
    # 五、統計基準說明
    # ────────────────────────────────────────────────────────────────────────
    doc.add_heading("五、統計基準說明", 1)
    add_image_safe(doc, shots.get("11_methodology"))
    doc.add_paragraph()
    doc.add_paragraph(
        "此頁籤內嵌 HTML 說明文件（/report-count-methodology.html），"
        "詳細解釋各統計指標的計算口徑，包含："
    )
    methodology_items = [
        "件數計算方式（含 / 不含 work_hours=0 的差異）",
        "工時欄位定義（work_hours vs work_min 的換算）",
        "飯店與商場資料來源差異（dazhi vs luqun）",
        "IHG 客房保養的特殊計算方式",
        "上期未結（prev_uncompleted）的定義",
    ]
    for item in methodology_items:
        doc.add_paragraph(f"• {item}", style="List Bullet")
    doc.add_paragraph()
    doc.add_paragraph("建議所有使用者在首次閱讀 Dashboard 前先瀏覽此說明，避免對數字解讀產生歧義。")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────────
    # 六、常見操作情境
    # ────────────────────────────────────────────────────────────────────────
    doc.add_heading("六、常見操作情境", 1)

    scenarios = [
        ("查看本月集團工務整體完成率",
         [
             "進入頁面，等待資料載入完成",
             "查看 KPI 卡片第一列的「完成率」（大字顯示，紅色＝偏低、橙色＝注意、綠色＝良好）",
             "同時確認「⚠️ 異常提醒」面板是否有完成率相關警示",
         ]),
        ("比較特定月份飯店 vs 商場的報修量",
         [
             "工務報修區段 → 調整「報修年月」選單為目標月份",
             "比對左側「飯店工務部」卡片與右側「商場工務報修」卡片的報修總數與結案率",
             "或展開「🏢 飯店 vs 商場比較表」查看更完整的對比數據",
         ]),
        ("找出工時最多 / 負荷最重的人員",
         [
             "展開「🧮 人員負荷與效率分析」Collapse",
             "查看「判斷」欄位標示「需關注」（紅色）或「工時偏高」（橙色）的人員",
             "點選 HR 欄標題可排序，快速找到工時最高者",
         ]),
        ("查詢特定人員某週的工作記錄",
         [
             "切換至「工作日誌」頁籤",
             "選擇「人員」模式（Segmented 按鈕）",
             "選擇目標週次的起迄日期（RangePicker）",
             "在人員下拉選單選取目標人員",
             "點選「查詢」按鈕，結果以日期分層展開",
         ]),
        ("了解某天哪個工項類別最忙",
         [
             "展開「飯店每日累計案件數」或「商場每日累計案件數」",
             "找到目標日期欄，查看各工項類別的數字",
             "點選數字可直接跳轉至對應模組詳細清單",
         ]),
        ("月底快速生成工作日誌 Excel 報表",
         [
             "切換至「工作日誌」頁籤 → 選「整月」模式",
             "選擇目標年月 → 點「查詢」",
             "等待資料載入完成後，點選「匯出 Excel」按鈕",
             "下載的 .xlsx 包含該月全部人員的工作記錄",
         ]),
    ]

    for i, (title, steps) in enumerate(scenarios, 1):
        p = doc.add_paragraph()
        p.add_run(f"情境 {i}：{title}").bold = True
        p.runs[0].font.color.rgb = PRIMARY
        for j, step in enumerate(steps, 1):
            doc.add_paragraph(f"　{j}. {step}")
        doc.add_paragraph()

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────────
    # 附錄：顏色對照表
    # ────────────────────────────────────────────────────────────────────────
    doc.add_heading("附錄：指標顏色對照表", 1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["指標 / 狀況", "顏色", "觸發條件"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(table.rows[0].cells[i], "1B3A5C")
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    color_rows = [
        ("完成率", "🟢 綠色", "≥ 80%"),
        ("完成率", "🟠 橙色", "50%–79%"),
        ("完成率", "🔴 紅色", "< 50%"),
        ("均工時/件（負荷判斷）", "🟢 正常", "< 2.5 HR"),
        ("均工時/件（負荷判斷）", "🟠 工時偏高", "2.5–3.0 HR"),
        ("均工時/件（負荷判斷）", "🔴 需關注", "≥ 3.0 HR"),
        ("每日工時（表格格）", "一般（深灰）", "> 0 且 < 4 HR"),
        ("每日工時（表格格）", "🟠 橙色", "4–8 HR"),
        ("每日工時（表格格）", "🔴 紅色", "≥ 8 HR"),
        ("待辦驗數 / 上期未結", "🟢 綠色", "= 0（無異常）"),
        ("待辦驗數 / 上期未結", "🟠 / 🔴", "> 0（需追蹤）"),
        ("最久未結案天數（進度條）", "🟢 綠色", "< 7 天"),
        ("最久未結案天數（進度條）", "🟠 橙色", "7–13 天"),
        ("最久未結案天數（進度條）", "🔴 紅色", "≥ 14 天"),
        ("班別 Tag", "彩色方塊（班別代碼）", "正常上班班別"),
        ("班別 Tag", "🔴 紅色 MinusCircle", "排休但有工作記錄"),
        ("班別 Tag", "🟡 黃色 FileUnknown", "班表查無記錄"),
    ]
    for row_data in color_rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    # 儲存
    doc.save(str(output_path))
    print(f"\n✅ 手冊已儲存：{output_path}")


# ── 主程式 ────────────────────────────────────────────────────────────────────
async def main():
    OUTPUT_DIR.mkdir(exist_ok=True)

    print("=" * 55)
    print("  集團決策 Dashboard 教學手冊生成器")
    print("=" * 55)
    print(f"帳號：{USERNAME}")
    print(f"截圖儲存位置：{OUTPUT_DIR}")
    print(f"手冊輸出位置：{MANUAL_PATH}")
    print()

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=["--disable-blink-features=AutomationControlled"],
        )
        context = await browser.new_context(
            viewport={"width": 1600, "height": 900},
            device_scale_factor=1.5,        # 高解析度截圖（2400×1350 px）
            locale="zh-TW",
            timezone_id="Asia/Taipei",
        )
        page = await context.new_page()

        shots = await capture_screenshots(page, OUTPUT_DIR)
        await browser.close()

    print(f"\n截圖完成，共 {len(shots)} 張")
    print("\n📄 生成 Word 手冊...")
    build_word_manual(shots, MANUAL_PATH)
    print("\n✅ 執行完成！請開啟：", MANUAL_PATH)


if __name__ == "__main__":
    asyncio.run(main())
