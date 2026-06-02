#!/usr/bin/env python3
"""
商場工務報修（/luqun-repair/dashboard）影音教學旁白腳本生成器
===========================================================
使用方式：
  python luqun_generate_video_scripts.py

輸出：商場工務報修_影音教學腳本.docx

本腳本架構與飯店工務報修（dazhi-repair）腳本相同，
已看過飯店版教學的觀眾可直接使用本腳本錄製商場版（差異僅在資料來源）。
"""

import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx"); sys.exit(1)

OUTPUT_PATH = Path(__file__).parent / "商場工務報修_影音教學腳本.docx"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
ORANGE  = RGBColor(0xFA, 0xAD, 0x14)
GRAY    = RGBColor(0x59, 0x59, 0x59)


# ── 輔助函式 ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def narration_block(doc, text: str):
    for part in text.split("\n\n"):
        part = part.strip()
        if not part:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(part); r.font.size = Pt(11.5)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "12"); left.set(qn("w:color"), "4BA8E8")
        pBdr.append(left); pPr.append(pBdr)


def label_box(doc, label: str, text: str, fill_hex: str):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    lc = table.rows[0].cells[0]
    set_cell_bg(lc, fill_hex)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label); lr.bold = True; lr.font.size = Pt(10)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = table.rows[0].cells[1]
    set_cell_bg(rc, "F8F9FA")
    rc.paragraphs[0].add_run(text).font.size = Pt(10.5)
    doc.add_paragraph()


def action_block(doc, text: str):  label_box(doc, "🖱 操作", text, "52C41A")
def hint_block(doc, text: str):    label_box(doc, "⏸ 提示", text, "FA8C16")


def timing_row(doc, t_range: str, label: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"[{t_range}]  ")
    r1.bold = True; r1.font.color.rgb = ORANGE; r1.font.size = Pt(10)
    r2 = p.add_run(label)
    r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = PRIMARY


def section_divider(doc):
    p = doc.add_paragraph("─" * 64)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── 腳本資料 ──────────────────────────────────────────────────────────────────
SCRIPTS = [

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第一集",
        "title": "功能簡介 & 查詢列 & KPI 7 卡",
        "duration": "建議時長：4–5 分鐘",
        "goal": "讓觀眾了解商場工務報修的七個頁籤架構、查詢列操作，以及 KPI 7 卡的意義與點擊功能。",
        "segments": [
            {
                "time": "00:00–00:30",
                "label": "功能簡介",
                "narration": (
                    "大家好，歡迎收看商場工務報修教學系列。"
                    "這個模組管理商場（春大直）工務部的所有報修案件，"
                    "從報修到結案的全程追蹤，包含費用統計和 PowerPoint 匯出。"
                    "\n\n"
                    "功能架構和飯店工務報修完全對應，"
                    "如果你已看過飯店版，只需知道這是商場的資料來源。"
                    "我們這集從頭完整介紹。"
                ),
                "actions": ["進入頁面，讓觀眾看到 Dashboard 全貌"],
                "hints": [],
            },
            {
                "time": "00:30–01:00",
                "label": "7 個頁籤概覽",
                "narration": (
                    "頁面有七個頁籤。"
                    "\n\n"
                    "Dashboard 是預設頁籤，整合 KPI 和圖表。"
                    "3.1 到 3.4 是四個統計分析視角。"
                    "金額統計是費用彙整，"
                    "最後的「報修清單總表」就是完整的商場報修案件清單，"
                    "也就是俗稱的「春大直報修清單」。"
                ),
                "actions": ["依序指向七個頁籤"],
                "hints": [],
            },
            {
                "time": "01:00–01:30",
                "label": "查詢列與匯出工具",
                "narration": (
                    "頁面頂部的查詢列控制所有頁籤的資料範圍。"
                    "左側年度選單選擇年份，右側月份可選全年或特定月份。"
                    "選好後點「查詢」更新，點「重設」回到預設。"
                    "\n\n"
                    "右側有紫色的「匯出 PowerPoint」按鈕和「連線測試」按鈕，"
                    "第五集會介紹匯出功能。"
                ),
                "actions": ["示範切換月份，點查詢"],
                "hints": ["建議選有積案資料的近期月份"],
            },
            {
                "time": "01:30–03:30",
                "label": "KPI 7 卡解讀",
                "narration": (
                    "Dashboard 頂部是七張 KPI 卡片，全部都可以點擊展開案件清單 Modal。"
                    "\n\n"
                    "第一張「本月相關案件」深藍色——上月遺留加上本月新報修件數。"
                    "\n\n"
                    "第二張「已完成件數」綠色——以「已驗收」為完成標準的累計件數。"
                    "\n\n"
                    "第三張「待辦驗件數」黃色——已修好等客戶驗收確認的商場案件。"
                    "\n\n"
                    "第四張「未完成件數」紅色——積案數，這個是最需要關注的指標。"
                    "\n\n"
                    "第五張「平均結案天數」藍色——商場工務結案效率指標。"
                    "\n\n"
                    "第六張「本月工時統計」青色——商場工務本月工時（hr），"
                    "副標有換算成工作天的參考值。"
                    "\n\n"
                    "第七張「客房報修件數」橙色——本期涉及商場客房的報修件數。"
                    "\n\n"
                    "每張卡片右上角有 ℹ 圖示，停留後顯示計算口徑的詳細說明。"
                ),
                "actions": [
                    "放大顯示七張卡片，逐一指向",
                    "停留在一張卡片的 ℹ 圖示顯示 Tooltip",
                ],
                "hints": [],
            },
            {
                "time": "03:30–04:15",
                "label": "點擊 KPI 展開案件清單",
                "narration": (
                    "點任何一張 KPI 卡片，跳出案件清單 Modal，"
                    "列出符合條件的每筆商場報修案件，"
                    "方便直接確認是哪些案件構成這個數字。"
                    "\n\n"
                    "下一集繼續看費用卡片和圖表。"
                ),
                "actions": [
                    "點「未完成件數」，展示 Modal",
                    "點 X 關閉",
                ],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第二集",
        "title": "費用 KPI & 圖表區",
        "duration": "建議時長：3–4 分鐘",
        "goal": "讓觀眾學會解讀商場報修的費用三卡和四個圖表。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "這集看費用 KPI 和圖表區，幫助管理者掌握商場報修費用結構和工務量趨勢。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:30",
                "label": "費用 KPI 3 張",
                "narration": (
                    "KPI 7 卡下方是三張費用卡片。"
                    "\n\n"
                    "左邊「委外+維修費用」紫色——累計到選定月的委外加維修費用合計，"
                    "副標顯示兩者各別金額。點擊展開費用明細 Modal，"
                    "列出所有有費用記錄的商場報修案件。"
                    "\n\n"
                    "中間「扣款費用」紅色——商場端的累計扣款費用，"
                    "點擊展開扣款事項明細，含每筆扣款原因說明。"
                    "\n\n"
                    "右邊「當月金額」青色——當月的委外+維修費、扣款費用、扣款專櫃、當月小計，"
                    "不可點擊，直接顯示數字。"
                ),
                "actions": [
                    "指向三張費用卡片",
                    "點「委外+維修費用」展示 Modal，再關閉",
                ],
                "hints": [],
            },
            {
                "time": "01:30–02:30",
                "label": "折線趨勢圖 & 類型圓餅圖",
                "narration": (
                    "費用卡片下方是兩個圖表。"
                    "\n\n"
                    "左邊是近 12 個月商場報修趨勢折線圖——"
                    "深藍線是報修件數，綠線是完成件數。"
                    "如果兩條線距離越來越大，代表商場積案在增加。"
                    "\n\n"
                    "右邊是報修類型分布圓餅圖——"
                    "各類型占比，停留顯示精確數字。"
                    "找出商場最常發生的報修類型，作為預防性保養的依據。"
                ),
                "actions": ["停留在折線圖和圓餅圖顯示 Tooltip"],
                "hints": [],
            },
            {
                "time": "02:30–03:15",
                "label": "樓層分布 & 狀況分布",
                "narration": (
                    "下方兩個水平長條圖。"
                    "\n\n"
                    "「發生樓層分布」——找出商場哪個樓層報修件數最多，"
                    "作為加強巡檢或預防保養的重點區域。"
                    "\n\n"
                    "「處理狀況分布」——顯示各狀態的件數，"
                    "綠色已完成、藍色進行中、黃色等待類、紅色有問題的狀態。"
                    "\n\n"
                    "好，費用和圖表介紹完了。下一集看 Top 10 和案件 Drawer。"
                ),
                "actions": ["指向兩個圖表說明"],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第三集",
        "title": "Top 10 快速摘要 & 案件詳情 Drawer",
        "duration": "建議時長：4 分鐘",
        "goal": "讓觀眾學會用三個 Top 10 清單快速鎖定問題案件，並掌握 Drawer 所有欄位（含維修圖片）。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "Dashboard 底部有三個 Top 10 快速摘要，點擊任何一筆都可以展開案件詳情 Drawer。"
                    "這集介紹三個清單和完整的 Drawer 欄位。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:20",
                "label": "三個 Top 10 清單",
                "narration": (
                    "左邊「未完成案件 Top 10」紅色——"
                    "依等待天數降序排列，讓管理者快速找到積壓最久的商場報修。"
                    "等待超過 30 天顯示紅色警示，7 到 29 天顯示橙色。"
                    "\n\n"
                    "中間「高費用案件 Top 10」紫色——"
                    "商場費用金額最高的案件，顯示金額和是否已結案。"
                    "\n\n"
                    "右邊「高工時案件 Top 10」藍色——"
                    "工時最高的商場報修案件，反映工務人員投入時間最多的工作。"
                    "\n\n"
                    "三個清單的每筆記錄都是可以點擊的，點下去就展開右側 Drawer。"
                ),
                "actions": [
                    "指向三個 Top 10 清單",
                    "指向等待天數的顏色 Tag 說明",
                ],
                "hints": ["建議選有積案的月份，讓清單有資料"],
            },
            {
                "time": "01:20–03:30",
                "label": "案件詳情 Drawer",
                "narration": (
                    "點擊任一 Top 10 清單的案件，右側滑出案件詳情 Drawer。"
                    "\n\n"
                    "標題顯示「報修詳情：報修編號」，"
                    "旁邊有「在 Ragic 查看」的按鈕，點擊在新分頁開啟商場 Ragic 的原始記錄。"
                    "\n\n"
                    "Drawer 的欄位包含：報修編號、標題（粗體）、報修人姓名、"
                    "報修類型（Tag）、發生樓層、發生時間、負責單位。"
                    "\n\n"
                    "工時與費用：花費工時、委外費用、維修費用、總費用（紫色粗體）。"
                    "\n\n"
                    "驗收結案：驗收者、結案人、結案時間、結案天數。"
                    "\n\n"
                    "扣款資訊：扣款事項和扣款費用（若有）。"
                    "\n\n"
                    "最下方是維修圖片——系統自動從資料庫載入商場的維修照片，"
                    "點縮圖展開 Lightbox 全螢幕預覽，支援左右切換。"
                    "\n\n"
                    "這個 Drawer 讓管理者不需要切換到 Ragic，"
                    "就能在 Dashboard 直接看到完整的商場報修記錄。"
                ),
                "actions": [
                    "點一筆 Top 10 展開 Drawer",
                    "指向各欄位說明",
                    "指向維修圖片，點縮圖示範 Lightbox",
                    "指向「在 Ragic 查看」連結",
                ],
                "hints": ["建議選有維修圖片的商場案件"],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第四集",
        "title": "統計分析 Tab（3.1 ~ 3.4）& 金額統計",
        "duration": "建議時長：4 分鐘",
        "goal": "讓觀眾掌握四個統計 Tab 和金額統計的用途，幫助深入分析商場工務報修數據。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "這集介紹五個統計分析頁籤，"
                    "幫助管理者從不同角度深入了解商場工務報修的狀況。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:00",
                "label": "3.1 報修統計",
                "narration": (
                    "切換到「3.1 報修」Tab。"
                    "\n\n"
                    "折線圖顯示全年各月商場的報修件數、完成件數、未完成件數趨勢。"
                    "點圖表中的特定月份，下方顯示那個月的商場案件明細清單。"
                    "\n\n"
                    "這個 Tab 適合觀察商場工務量的季節性規律——"
                    "例如商場旺季是否報修件數特別多，或者某幾個月積案特別嚴重。"
                ),
                "actions": [
                    "切換到「3.1 報修」Tab",
                    "點圖表一個月份展示明細",
                ],
                "hints": [],
            },
            {
                "time": "01:00–01:35",
                "label": "3.2 結案時間統計",
                "narration": (
                    "「3.2 結案時間」Tab 分析商場報修從申請到結案所花的天數分布。"
                    "以區間顯示：0 到 7 天、8 到 30 天等。"
                    "\n\n"
                    "如果大部分案件集中在 0 到 7 天，"
                    "代表商場工務部處理效率很好。"
                    "如果 30 天以上的比例偏高，需要找出拖延原因。"
                ),
                "actions": ["切換到「3.2 結案時間」Tab，指向各區間長條"],
                "hints": [],
            },
            {
                "time": "01:35–02:10",
                "label": "3.3 報修類型統計",
                "narration": (
                    "「3.3 報修類型」Tab 分析商場各報修類型的件數占比。"
                    "\n\n"
                    "找出商場最常出現的報修類型，"
                    "作為商場預防性保養計劃的資源配置依據。"
                    "例如空調系統報修最多，可以考慮增加空調的定期保養頻率。"
                ),
                "actions": ["切換到「3.3 報修類型」Tab"],
                "hints": [],
            },
            {
                "time": "02:10–02:45",
                "label": "3.4 本月客房報修表",
                "narration": (
                    "「3.4 本月客房報修表」依商場客房房號列出本月的報修記錄。"
                    "\n\n"
                    "可以找出哪幾間商場單元本月報修次數特別多，"
                    "判斷是否有需要較大規模整修的區域。"
                ),
                "actions": ["切換到「3.4 本月客房報修表」Tab"],
                "hints": [],
            },
            {
                "time": "02:45–03:15",
                "label": "金額統計",
                "narration": (
                    "「金額統計」Tab 以年度彙整所有有費用記錄的商場報修案件，"
                    "顯示委外費用、維修費用、扣款費用的月份分布和全年合計。"
                    "\n\n"
                    "適合年度費用報表或向管理層說明商場報修費用結構。"
                    "\n\n"
                    "好，五個統計 Tab 介紹完了。下一集是最後一集，看報修清單總表和匯出功能。"
                ),
                "actions": ["切換到「金額統計」Tab"],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第五集",
        "title": "報修清單總表 & 匯出 & 常見情境",
        "duration": "建議時長：3–4 分鐘",
        "goal": "讓觀眾學會用報修清單總表做多條件篩選，以及匯出 Excel 和 PowerPoint，並示範兩個情境。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場",
                "narration": (
                    "最後一集介紹「報修清單總表」——商場版的完整案件清單，"
                    "以及兩種匯出功能。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:40",
                "label": "報修清單總表",
                "narration": (
                    "切換到「報修清單總表」Tab。"
                    "這裡是商場工務報修的完整清單，正式名稱叫「春大直-報修清單總表」。"
                    "\n\n"
                    "頂部有多條件篩選：年月、報修類型、樓層、負責單位、狀態、關鍵字。"
                    "可以組合使用，例如篩選「衛浴類型、3 樓、未結案」，"
                    "快速找出特定情境的商場報修。"
                    "\n\n"
                    "清單欄位包含報修編號、標題、樓層、報修時間、狀態 Tag、"
                    "完工時間、結案天數、工時、費用。"
                    "\n\n"
                    "點擊任一列，右側展開案件詳情 Drawer，和 Top 10 展開的 Drawer 相同。"
                    "\n\n"
                    "右上角「匯出 Excel」按鈕，下載目前篩選結果為 .xlsx 檔案，"
                    "檔名帶年月條件，方便歸檔。"
                ),
                "actions": [
                    "切換到「報修清單總表」Tab",
                    "示範組合篩選條件",
                    "點一列展示 Drawer",
                    "點「匯出 Excel」按鈕",
                ],
                "hints": [],
            },
            {
                "time": "01:40–02:10",
                "label": "匯出 PowerPoint",
                "narration": (
                    "頁頭右側有「匯出 PowerPoint」紫色按鈕。"
                    "確認查詢列已選定特定月份（不能選全年），點擊後等待下載，"
                    "生成含商場 KPI 的 .pptx 簡報，適合月報或主管匯報場合。"
                ),
                "actions": [
                    "確認選定月份",
                    "點匯出 PowerPoint，展示進度條",
                ],
                "hints": [],
            },
            {
                "time": "02:10–02:45",
                "label": "情境示範",
                "narration": (
                    "情境一：月底確認商場積案。"
                    "Dashboard → 「未完成件數」KPI 點擊展開 Modal，"
                    "查看「未完成案件 Top 10」找等待最久的案件，點入 Drawer 聯絡負責人員。"
                    "\n\n"
                    "情境二：匯出本月商場報修 Excel 給上級。"
                    "報修清單總表 → 確認年月篩選正確 → 點「匯出 Excel」，"
                    "30 秒以內完成下載。"
                ),
                "actions": ["示範情境一的流程"],
                "hints": [],
            },
            {
                "time": "02:45–03:20",
                "label": "結尾",
                "narration": (
                    "商場工務報修的五集教學全部結束了。"
                    "\n\n"
                    "快速回顧：第一集學 KPI 7 卡和查詢列，"
                    "第二集學費用卡和圖表，"
                    "第三集學 Top 10 和案件 Drawer，"
                    "第四集學五個統計分析 Tab，"
                    "第五集學清單總表、Excel/PPT 匯出和操作情境。"
                    "\n\n"
                    "這個模組和飯店工務報修是對應的，熟悉一個就能快速上手另一個。謝謝收看！"
                ),
                "actions": ["畫面停在 Dashboard 做結尾"],
                "hints": [],
            },
        ],
    },
]


# ── Word 文件生成 ─────────────────────────────────────────────────────────────
def build_doc(scripts: list, output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height   = Inches(11.69); sec.page_width    = Inches(8.27)
    sec.top_margin    = Inches(0.85);  sec.bottom_margin = Inches(0.85)
    sec.left_margin   = Inches(1.1);   sec.right_margin  = Inches(1.1)

    # 封面
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_heading("商場工務報修", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY; t.runs[0].font.size = Pt(26)
    t2 = doc.add_heading("影音教學旁白腳本", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT; t2.runs[0].font.size = Pt(20)
    doc.add_paragraph()
    sub = doc.add_paragraph(
        f"共五集 · 春大直 · 路由：/luqun-repair/dashboard\n{datetime.now().strftime('%Y 年 %m 月 %d 日')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    doc.add_paragraph()

    # 使用說明
    doc.add_heading("腳本使用說明", 1).runs[0].font.color.rgb = PRIMARY
    for tip in [
        "🎬 旁白：錄製時直接念，可依個人習慣調整，不必逐字照念。",
        "🖱 操作：錄製時同步要做的滑鼠動作。",
        "⏸ 提示：錄製前的準備建議，不需要說出來。",
        "[時間軸]：僅供參考，依實際節奏調整。",
        "建議選有積案（未完成件數 > 0）和費用記錄的月份，示範效果最佳。",
        "已看過飯店工務報修（dazhi）教學的觀眾，此腳本開頭可直接說「架構和飯店版相同，資料來源是商場」。",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(tip).font.size = Pt(10.5)

    doc.add_page_break()

    for script in scripts:
        h = doc.add_heading(f"{script['ep']}  {script['title']}", 1)
        h.runs[0].font.color.rgb = PRIMARY; h.runs[0].font.size = Pt(16)

        info = doc.add_paragraph()
        r1 = info.add_run(f"⏱ {script['duration']}　　")
        r1.font.size = Pt(10); r1.font.color.rgb = ORANGE; r1.bold = True

        gp = doc.add_paragraph()
        gp.paragraph_format.space_after = Pt(8)
        gp.add_run("學習目標：").bold = True
        gp.runs[0].font.color.rgb = ACCENT
        gp.add_run(script["goal"])

        section_divider(doc)

        for seg in script["segments"]:
            timing_row(doc, seg["time"], seg["label"])
            narration_block(doc, seg["narration"])
            for act in seg.get("actions", []):
                action_block(doc, act)
            for hint in seg.get("hints", []):
                hint_block(doc, hint)
            doc.add_paragraph()

        doc.add_page_break()

    # 附錄
    doc.add_heading("附錄：處理狀況顏色速查", 1).runs[0].font.color.rgb = PRIMARY
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["處理狀況", "Tag 顏色", "代表意義"]):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1B3A5C")
    for status, color, desc in [
        ("已驗收 / 已結案 / 完修", "🟢 success 綠",   "案件已完成"),
        ("處理中",                  "🔵 processing 藍","執行中"),
        ("待維修 / 待驗收",         "🟡 warning 黃",   "等待下一步"),
        ("待協調",                  "🟠 orange 橙",    "需協調"),
        ("待排除",                  "🔴 error 紅",     "需即刻處理"),
    ]:
        row = table.add_row()
        row.cells[0].text = status; row.cells[1].text = color; row.cells[2].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 腳本手冊已儲存：{output_path}")


if __name__ == "__main__":
    build_doc(SCRIPTS, OUTPUT_PATH)
    print(f"執行完成！請開啟：{OUTPUT_PATH}")
