#!/usr/bin/env python3
"""
全棟例行維護（/mall/full-building-maintenance）教學手冊自動生成腳本
=================================================================
使用方式：
  pip install playwright python-docx
  playwright install chromium
  python fullbldg_pm_generate_manual.py

輸出：
  manual_screenshots_fullbldg_pm/
  全棟例行維護教學手冊.docx

與商場週期保養表差異：
  1. Dashboard 多一張「保養時間」KPI 卡（共六卡）
  2. 每日巡檢表：只有單日 DatePicker，且標示「本地同步開發中」
  3. 年度計劃表格子多 no_frequency（∅）、no_data（！）兩種狀態
"""

import asyncio
import sys
from pathlib import Path
from datetime import datetime

try:
    from playwright.async_api import async_playwright
except ImportError:
    print("❌ 請先執行：pip install playwright && playwright install chromium"); sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx"); sys.exit(1)

# ── 設定 ──────────────────────────────────────────────────────────────────────
BASE_URL    = "http://localhost:5173"
TARGET_URL  = f"{BASE_URL}/mall/full-building-maintenance"
OUTPUT_DIR  = Path(__file__).parent / "manual_screenshots_fullbldg_pm"
MANUAL_PATH = Path(__file__).parent / "全棟例行維護教學手冊.docx"
USERNAME    = "admin"
PASSWORD    = "Admin@2026"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
GREEN   = RGBColor(0x52, 0xC4, 0x1A)
ORANGE  = RGBColor(0xFA, 0xAD, 0x14)
RED     = RGBColor(0xCF, 0x13, 0x22)
GRAY    = RGBColor(0x59, 0x59, 0x59)

CONTENT_X  = 185
CONTENT_W  = 1390
VIEWPORT_H = 900


# ── 截圖邏輯 ──────────────────────────────────────────────────────────────────
async def capture_screenshots(page, output_dir: Path) -> dict:
    shots = {}

    async def snap(name: str, scroll_y: int = 0, clip=None, wait_ms: int = 800):
        try:
            await page.evaluate(f"window.scrollTo(0, {scroll_y})")
            await page.wait_for_timeout(wait_ms)
            path = output_dir / f"{name}.png"
            await page.screenshot(path=str(path), full_page=False,
                                  **({"clip": clip} if clip else {}))
            shots[name] = path
            print(f"   ✓ {name}")
        except Exception as e:
            print(f"   ✗ {name}：{e}")

    async def click_tab(tab_text: str, wait_ms: int = 2000):
        try:
            tab = page.locator("div.ant-tabs-tab").filter(has_text=tab_text).first
            await tab.click()
            await page.wait_for_timeout(wait_ms)
        except Exception as e:
            print(f"   ⚠ Tab [{tab_text}] 失敗：{e}")

    # ── 登入 ──────────────────────────────────────────────────────────────────
    print("🔐 登入中...")
    await page.goto(f"{BASE_URL}/login")
    await page.wait_for_load_state("domcontentloaded")
    await page.wait_for_timeout(800)
    await page.locator("input[placeholder*='帳號'], input[type='text']").first.fill(USERNAME)
    await page.locator("input[type='password']").first.fill(PASSWORD)
    await page.locator("button[type='submit'], button:has-text('登入')").first.click()
    try:
        await page.wait_for_url(f"{BASE_URL}/dashboard", timeout=15000)
        print("   ✓ 登入成功")
    except Exception:
        print(f"   ✓ 已跳轉至 {page.url}")

    # ── 導航 ──────────────────────────────────────────────────────────────────
    await page.goto(TARGET_URL)
    await page.wait_for_load_state("networkidle")
    try:
        await page.locator(".ant-statistic-content-value").first.wait_for(timeout=12000)
        await page.wait_for_timeout(2000)
        print("   ✓ 頁面載入完成")
    except Exception:
        await page.wait_for_timeout(5000)

    # ════════════════════════════════════════════════════════════
    # Tab 1：Dashboard（KPI 六卡）
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：Dashboard Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await page.wait_for_timeout(600)

    # 1. KPI 六卡（整列）
    await snap("01_kpi_six",
               clip={"x": CONTENT_X, "y": 55, "width": CONTENT_W, "height": 380})

    # 2. 完成率進度條 + 圖表
    await snap("02_charts",
               clip={"x": CONTENT_X, "y": 430, "width": CONTENT_W, "height": 370})

    # 3. 預警區
    await snap("03_alerts",
               scroll_y=800,
               clip={"x": CONTENT_X, "y": 800, "width": CONTENT_W, "height": 360})

    # ════════════════════════════════════════════════════════════
    # Tab 2：每日巡檢表（開發中版本）
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：每日巡檢表 Tab（開發中）")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("每日巡檢表", wait_ms=2500)
    await snap("04_daily_form",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 600})

    # ════════════════════════════════════════════════════════════
    # Tab 3：每月維護
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：每月維護 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("每月維護", wait_ms=2500)
    await snap("05_monthly",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab 4：每季維護
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：每季維護 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("每季維護", wait_ms=2500)
    await snap("06_quarterly",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab 5：每年維護
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：每年維護 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("每年維護", wait_ms=2500)
    await snap("07_yearly",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab 6：排程管理
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：排程管理 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("排程管理", wait_ms=2500)
    await snap("08_schedule",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 620})

    # ════════════════════════════════════════════════════════════
    # Tab 7：年度計劃表
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：年度計劃表 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("年度計劃表", wait_ms=2500)
    await snap("09_annual",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 600})

    # ════════════════════════════════════════════════════════════
    # Tab 8：批次清單
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：批次清單 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("批次清單", wait_ms=2000)
    await snap("10_batch_list",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 550})

    return shots


# ── Word 手冊輔助函式 ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def add_image_safe(doc, path, width_inches: float = 5.8):
    if isinstance(path, Path) and path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph("[ 截圖未能取得，請參考系統實際畫面 ]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = GRAY; p.runs[0].font.italic = True


def add_tip(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run("💡 提示："); r.bold = True; r.font.color.rgb = PRIMARY
    p.add_run(text)


def diff_box(doc, text: str, bg="E3F2FD"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    c = table.rows[0].cells[0]
    set_cell_bg(c, bg)
    c.paragraphs[0].add_run(text).font.size = Pt(10.5)
    doc.add_paragraph()


def hdr_table(doc, headers, bg="1B3A5C"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, bg)
    return table


def h1(doc, t):
    h = doc.add_heading(t, 1); h.runs[0].font.color.rgb = PRIMARY; return h

def h2(doc, t):
    h = doc.add_heading(t, 2); h.runs[0].font.color.rgb = ACCENT; return h


# ── Word 手冊建構 ─────────────────────────────────────────────────────────────
def build_word_manual(shots: dict, output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11.69); sec.page_width  = Inches(8.27)
    sec.top_margin  = Inches(0.9);   sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1.1);   sec.right_margin  = Inches(1.1)

    # ── 封面 ──────────────────────────────────────────────────────────────────
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_heading("全棟例行維護", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY; t.runs[0].font.size = Pt(28)
    t2 = doc.add_heading("教學操作手冊", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT; t2.runs[0].font.size = Pt(22)
    doc.add_paragraph()
    sub = doc.add_paragraph(
        f"路由：/mall/full-building-maintenance\n{datetime.now().strftime('%Y 年 %m 月 %d 日')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    doc.add_page_break()

    # ── 一、功能簡介 ──────────────────────────────────────────────────────────
    h1(doc, "一、功能簡介")
    p = doc.add_paragraph()
    p.add_run("全棟例行維護").bold = True
    p.add_run(
        "（路徑：商場管理 ▶ 全棟例行維護，路由：/mall/full-building-maintenance）"
        "管理整棟大樓的例行維護保養排程與執行記錄，"
        "架構與商場週期保養表幾乎相同，但有以下三點差異。"
    )
    doc.add_paragraph()

    diff_box(doc,
        "⭐ 與商場週期保養表的主要差異：\n"
        "1. Dashboard KPI 共六張（商場版五張），多了「保養時間」（實際工時）\n"
        "2. 每日巡檢表：只有單日 DatePicker 篩選，且目前標示「本地同步功能開發中」\n"
        "3. 年度計劃表格子多 no_frequency（∅）和 no_data（！）兩種狀態",
        bg="E3F2FD")

    tab_info = [
        ("Dashboard",   "KPI 六卡 + 圖表 + 預警（比商場版多一張「保養時間」）"),
        ("每日巡檢表",  "⚠ 本地同步開發中 — 單日 DatePicker 篩選（無年月選單）"),
        ("每月維護",    "月統計 + 年度矩陣 + 未完成說明（同商場版）"),
        ("每季維護",    "Q1-Q4 季度統計（同商場版）"),
        ("每年維護",    "年度統計 + 季度分布（同商場版）"),
        ("排程管理",    "排程管理 + 產生排程（同商場版）"),
        ("年度計劃表",  "12個月矩陣（比商場版多兩種格子狀態）"),
        ("批次清單",    "歷史批次列表（同商場版）"),
    ]
    table = hdr_table(doc, ["頁籤", "說明"])
    for name, desc in tab_info:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ── 二、Dashboard — KPI 六卡 ─────────────────────────────────────────────
    h1(doc, "二、Dashboard — KPI 六卡")
    add_image_safe(doc, shots.get("01_kpi_six"))
    doc.add_paragraph()
    diff_box(doc,
        "⭐ 全棟版 Dashboard 有六張 KPI 卡，商場版只有五張。\n"
        "差異在於第六張「保養時間」——顯示實際完成的保養工時（小時），"
        "對應 Ragic 欄位 actual_minutes / 60。",
        bg="E8F5E9")

    kpi6_rows = [
        ("有效項目",   "當月應執行的保養工項總數",           "深藍"),
        ("已完成",     "完成件數（含完成率%）",               "綠色"),
        ("逾期件數",   "過排定日期未完成的件數",              "深紅"),
        ("異常待追蹤", "有異常記錄需追蹤的件數",              "紫色"),
        ("預估工時",   "計劃工時（planned_minutes / 60）",   "藍色"),
        ("保養時間",   "實際保養工時（actual_minutes / 60）⭐全棟版獨有", "綠色"),
    ]
    table = hdr_table(doc, ["KPI", "說明", "顏色"])
    for name, desc, color in kpi6_rows:
        row = table.add_row()
        for i, v in enumerate([name, desc, color]):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    h2(doc, "2.1  圖表 & 預警區")
    add_image_safe(doc, shots.get("02_charts"))
    doc.add_paragraph("各類別完成率 Bar Chart + 狀態分布 Donut 圖（同商場版）。")
    doc.add_paragraph()
    add_image_safe(doc, shots.get("03_alerts"))
    doc.add_paragraph("逾期項目 Top 10 + 待執行項目 + 本月批次快速入口（同商場版）。")
    doc.add_page_break()

    # ── 三、每日巡檢表（開發中） ──────────────────────────────────────────────
    h1(doc, "三、每日巡檢表（整棟巡檢，目前開發中）")
    add_image_safe(doc, shots.get("04_daily_form"))
    doc.add_paragraph()
    diff_box(doc,
        "⚠ 重要說明：全棟例行維護的「每日巡檢表」Tab 目前尚未完成本地資料庫同步接通。\n"
        "頁面顯示藍色提示橫幅：「整棟巡檢本地同步功能開發中，"
        "目前每日巡檢表尚未接通本地 DB，模板欄位已備妥。"
        "請至 Ragic 系統填寫各樓層巡檢表單，接通後資料將自動顯示於此。」\n\n"
        "與商場週期保養表版本的另一差異：只有單日 DatePicker 篩選（無年月選單），"
        "每次查詢特定日期的巡檢記錄。",
        bg="FFF8E1")

    h2(doc, "3.1  目前可用功能")
    doc.add_paragraph("選擇日期 → 點「重新整理」→ 若 Ragic 有該日巡檢批次資料，會嘗試從後端取得並顯示。")
    doc.add_paragraph()

    h2(doc, "3.2  表格欄位（與商場版相同）")
    col_rows = [
        ("樓層",          "巡檢樓層（跨列合併）"),
        ("項目",          "巡檢設備/區域（同項目列合併）"),
        ("檢查內容",      "具體檢查項目說明"),
        ("運轉狀況（結果）","Tag：正常（綠）/ 異常（紅）/ 待處理（黃）/ 未巡檢（—）"),
        ("實際巡檢人員",  "執行巡檢的人員姓名"),
        ("異常說明",      "異常時的紅色詳細說明"),
        ("時間（分）",    "巡檢耗時分鐘數（同項目列合併）"),
    ]
    table = hdr_table(doc, ["欄位", "說明"])
    for name, desc in col_rows:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ── 四、每月/季/年維護（同商場版） ────────────────────────────────────────
    h1(doc, "四、每月 / 每季 / 每年維護統計（同商場版）")
    doc.add_paragraph("Tab 3、4、5 功能與商場週期保養表完全相同，以下為截圖確認：")
    doc.add_paragraph()

    h2(doc, "4.1  每月維護")
    add_image_safe(doc, shots.get("05_monthly"))
    doc.add_paragraph("年度矩陣 + 單月鑽取（上月累計 / 本月統計 / 未完成說明）。")
    doc.add_paragraph()

    h2(doc, "4.2  每季維護")
    add_image_safe(doc, shots.get("06_quarterly"))
    doc.add_paragraph("Q1-Q4 選擇卡片 + 季度 KPI。")
    doc.add_paragraph()

    h2(doc, "4.3  每年維護")
    add_image_safe(doc, shots.get("07_yearly"))
    doc.add_paragraph("全年 KPI + 季度分布表。")
    doc.add_page_break()

    # ── 五、排程管理（同商場版） ─────────────────────────────────────────────
    h1(doc, "五、排程管理（同商場版）")
    add_image_safe(doc, shots.get("08_schedule"))
    doc.add_paragraph()
    doc.add_paragraph(
        "篩選工具列 + 產生排程按鈕 + 排程表格，操作方式與商場版相同。\n"
        "排程狀態：已完成（綠）/ 進行中（藍）/ 逾期（深紅）/ 待執行（橙）/ 未排定（黃）。"
    )
    add_tip(doc, "每月初點「產生排程」，為當月應執行的全棟保養項目批次建立排程記錄。")
    doc.add_page_break()

    # ── 六、年度計劃表（比商場版多兩種格子狀態）─────────────────────────────
    h1(doc, "六、年度計劃表")
    add_image_safe(doc, shots.get("09_annual"))
    doc.add_paragraph()
    diff_box(doc,
        "⭐ 全棟版的年度計劃表比商場版多兩種格子狀態：\n"
        "• ∅（no_frequency）：該保養項目無此月份的頻率設定，不需保養，淺灰底色\n"
        "• ！（no_data）：應保養但尚無資料，粉紅底色，需至 Ragic 確認",
        bg="E3F2FD")

    h2(doc, "年度計劃表格子狀態完整說明")
    annual_rows = [
        ("✅", "completed",    "已完成",   "綠底 #f6ffed"),
        ("🔴", "overdue",      "逾期",     "紅底 #fff1f0"),
        ("🔵", "in_progress",  "進行中",   "藍底 #e6f4ff"),
        ("⭕", "scheduled",    "待執行",   "橙底 #fff7e6"),
        ("?",  "unscheduled",  "未排定",   "黃底 #fffbe6"),
        ("─",  "non_month",   "非本月",   "淡灰底 #fafafa"),
        ("！", "no_data",      "應保養但無資料 ⭐", "粉紅底 #fff0f6"),
        ("∅",  "no_frequency", "無此月份頻率 ⭐", "淺灰底 #f5f5f5"),
    ]
    table = hdr_table(doc, ["圖示", "狀態碼", "說明", "格子顏色"])
    for icon, code, desc, color in annual_rows:
        row = table.add_row()
        for i, v in enumerate([icon, code, desc, color]):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    add_tip(doc, "點擊任一格子，右側滑出 Drawer 明細。若格子顯示「！」警示，表示應執行月份尚無資料，需至 Ragic 確認或建立排程。")
    doc.add_page_break()

    # ── 七、批次清單 ──────────────────────────────────────────────────────────
    h1(doc, "七、批次清單（同商場版）")
    add_image_safe(doc, shots.get("10_batch_list"))
    doc.add_paragraph()
    doc.add_paragraph(
        "歷史保養批次列表，欄位同商場版：保養單號（可點跳轉）、保養月份、"
        "批次狀態 Tag、完成率進度條、逾期件數（紅色）、異常件數（紫色）。"
    )
    doc.add_page_break()

    # ── 八、常見操作情境 ──────────────────────────────────────────────────────
    h1(doc, "八、常見操作情境")
    scenarios = [
        ("確認本月全棟保養的實際完成工時",
         ["Dashboard Tab → 查看第六張 KPI 卡「保養時間」",
          "與「預估工時」比較：實際 < 預估可能代表有項目未完成",
          "查看逾期件數確認是否有漏做的項目"]),
        ("月初建立全棟保養排程",
         ["排程管理 Tab → 選本月 → 點「產生排程」",
          "回到 Dashboard 確認有效項目卡片數量已更新"]),
        ("查看年度計劃表並確認 ！ 警示",
         ["切換到「年度計劃表」Tab",
          "找到顯示「！」的格子（粉紅底色）",
          "點擊格子確認 Drawer 提示內容，確認是否需補建排程或前往 Ragic 處理"]),
        ("確認某月的批次完成狀態",
         ["切換到「批次清單」Tab → 選目標年份",
          "找到對應月份，查看完成率進度條和逾期件數",
          "點保養單號進入批次詳細頁逐項確認"]),
    ]
    for i, (title, steps) in enumerate(scenarios, 1):
        p = doc.add_paragraph()
        p.add_run(f"情境 {i}：{title}").bold = True
        p.runs[0].font.color.rgb = PRIMARY
        for step in steps:
            doc.add_paragraph(f"　→ {step}")
        doc.add_paragraph()

    # ── 附錄 ──────────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "附錄：與商場週期保養表差異彙整")
    table = hdr_table(doc, ["功能", "商場週期保養表", "全棟例行維護"])
    diff_rows = [
        ("路由",          "/mall/periodic-maintenance",  "/mall/full-building-maintenance"),
        ("Dashboard KPI", "五張（無保養時間）",           "六張（含「保養時間」實際工時）"),
        ("每日巡檢表",    "月份模式 + 單日模式（已接通）", "單日模式（本地同步開發中）"),
        ("年度計劃表格子", "6 種狀態",                    "8 種狀態（+no_data + no_frequency）"),
        ("其他 Tab",      "完全相同",                     "完全相同"),
    ]
    for func, mall, full in diff_rows:
        row = table.add_row()
        row.cells[0].text = func; row.cells[1].text = mall; row.cells[2].text = full
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 手冊已儲存：{output_path}")


# ── 主程式 ────────────────────────────────────────────────────────────────────
async def main():
    OUTPUT_DIR.mkdir(exist_ok=True)
    print("=" * 55)
    print("  全棟例行維護 教學手冊生成器")
    print("=" * 55)
    print(f"目標：{TARGET_URL}")
    print(f"截圖：{OUTPUT_DIR}")
    print(f"手冊：{MANUAL_PATH}\n")

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=["--disable-blink-features=AutomationControlled"],
        )
        context = await browser.new_context(
            viewport={"width": 1600, "height": VIEWPORT_H},
            device_scale_factor=1.5,
            locale="zh-TW",
            timezone_id="Asia/Taipei",
        )
        page = await context.new_page()
        shots = await capture_screenshots(page, OUTPUT_DIR)
        await browser.close()

    print(f"\n截圖完成，共 {len(shots)} 張")
    print("\n📄 生成 Word 手冊...")
    build_word_manual(shots, MANUAL_PATH)
    print(f"\n✅ 執行完成！請開啟：{MANUAL_PATH}")


if __name__ == "__main__":
    asyncio.run(main())
