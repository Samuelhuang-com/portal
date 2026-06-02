#!/usr/bin/env python3
"""
商場管理 Dashboard（/mall/overview）影音教學旁白腳本生成器
=========================================================
使用方式：
  python mall_overview_generate_video_scripts.py

輸出：商場管理Dashboard_影音教學腳本.docx
"""

import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx"); sys.exit(1)

OUTPUT_PATH = Path(__file__).parent / "商場管理Dashboard_影音教學腳本.docx"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
ORANGE  = RGBColor(0xFA, 0xAD, 0x14)
GRAY    = RGBColor(0x59, 0x59, 0x59)


# ── 輔助函式 ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def narration_block(doc, text: str):
    for part in text.split("\n\n"):
        part = part.strip()
        if not part:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(part); r.font.size = Pt(11.5)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "12"); left.set(qn("w:color"), "4BA8E8")
        pBdr.append(left); pPr.append(pBdr)


def label_box(doc, label: str, text: str, fill_hex: str):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    lc = table.rows[0].cells[0]
    set_cell_bg(lc, fill_hex)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label); lr.bold = True; lr.font.size = Pt(10)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = table.rows[0].cells[1]
    set_cell_bg(rc, "F8F9FA")
    rc.paragraphs[0].add_run(text).font.size = Pt(10.5)
    doc.add_paragraph()


def action_block(doc, text: str):  label_box(doc, "🖱 操作", text, "52C41A")
def hint_block(doc, text: str):    label_box(doc, "⏸ 提示", text, "FA8C16")


def timing_row(doc, t_range: str, label: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"[{t_range}]  ")
    r1.bold = True; r1.font.color.rgb = ORANGE; r1.font.size = Pt(10)
    r2 = p.add_run(label)
    r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = PRIMARY


def section_divider(doc):
    p = doc.add_paragraph("─" * 64)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── 腳本資料 ──────────────────────────────────────────────────────────────────
SCRIPTS = [

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第一集",
        "title": "進入系統 & 篩選列 & 7 來源卡片",
        "duration": "建議時長：4 分鐘",
        "goal": "讓觀眾了解商場管理 Dashboard 的六個頁籤架構、篩選列操作，以及兩排共 7 個來源卡片各自代表什麼。",
        "segments": [
            {
                "time": "00:00–00:30",
                "label": "功能簡介",
                "narration": (
                    "大家好，歡迎收看商場管理 Dashboard 教學系列。"
                    "這個 Dashboard 整合了商場所有工務作業的即時狀態，"
                    "把七個不同的工務來源集中在一個畫面，"
                    "讓管理者一眼掌握整體的健康度。"
                    "\n\n"
                    "頁面有六個頁籤：A Dashboard 是預設頁籤，"
                    "B 每日累計、C 每月累計、D 每年累計是工時分析，"
                    "最後是人員工時% 和人員排名。"
                    "這集先從 Tab A 開始，了解篩選列和 7 個來源卡片。"
                ),
                "actions": ["進入頁面，指向六個頁籤"],
                "hints": [],
            },
            {
                "time": "00:30–01:15",
                "label": "篩選列操作",
                "narration": (
                    "Tab A 頂部的篩選列有幾個重要控制項。"
                    "\n\n"
                    "年度和月份選單控制七個來源卡片的資料範圍，"
                    "改了之後各來源會各自重新載入，通常需要三到八秒。"
                    "月份可以選「全年」看全年彙整，或選特定月份看當月資料。"
                    "\n\n"
                    "旁邊有一個巡檢日期的 DatePicker，"
                    "這個專門控制「商場工務巡檢」的日別統計——"
                    "選不同的日期，巡檢卡片的資料就會更新到那一天的狀況。"
                    "旁邊的「今日」按鈕可以快速切回今天。"
                    "\n\n"
                    "右上角的紫色「匯出 PowerPoint」按鈕，"
                    "可以把當前的 KPI 資料匯出為簡報檔，第五集會介紹。"
                ),
                "actions": [
                    "示範切換月份，指向各卡片更新",
                    "指向巡檢日期 DatePicker",
                    "指向匯出 PowerPoint 按鈕",
                ],
                "hints": ["建議選有完整資料的近期月份"],
            },
            {
                "time": "01:15–02:30",
                "label": "彙總 KPI 列",
                "narration": (
                    "七個來源卡片上方是五張彙總 KPI 卡，"
                    "匯總四大主要來源（商場例行維護、全棟例行維護、商場工務巡檢、商場工務報修）的數字。"
                    "\n\n"
                    "「本期總工項」是四個來源的工項數加總，深藍色。"
                    "「已完成工項」是四個來源的完成件數加總，綠色。"
                    "「本期工時合計」旁邊有問號圖示，停留後說明各來源的工時計算口徑。"
                    "「異常/未完成」大於零顯示紅色，等於零顯示綠色加「全部正常」標籤。"
                    "「逾期未完成」主要來自例行維護的逾期項目，大於零顯示深紅色。"
                ),
                "actions": [
                    "指向五張 KPI 卡依序說明",
                    "停留在工時合計的問號圖示顯示 Tooltip",
                ],
                "hints": [],
            },
            {
                "time": "02:30–04:00",
                "label": "7 來源狀態卡片",
                "narration": (
                    "7 個來源卡片分兩排排列。"
                    "\n\n"
                    "第一排是維護和巡檢類，4 張：商場例行維護（深藍）、全棟例行維護（藍色）、"
                    "商場工務巡檢（紫色）、整棟巡檢（綠色）。"
                    "\n\n"
                    "第二排是報修和任務類，3 張：商場工務報修（橙色）、商場主管交辦（深紅）、商場緊急事件（紅色）。"
                    "\n\n"
                    "每張卡片顯示：工項件數、已完成件數、完成率（%），以及異常件數和逾期件數。"
                    "\n\n"
                    "要特別說明的是主管交辦和緊急事件這兩張——"
                    "這兩類任務通常不計完成率，所以卡片上只顯示件數和工時，沒有完成率數字。"
                    "\n\n"
                    "點擊任何一張卡片，都可以跳轉到對應的詳細清單頁面。"
                    "\n\n"
                    "好，這集七個來源卡片介紹完了。下一集看圖表區和費用摘要。"
                ),
                "actions": [
                    "逐一指向兩排卡片",
                    "指向主管交辦和緊急事件卡片說明特殊性",
                    "點一張卡片示範跳轉，再按上一頁回來",
                ],
                "hints": ["建議選有警示（橙色或紅色）的月份，讓卡片顏色對比更清楚"],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第二集",
        "title": "圖表區 & 費用摘要 & 匯出 PowerPoint",
        "duration": "建議時長：3 分鐘",
        "goal": "讓觀眾學會解讀四個圖表和費用摘要，以及匯出 PowerPoint 的操作流程。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "這集繼續看 Tab A 的圖表區和費用摘要，以及 PowerPoint 匯出功能。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:40",
                "label": "四個圖表",
                "narration": (
                    "來源卡片下方是四個圖表，我們依序看。"
                    "\n\n"
                    "左上角：「各來源工項數比較」水平長條圖——"
                    "每列是一個來源，深色長條是總數，淺色是已完成數，"
                    "可以快速比對各來源的工作量大小和完成進度。"
                    "\n\n"
                    "右上角：「各來源完成率」水平長條圖——"
                    "同樣一個來源一列，顯示完成率百分比。"
                    "維護類的完成率通常高一些，報修類可能較低，"
                    "對比兩張圖可以同時掌握量和率。"
                    "\n\n"
                    "左下角：「工時來源占比」圓餅圖——"
                    "顯示各來源工時的比例，停留在各區塊可看到精確工時數和佔比。"
                    "\n\n"
                    "右下角：「商場工務報修 12 個月趨勢」折線圖——"
                    "深色線是報修件數，淺色線是結案件數，"
                    "觀察兩條線的差距可以判斷積案是否在增加。"
                ),
                "actions": [
                    "停留在各圖表顯示 Tooltip",
                    "指向折線圖說明兩條線的差距意義",
                ],
                "hints": [],
            },
            {
                "time": "01:40–02:20",
                "label": "費用摘要",
                "narration": (
                    "圖表下方是費用摘要，資料來自商場工務報修（陸群）。"
                    "顯示委外加維修費用的年度累計、扣款費用，以及當月金額小計。"
                    "\n\n"
                    "如果需要更詳細的費用分析，可以點擊「商場工務報修」來源卡片，"
                    "跳轉到報修模組的詳細頁面，那裡有更完整的費用 KPI 和明細。"
                ),
                "actions": ["指向費用摘要區塊"],
                "hints": [],
            },
            {
                "time": "02:20–03:00",
                "label": "匯出 PowerPoint",
                "narration": (
                    "頁頭右側有一個紫色的「匯出 PowerPoint」按鈕。"
                    "\n\n"
                    "使用前要注意：月份必須選定特定月份，不能選「全年」，否則按鈕無效。"
                    "\n\n"
                    "確認月份後點擊，系統開始生成 .pptx 簡報，"
                    "進度條顯示生成進度，完成後瀏覽器自動下載。"
                    "\n\n"
                    "PPT 內容會包含 KPI 摘要和各來源的狀態資料，"
                    "適合直接拿去月報或向主管匯報。"
                    "\n\n"
                    "好，Tab A 介紹完了。下一集看三個累計分析 Tab。"
                ),
                "actions": [
                    "確認月份已選特定月份",
                    "點「匯出 PowerPoint」展示進度條",
                ],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第三集",
        "title": "Tab B/C/D 每日 / 每月 / 每年累計",
        "duration": "建議時長：4 分鐘",
        "goal": "讓觀眾掌握三個累計分析 Tab 的用途、篩選操作方式，以及工時表格的閱讀方式。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "這集介紹三個工時累計分析 Tab：B 每日累計、C 每月累計、D 每年累計。"
                    "這三個 Tab 都有各自獨立的篩選器，和 Tab A 的篩選互不影響。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:40",
                "label": "Tab B — 每日累計",
                "narration": (
                    "切換到「B. 每日累計」Tab。"
                    "第一次切換時系統才開始載入資料，通常需要一到兩秒。"
                    "\n\n"
                    "這張表格以商場工務報修的每日資料為基礎，"
                    "列是五大工項類別——現場報修、上級交辦、緊急事件、例行維護、每日巡檢，"
                    "欄是當月每一天（日期加星期）。"
                    "\n\n"
                    "每個格子顯示當天的案件件數，零值顯示橫線。"
                    "最右邊是月合計和占比。"
                    "\n\n"
                    "左上角年月選單可以切換查詢月份，"
                    "這裡的篩選是獨立的，改了不會影響 Tab A 的資料。"
                    "\n\n"
                    "工時的計算口徑以「occupied_at」欄位歸屬日期，"
                    "和 Tab A 商場工務報修卡片的計算口徑一致。"
                ),
                "actions": [
                    "切換到「B. 每日累計」Tab",
                    "等待資料載入",
                    "示範切換月份篩選",
                ],
                "hints": [],
            },
            {
                "time": "01:40–02:40",
                "label": "Tab C — 每月累計",
                "narration": (
                    "切換到「C. 每月累計」Tab。"
                    "\n\n"
                    "這張矩陣以一整年為視角，"
                    "列是五大工項類別，欄是 1 月到 12 月。"
                    "每個格子顯示該月份的工時數。"
                    "\n\n"
                    "未來月份的格子顯示橫線。"
                    "可以點欄位標題排序，快速找出工時最高的月份或類別。"
                    "\n\n"
                    "左上角年份選單控制要查看哪個年度，"
                    "可以切換到不同年份做跨年比較。"
                    "\n\n"
                    "這個 Tab 很適合做年度工務量分布的分析，"
                    "找出哪個月份工作量特別集中，或者哪個工項類別全年佔比最大。"
                ),
                "actions": [
                    "切換到「C. 每月累計」Tab",
                    "指向未來月份橫線",
                    "切換年份示範",
                ],
                "hints": [],
            },
            {
                "time": "02:40–03:20",
                "label": "Tab D — 每年累計",
                "narration": (
                    "切換到「D. 每年累計」Tab。"
                    "\n\n"
                    "這個視角以年度為單位，顯示各月份工時的累計走勢（Running Total），"
                    "適合觀察全年工務量的月份分布規律。"
                    "\n\n"
                    "例如可以看出哪幾個月工作量特別重，"
                    "或者旺季和淡季的工時差異有多大。"
                    "切換年份可以比較不同年度的走勢。"
                    "\n\n"
                    "好，三個累計 Tab 介紹完了。下一集看人員工時分析。"
                ),
                "actions": [
                    "切換到「D. 每年累計」Tab",
                    "切換年份示範比較",
                ],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第四集",
        "title": "人員工時% & 人員排名 & 常見情境",
        "duration": "建議時長：4 分鐘",
        "goal": "讓觀眾掌握人員工時分析的兩個 Tab，並示範四個最常用的操作情境。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場",
                "narration": (
                    "這是商場管理 Dashboard 教學系列的最後一集。"
                    "我們來看人員工時分析的兩個 Tab，以及常見操作情境。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:20",
                "label": "人員工時%",
                "narration": (
                    "切換到「人員工時%」Tab。"
                    "\n\n"
                    "這張表格是熱度分析視圖——"
                    "列是五大工項類別（現場報修、上級交辦、緊急事件、例行維護、每日巡檢），"
                    "欄是工時排名前 15 的人員。"
                    "\n\n"
                    "每個格子顯示「這位人員在這個類別的工時，"
                    "佔他個人總工時的百分比」。"
                    "\n\n"
                    "顏色的意義：紅色大於等於三十%，代表這個人高度集中在這個類別；"
                    "橙色是十五到三十%；綠色是有貢獻但比例較低；灰色橫線代表沒有工時記錄。"
                    "\n\n"
                    "左上角年份選單控制查詢年度。"
                    "這個 Tab 和「人員排名」共用同一組資料，"
                    "切換到其中一個後，另一個也不需要再等待。"
                ),
                "actions": [
                    "切換到「人員工時%」Tab",
                    "指向紅色格子說明意義",
                ],
                "hints": [],
            },
            {
                "time": "01:20–02:10",
                "label": "人員排名",
                "narration": (
                    "切換到「人員排名」Tab。"
                    "\n\n"
                    "頁面分兩部分：上方是堆疊長條圖，下方是排名表格。"
                    "\n\n"
                    "堆疊長條圖以人員為橫軸，顯示每位人員的工時來源結構——"
                    "橙色是現場報修、深藍是例行維護、紫色是每日巡檢。"
                    "長條越長代表工時越多，不同顏色的占比代表工作類型的分布。"
                    "\n\n"
                    "下方排名表格以工時降序排列，"
                    "欄位包含排名、姓名、總工時、占比，以及各類別工時的分解。"
                    "\n\n"
                    "搭配使用：先在排名找到工時最重的人，"
                    "再切回「人員工時%」找到那個人的欄位，"
                    "看看他的工時集中在哪個作業類別。"
                ),
                "actions": [
                    "切換到「人員排名」Tab",
                    "指向堆疊長條圖的各顏色說明",
                    "指向排名表格第一名",
                ],
                "hints": [],
            },
            {
                "time": "02:10–02:45",
                "label": "情境示範 1：月初健康快速確認",
                "narration": (
                    "情境一：月初快速確認商場工務整體狀況。"
                    "\n\n"
                    "進入頁面，年月選本月，等待 7 張卡片全部載入。"
                    "掃視卡片顏色：全部綠色代表良好；有橙色或紅色的需要追蹤。"
                    "\n\n"
                    "同時確認彙總 KPI 的「異常/未完成」和「逾期未完成」是否為零，"
                    "如果不是零，點出問題的來源卡片跳轉到詳細清單處理。"
                ),
                "actions": ["示範月份切換 + 掃視卡片"],
                "hints": [],
            },
            {
                "time": "02:45–03:15",
                "label": "情境示範 2：找出工時最重的人員",
                "narration": (
                    "情境二：主管想了解本年度哪位工務人員工時最多。"
                    "\n\n"
                    "切換到「人員排名」Tab，第一名就是工時最多的人員。"
                    "再切到「人員工時%」找到同一個人的欄位，"
                    "看他的工時分布在哪些作業類別，"
                    "判斷是否工作類型太集中需要調配。"
                ),
                "actions": [
                    "人員排名 Tab 找第一名",
                    "切到人員工時% 找同一人",
                ],
                "hints": [],
            },
            {
                "time": "03:15–03:45",
                "label": "結尾",
                "narration": (
                    "商場管理 Dashboard 的四集教學到這邊全部結束了。"
                    "\n\n"
                    "快速回顧：第一集學七個來源卡片和篩選列，"
                    "第二集學圖表區、費用摘要和 PPT 匯出，"
                    "第三集學每日/每月/每年三個累計 Tab，"
                    "第四集學人員工時分析和操作情境。"
                    "\n\n"
                    "這個 Dashboard 和飯店管理 Dashboard 的架構高度對應，"
                    "學會一個，另一個也能快速上手。謝謝收看！"
                ),
                "actions": ["畫面停在 Tab A Dashboard 做結尾"],
                "hints": [],
            },
        ],
    },
]


# ── Word 文件生成 ─────────────────────────────────────────────────────────────
def build_doc(scripts: list, output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height   = Inches(11.69); sec.page_width    = Inches(8.27)
    sec.top_margin    = Inches(0.85);  sec.bottom_margin = Inches(0.85)
    sec.left_margin   = Inches(1.1);   sec.right_margin  = Inches(1.1)

    # 封面
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_heading("商場管理 Dashboard", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY; t.runs[0].font.size = Pt(26)
    t2 = doc.add_heading("影音教學旁白腳本", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT; t2.runs[0].font.size = Pt(20)
    doc.add_paragraph()
    sub = doc.add_paragraph(
        f"共四集 · 路由：/mall/overview\n{datetime.now().strftime('%Y 年 %m 月 %d 日')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    doc.add_paragraph()

    # 使用說明
    doc.add_heading("腳本使用說明", 1).runs[0].font.color.rgb = PRIMARY
    for tip in [
        "🎬 旁白：錄製時直接念，可依個人習慣調整，不必逐字照念。",
        "🖱 操作：錄製時同步要做的滑鼠動作。",
        "⏸ 提示：錄製前的準備建議，不需要說出來。",
        "[時間軸]：僅供參考，依實際節奏調整。",
        "建議選有各種顏色卡片（有警示）的月份，視覺示範效果最佳。",
        "若已錄製過「飯店管理 Dashboard」教學，可在開頭說明兩者架構對應，縮短說明時間。",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(tip).font.size = Pt(10.5)

    doc.add_page_break()

    for script in scripts:
        h = doc.add_heading(f"{script['ep']}  {script['title']}", 1)
        h.runs[0].font.color.rgb = PRIMARY; h.runs[0].font.size = Pt(16)

        info = doc.add_paragraph()
        r1 = info.add_run(f"⏱ {script['duration']}　　")
        r1.font.size = Pt(10); r1.font.color.rgb = ORANGE; r1.bold = True

        gp = doc.add_paragraph()
        gp.paragraph_format.space_after = Pt(8)
        gp.add_run("學習目標：").bold = True
        gp.runs[0].font.color.rgb = ACCENT
        gp.add_run(script["goal"])

        section_divider(doc)

        for seg in script["segments"]:
            timing_row(doc, seg["time"], seg["label"])
            narration_block(doc, seg["narration"])
            for act in seg.get("actions", []):
                action_block(doc, act)
            for hint in seg.get("hints", []):
                hint_block(doc, hint)
            doc.add_paragraph()

        doc.add_page_break()

    # 附錄
    doc.add_heading("附錄：7 來源顏色與跳轉速查", 1).runs[0].font.color.rgb = PRIMARY
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["來源名稱", "顏色", "跳轉路徑"]):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1B3A5C")

    rows_data = [
        ("商場例行維護", "深藍 #1B3A5C",  "/mall/periodic-maintenance"),
        ("全棟例行維護", "藍色 #4BA8E8",  "/mall/full-building-maintenance"),
        ("商場工務巡檢", "紫色 #722ED1",  "/mall-facility-inspection/dashboard"),
        ("整棟巡檢",     "綠色 #52C41A",  "/full-building-inspection/dashboard"),
        ("商場工務報修", "橙色 #FA8C16",  "/luqun-repair/dashboard"),
        ("商場主管交辦", "深紅 #C0392B",  "/mall/other-tasks"),
        ("商場緊急事件", "紅色 #D4380D",  "/mall/other-tasks"),
    ]
    for name, color, route in rows_data:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = color; row.cells[2].text = route
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 腳本手冊已儲存：{output_path}")


if __name__ == "__main__":
    build_doc(SCRIPTS, OUTPUT_PATH)
    print(f"執行完成！請開啟：{OUTPUT_PATH}")
