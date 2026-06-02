#!/usr/bin/env python3
"""
集團決策 Dashboard 影音教學旁白腳本生成器
=========================================
使用方式（已安裝 python-docx 即可直接執行）：
  python generate_video_scripts.py

輸出：集團決策Dashboard_影音教學腳本.docx
"""

import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx")
    sys.exit(1)

OUTPUT_PATH = Path(__file__).parent / "集團決策Dashboard_影音教學腳本.docx"

PRIMARY  = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT   = RGBColor(0x4B, 0xA8, 0xE8)
SUCCESS  = RGBColor(0x52, 0xC4, 0x1A)
WARNING  = RGBColor(0xFA, 0xAD, 0x14)
DANGER   = RGBColor(0xCF, 0x13, 0x22)
GRAY     = RGBColor(0x59, 0x59, 0x59)
PURPLE   = RGBColor(0x72, 0x2E, 0xD1)
BG_LIGHT = "EBF4FB"
BG_HINT  = "FFF8E6"
BG_ACT   = "F0F9F0"


# ── 輔助函式 ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading1(doc, text: str):
    h = doc.add_heading(text, 1)
    h.runs[0].font.color.rgb = PRIMARY
    return h


def add_heading2(doc, text: str):
    h = doc.add_heading(text, 2)
    h.runs[0].font.color.rgb = ACCENT
    return h


def add_heading3(doc, text: str):
    h = doc.add_heading(text, 3)
    h.runs[0].font.color.rgb = GRAY
    return h


def label_box(doc, label: str, text: str, bg: str, label_color: RGBColor):
    """帶彩色標籤的提示框（用表格模擬）"""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_bg(table.rows[0].cells[0], "1B3A5C" if label_color == PRIMARY else
                "4BA8E8" if label_color == ACCENT else
                "FFF8E6" if bg == BG_HINT else
                "F0F9F0")
    lc = table.rows[0].cells[0]
    lc.width = Inches(0.9)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rc = table.rows[0].cells[1]
    set_cell_bg(rc, bg)
    rp = rc.paragraphs[0]
    rp.add_run(text).font.size = Pt(10.5)
    doc.add_paragraph()


def narration_block(doc, text: str):
    """旁白正文（帶左藍線裝飾）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    # 左側藍線用段落邊框
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "4BA8E8")
    pBdr.append(left)
    pPr.append(pBdr)
    return p


def action_block(doc, text: str):
    """操作提示（綠底）"""
    label_box(doc, "🖱 操作", text, BG_ACT, SUCCESS)


def pause_block(doc, text: str):
    """暫停提示（黃底）"""
    label_box(doc, "⏸ 提示", text, BG_HINT, WARNING)


def timing_row(doc, minutes: str, label: str):
    """時間軸標示列"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"[{minutes}]  ")
    r1.bold = True
    r1.font.color.rgb = WARNING
    r1.font.size = Pt(10)
    r2 = p.add_run(label)
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY


def section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 62)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(9)


# ── 腳本內容 ──────────────────────────────────────────────────────────────────

SCRIPTS = [

    # ══════════════════════════════════════════════════════════════════════════
    {
        "ep": "第一集",
        "title": "進入系統 & KPI 指標解讀",
        "duration": "建議時長：3–4 分鐘",
        "goal": "讓觀眾知道如何進入頁面，並快速讀懂 16 個 KPI 卡片的意義與顏色警示。",
        "segments": [
            {
                "time": "00:00–00:30",
                "label": "開場白",
                "narration": (
                    "大家好，歡迎收看集團決策 Dashboard 教學系列。"
                    "這套系統整合了飯店工務部和商場工務報修的即時數據，"
                    "讓管理層可以在一個畫面裡掌握全集團的工務狀態。"
                    "這集我們先從「怎麼進入」和「KPI 卡片怎麼看」開始。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:30–01:00",
                "label": "進入頁面",
                "narration": (
                    "登入系統之後，可以看到左側的側邊欄。"
                    "最上方有一個標星號的選項——「★ 集團決策 Dashboard」，"
                    "這是我們這次的主角。點一下，系統就會開始自動載入最新的資料。"
                    "通常需要等個三到五秒，右上角出現「更新於某時某分」的字樣，"
                    "就代表資料已經全部載入完成了。"
                ),
                "actions": ["點選左側側邊欄「★ 集團決策 Dashboard」", "等待右上角出現更新時間戳"],
                "hints": ["建議錄製前先確認後端服務已啟動（uvicorn 正在運行）"],
            },
            {
                "time": "01:00–01:30",
                "label": "頁面三個頁籤",
                "narration": (
                    "頁面頂部有三個頁籤：集團工務概覽、工作日誌、統計基準說明。"
                    "這集我們先聚焦在第一個——集團工務概覽。"
                    "其他兩個我們後面的集數會分別介紹。"
                ),
                "actions": ["指向頁籤列，依序點一下各頁籤讓觀眾知道位置"],
                "hints": [],
            },
            {
                "time": "01:30–02:30",
                "label": "第一列 KPI：核心工務指標",
                "narration": (
                    "現在來看 KPI 卡片。第一列一共有八個指標，我們從左到右看過去。"
                    "\n\n"
                    "「本月總案件」——飯店加商場的當月報修案件總數，顯示深藍色。"
                    "\n\n"
                    "「本月總工時」——特別注意，這裡的工時只計算有填寫工時記錄的案件，"
                    "旁邊有個 ℹ 圖示，停留在上面就會看到詳細說明。"
                    "\n\n"
                    "「完成件數」固定是綠色，「未完成件數」如果大於零就會變紅色，等於零才是綠色。"
                    "\n\n"
                    "「完成率」這張卡最重要——大於等於八成是綠色代表良好，"
                    "五到八成之間是橙色要注意，低於五成就變紅色，需要立刻關注。"
                    "\n\n"
                    "最右邊兩張是「飯店案件占比」和「商場案件占比」，"
                    "一眼就能知道工務量的分布重心在哪邊。"
                ),
                "actions": [
                    "滑鼠停留在各 KPI 卡片上，讓觀眾看清楚數字",
                    "停留在「本月總工時」ℹ 圖示上，展示 Tooltip",
                ],
                "hints": ["建議選一個有異常（未完成件數 > 0）的月份，讓顏色警示更清楚"],
            },
            {
                "time": "02:30–03:00",
                "label": "第二列 KPI：待辦驗 & 上期未結",
                "narration": (
                    "第二列四張卡片是追蹤歷史問題用的。"
                    "\n\n"
                    "「待辦驗數」代表已經修好、但還在等客戶驗收確認的件數，"
                    "大於零會亮橙色——這些案件需要主動去催客戶簽認。"
                    "\n\n"
                    "「上期未結」就更嚴重了，是上個月就沒結案、拖到這個月的件數，"
                    "大於零會亮紅色，是需要優先追蹤的清單。"
                ),
                "actions": ["放大顯示第二列四張卡片"],
                "hints": [],
            },
            {
                "time": "03:00–03:30",
                "label": "第三列 KPI：上級交辦 & 緊急事件",
                "narration": (
                    "第三列是針對主管交辦任務和緊急事件的統計。"
                    "紫色是上級交辦，紅色是緊急事件，分別顯示件數和工時。"
                    "這兩類任務是跨飯店和商場的，"
                    "在後面的工作日誌章節還會再更詳細地介紹查詢方式。"
                    "\n\n"
                    "好，這集的 KPI 解讀到這邊。下一集我們來看工務報修摘要和趨勢圖。"
                ),
                "actions": ["指向第三列，展示紫色和紅色卡片"],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "ep": "第二集",
        "title": "工務報修摘要 & 趨勢圖解讀",
        "duration": "建議時長：3–4 分鐘",
        "goal": "讓觀眾學會切換年月查看歷史數據、解讀摘要卡片與趨勢圖。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "上一集我們看了 KPI 卡片的意義。"
                    "這集繼續往下，來看工務報修摘要區塊，以及近十二個月的趨勢圖。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–00:50",
                "label": "年月篩選器",
                "narration": (
                    "往下滾可以看到「工務報修」這個區塊，旁邊標示了目前顯示的月份。"
                    "左邊有兩個下拉選單——年份和月份。"
                    "切換這兩個選單，底下的摘要卡片就會跟著更新，"
                    "讓你比較不同月份的報修狀況。"
                    "要注意，這個篩選只影響工務報修區塊，KPI 卡片還是顯示當月。"
                ),
                "actions": ["示範切換到上個月，讓卡片數字更新"],
                "hints": ["若切換月份後數字沒有變化，等一到兩秒再確認"],
            },
            {
                "time": "00:50–02:00",
                "label": "飯店工務部摘要卡",
                "narration": (
                    "左邊是飯店工務部的摘要卡。"
                    "先看最上面一排數字：報修總數、已結案、待辦驗數、未結案、結案率、"
                    "平均結案天數、工時。"
                    "\n\n"
                    "結案率下面有一條進度條，顏色和上面的百分比一致。"
                    "如果有最高報修類型，會顯示藍色的 Tag，例如「↑ 建築 16 件」，"
                    "讓你一眼知道這個月哪一類型的報修最多。"
                    "\n\n"
                    "如果有超過十四天都還沒結案的，會顯示紅色的警示 Tag。"
                    "\n\n"
                    "卡片底部最重要——有一句話的智慧摘要，"
                    "系統自動判斷目前狀況然後用一句話說出來，"
                    "比如「目前 210 件未結案，主要集中於建築類型」。"
                    "\n\n"
                    "再往下是當月費用小計，包含委外維修費、扣款費用、當月小計。"
                    "\n\n"
                    "右上角有「查看詳情 ›」的按鈕，點下去會直接跳轉到飯店工務詳細清單頁面，"
                    "不需要再去側邊欄找。"
                ),
                "actions": [
                    "放大飯店工務部卡片，逐一指向各數字",
                    "點一下「查看詳情 ›」示範跳轉，再按上一頁回來",
                ],
                "hints": ["建議選有異常（未結案 > 0、有警示 Tag）的月份，說服力更強"],
            },
            {
                "time": "02:00–02:40",
                "label": "商場工務報修摘要 & 工項類別比較表",
                "narration": (
                    "中間是商場工務報修，結構和飯店完全一樣，這邊就不重複說明了。"
                    "\n\n"
                    "右邊是「飯店╱商場工項類別比較」表格。"
                    "這張表把現場報修、上級交辦、緊急事件、例行維護、每日巡檢，"
                    "用飯店和商場並排的方式呈現，最右邊有占比，"
                    "可以很快看出哪一類工項佔了最多比例。"
                    "如果占比超過 50% 就會顯示紅色，代表工項集中度很高。"
                ),
                "actions": ["指向商場卡片，再指向右側比較表"],
                "hints": [],
            },
            {
                "time": "02:40–03:30",
                "label": "近 12 個月趨勢圖 & 圓餅圖",
                "narration": (
                    "繼續往下滾，看到趨勢圖和圓餅圖。"
                    "\n\n"
                    "左邊的折線圖顯示近十二個月的報修件數趨勢——"
                    "深藍線是報修件數，綠線是完成件數。"
                    "當兩條線越來越近，代表結案效率在提升；"
                    "如果深藍線一直遠高於綠線，就要特別留意積案問題。"
                    "\n\n"
                    "右邊的圓餅圖顯示當月報修類型的分布，"
                    "例如衛浴四成、空調一成四，"
                    "可以幫助判斷資源要重點投入在哪個設施類型。"
                    "\n\n"
                    "下面還有一組是商場的，結構相同。"
                    "\n\n"
                    "好，工務報修摘要和趨勢圖的解說到這邊。"
                    "下一集我們會看每日累計表和六個明細分析面板。"
                ),
                "actions": [
                    "滑鼠停留在折線圖的數據點上，讓 Tooltip 顯示",
                    "滑鼠停留在圓餅圖各區塊上，讓標籤顯示",
                ],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "ep": "第三集",
        "title": "每日累計表 & 六個明細分析面板",
        "duration": "建議時長：5–6 分鐘",
        "goal": "教觀眾展開每日累計表讀數字，以及六個分析面板各自的用途與閱讀方式。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "這集是整個 Dashboard 最深度的一集，"
                    "我們要看每日累計表，還有六個可以折疊的分析面板。"
                    "建議管理層或需要深入追蹤的同仁特別看這集。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:30",
                "label": "飯店 / 商場每日累計案件數",
                "narration": (
                    "繼續往下滾，會看到兩個折疊面板：「飯店每日累計案件數」和「商場每日累計案件數」。"
                    "預設是收合的，點一下標題就可以展開。"
                    "\n\n"
                    "右上角有一個「⊕ 全展開」按鈕，一次把飯店和商場兩張表都打開，"
                    "方便對比。點「⊖ 全收合」就會全部折疊回去。"
                    "\n\n"
                    "表格的欄是當月的每一天，列是工項類別。"
                    "每個格子顯示當天的案件件數。"
                    "\n\n"
                    "要特別說明的是「上級交辦」和「緊急事件」這兩列，"
                    "因為這兩種任務目前系統沒有每日明細的紀錄，"
                    "所以只能顯示月合計，格子裡會看到橘色的數字加上「月計」的小標籤。"
                    "\n\n"
                    "另外，工項類別的名稱是可以點的，"
                    "例如點「現場報修」會直接跳轉到飯店工務的詳細清單頁面，"
                    "非常方便，不用再去側邊欄找。"
                ),
                "actions": [
                    "點「⊕ 全展開」展開兩張表",
                    "指向各工項類別名稱",
                    "點一下「現場報修」示範跳轉，再按上一頁回來",
                ],
                "hints": ["上級交辦/緊急事件的月計標籤是橙色，提醒觀眾注意"],
            },
            {
                "time": "01:30–02:20",
                "label": "📅 每日累計工時表",
                "narration": (
                    "繼續往下，是六個明細分析面板。我們依序看。"
                    "\n\n"
                    "第一個是「每日累計工時表」，顯示當月每天各工項類別的工時，單位是 HR。"
                    "工時的數字有顏色標記：大於等於八小時會顯示紅色，"
                    "四到八小時之間是橙色，方便快速找出工作量異常偏高的那幾天。"
                    "\n\n"
                    "注意，這張表只有在選了「特定月份」的時候才會顯示，"
                    "如果你選的是全年模式，它會提示你選擇月份再看。"
                ),
                "actions": ["展開此面板，指向紅色/橙色的數字"],
                "hints": [],
            },
            {
                "time": "02:20–02:50",
                "label": "📆 每月累計工時表",
                "narration": (
                    "第二個是「每月累計工時表」，顯示選定年度全年十二個月的工時。"
                    "未來的月份會顯示橫線，代表還沒有資料。"
                    "可以點欄位標題排序，找出工時最高的月份。"
                ),
                "actions": ["展開此面板，點 TOTAL 欄排序示範"],
                "hints": [],
            },
            {
                "time": "02:50–03:50",
                "label": "🧮 人員負荷與效率分析",
                "narration": (
                    "第三個是「人員負荷與效率分析」，這個面板對管理者很有用。"
                    "\n\n"
                    "表格列出每位工務人員的總工時、件數、還有「均工時╱件」這個指標，"
                    "也就是平均每一件工作花了多少小時。"
                    "\n\n"
                    "最右邊有一個「判斷」欄，系統會自動幫你標示："
                    "均工時大於等於三小時是紅色「需關注」，"
                    "二點五到三小時是橙色「工時偏高」，"
                    "低於二點五小時是綠色「正常」。"
                    "\n\n"
                    "如果看到紅色或橙色的人員，可能代表這個人工作量過重，"
                    "或是承接了一些比較複雜耗時的案件，需要進一步了解原因。"
                    "\n\n"
                    "欄位標題上都有 ℹ 圖示，停留後會說明每個欄位的計算方式，"
                    "如果對某個數字有疑問，先看那個說明。"
                ),
                "actions": [
                    "展開此面板",
                    "停留在「均工時/件」ℹ 圖示上",
                    "指向紅色的「需關注」Tag",
                ],
                "hints": [],
            },
            {
                "time": "03:50–04:20",
                "label": "🏢 飯店 vs 商場比較表",
                "narration": (
                    "第四個是「飯店 vs 商場比較表」。"
                    "這張表把飯店工務部和商場工務報修並排，"
                    "可以同時比較案件數、工時、完成件數、未完成件數、完成率、主要工項。"
                    "最底列是集團合計。"
                    "如果想快速知道這個月哪邊壓力比較大，看這張表就很直觀。"
                ),
                "actions": ["展開此面板，指向底部合計列"],
                "hints": [],
            },
            {
                "time": "04:20–04:50",
                "label": "📊 工項類別 × 單位矩陣",
                "narration": (
                    "第五個是「工項類別 × 單位矩陣」。"
                    "這張表更細一層，把每個工項類別，"
                    "拆開顯示飯店和商場各自的件數和工時，還有件數占比。"
                    "如果某個工項的占比很高，可以考慮是否需要調整人力配置。"
                ),
                "actions": ["展開此面板，指向件占比欄位"],
                "hints": [],
            },
            {
                "time": "04:50–05:30",
                "label": "⚠️ 異常提醒（預設展開）",
                "narration": (
                    "最後一個——也是唯一預設就展開的——是「異常提醒」面板。"
                    "\n\n"
                    "系統會自動分析當月數據，如果有以下狀況就會跳出警示："
                    "未完成件數大於零、完成率低於六成、"
                    "某一工項類別佔比超過六成、人員工時超過八十小時、"
                    "某天工時暴增超過月均值的兩倍。"
                    "\n\n"
                    "每條警示都有紅橙藍三種等級，紅色最需要立刻處理。"
                    "\n\n"
                    "如果當月一切正常，就會顯示一個綠色的「✅ 本月無異常警示，工務運作正常」。"
                    "\n\n"
                    "這個面板建議每次進 Dashboard 的時候第一個先看，"
                    "有警示就往上面對應的指標查原因。"
                    "\n\n"
                    "好，這集六個分析面板介紹完了。下一集我們看工作日誌。"
                ),
                "actions": [
                    "展開此面板，逐一指向各警示",
                    "如果當月無警示，換到有警示的月份示範",
                ],
                "hints": ["建議先切換到有多個警示的月份再錄製這個段落"],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "ep": "第四集",
        "title": "工作日誌四種查詢模式",
        "duration": "建議時長：4–5 分鐘",
        "goal": "教觀眾用單日、區間、整月、人員四種模式查詢工作記錄，並示範 Excel 匯出。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "這集介紹第二個頁籤——工作日誌。"
                    "工作日誌整合了三個來源的逐筆工作記錄："
                    "飯店工務部、商場工務報修、以及主管交辦和緊急事件。"
                    "你可以查任意日期、任意人員的工作明細。"
                ),
                "actions": ["點選「工作日誌」頁籤"],
                "hints": [],
            },
            {
                "time": "00:20–00:40",
                "label": "摘要卡片說明",
                "narration": (
                    "頁籤最上方有兩張摘要卡片。"
                    "第一張是主管交辦和緊急事件的當月統計，"
                    "顯示件數、工時，還有飯店和商場的拆分數字。"
                    "第二張是現場報修、例行維護、每日巡檢的統計。"
                    "這兩張卡片的數字是從下方工作日誌表格直接加總過來的，"
                    "所以查不同日期範圍，上方的數字也會跟著變動。"
                ),
                "actions": ["指向兩張摘要卡片"],
                "hints": [],
            },
            {
                "time": "00:40–01:30",
                "label": "單日模式",
                "narration": (
                    "查詢模式有四種，用上方的切換按鈕選擇。"
                    "先看「單日」模式——選擇一個日期，點「查詢」，"
                    "就會顯示當天所有人員的工作記錄。"
                    "\n\n"
                    "結果是以人員為單位展開的，"
                    "每個人的名字旁邊有一個彩色的小 Tag，那是班別代碼——"
                    "比如 A 班、B 班，游標停留上去會顯示班別名稱。"
                    "\n\n"
                    "如果看到紅色的 MinusCircle 圖示，代表這個人當天是排休，"
                    "但還是有工作紀錄，系統會自動提醒你。"
                    "如果看到黃色的問號圖示，代表班表裡找不到這個人的排班資料，"
                    "需要確認班表是否有登記完整。"
                    "\n\n"
                    "點一下人員名稱旁邊的三角形可以展開或收合他的工作列表。"
                    "右上角有「全部縮合」和「全部展開」按鈕，方便快速切換。"
                ),
                "actions": [
                    "選一個有資料的日期，點查詢",
                    "展開一位人員的記錄",
                    "停留在班別 Tag 上顯示 Tooltip",
                    "如有異常 Tag 請指向說明",
                ],
                "hints": ["選一個人員較多的平日，資料會比較豐富"],
            },
            {
                "time": "01:30–02:00",
                "label": "工作日誌表格 & Drawer 明細",
                "narration": (
                    "每一列工作記錄包含：工項類別、工作事項、預估耗時、開始結束時間、"
                    "工時、備註和回報事項。"
                    "\n\n"
                    "工作事項前面有一個「飯」或「商」的小標籤，"
                    "代表這筆記錄是來自飯店還是商場的系統。"
                    "\n\n"
                    "點一下任何一列，右側會滑出一個明細面板，"
                    "顯示原始的 Ragic 來源資料，包含所有欄位。"
                    "如果有維修圖片，底部也會一起顯示。"
                    "點標題的「在 Ragic 查看」連結，可以直接跳轉到 Ragic 系統的原始記錄。"
                ),
                "actions": [
                    "點一列工作記錄，展開右側 Drawer",
                    "指向「在 Ragic 查看」連結",
                    "如有圖片，示範點擊預覽",
                ],
                "hints": [],
            },
            {
                "time": "02:00–02:40",
                "label": "區間 / 整月模式",
                "narration": (
                    "切換到「區間」模式——選擇起迄日期，可以跨多天甚至跨月。"
                    "查詢後結果按日期分層，點一天展開就能看那天的人員記錄。"
                    "\n\n"
                    "「整月」模式和區間類似，但直接選年月，系統自動抓整月資料，"
                    "省去手動選起迄日期的步驟，做月報的時候很方便。"
                ),
                "actions": [
                    "切換到「區間」模式，選一個小週期",
                    "切換到「整月」模式，選當月",
                ],
                "hints": [],
            },
            {
                "time": "02:40–03:20",
                "label": "人員模式",
                "narration": (
                    "第四個是「人員」模式，這個最適合追蹤特定人員的工作狀況。"
                    "\n\n"
                    "先選日期區間，再從下拉選單選擇人員，點查詢。"
                    "結果會把這個人在這段期間的所有工作記錄，"
                    "依日期分層集中顯示，不需要一天一天去翻。"
                    "\n\n"
                    "如果主管想了解某位師傅這個月做了哪些工作，"
                    "或者要準備績效評估的資料，用這個模式最快。"
                ),
                "actions": [
                    "切換到「人員」模式",
                    "選日期區間 + 人員",
                    "點查詢展示結果",
                ],
                "hints": [],
            },
            {
                "time": "03:20–03:50",
                "label": "Excel 匯出",
                "narration": (
                    "查詢完成後，右邊會出現一個綠色外框的「匯出 Excel」按鈕。"
                    "點下去就會下載一個 .xlsx 檔案，"
                    "檔名會自動帶上查詢的日期範圍，人員模式還會加上人員名稱。"
                    "\n\n"
                    "這個功能很適合拿來做月度工作日誌報表，"
                    "不用手動整理，直接匯出就能用。"
                    "\n\n"
                    "好，工作日誌的四種模式介紹完了。最後一集我們看異常情境的處理流程。"
                ),
                "actions": [
                    "確保有查詢結果後點「匯出 Excel」",
                    "示範下載的檔名格式",
                ],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════════════
    {
        "ep": "第五集",
        "title": "異常提醒處理 & 常見操作情境",
        "duration": "建議時長：3–4 分鐘",
        "goal": "教觀眾如何從異常提醒出發追蹤問題，並示範六個常用情境的操作流程。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場",
                "narration": (
                    "這是教學系列的最後一集。"
                    "我們來看當系統出現異常提醒的時候，實際的處理流程是什麼，"
                    "以及幾個管理者最常用的操作情境。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:30",
                "label": "從異常提醒出發的追蹤流程",
                "narration": (
                    "打開集團工務概覽，先滾到最底部看「⚠️ 異常提醒」面板。"
                    "\n\n"
                    "假設今天看到這樣一條警示：「未完成件數警示：239 件尚未結案」。"
                    "這條是紅色的，需要立刻處理。"
                    "\n\n"
                    "步驟一：往上看第一列 KPI 的「完成率」，確認目前是百分之多少。"
                    "如果完成率低於六成，很可能還會有一條「完成率偏低」的警示。"
                    "\n\n"
                    "步驟二：往下看摘要卡片，確認未結案主要是集中在飯店還是商場。"
                    "智慧摘要那句話會告訴你主要的類型，例如「主要集中於建築類型」。"
                    "\n\n"
                    "步驟三：點摘要卡片的「查看詳情 ›」，跳轉到對應的詳細清單頁面，"
                    "直接篩選「未結案」狀態，看看是哪些案件、誰負責、已拖多久。"
                    "\n\n"
                    "這個流程適用於大部分的異常處理場景，"
                    "從警示 → 確認範圍 → 鎖定清單 → 分派追蹤。"
                ),
                "actions": [
                    "展開異常提醒面板，指向一條紅色警示",
                    "往上看 KPI 完成率",
                    "回到摘要卡，點「查看詳情 ›」跳轉",
                ],
                "hints": ["建議切到有多條警示的月份，讓流程更完整"],
            },
            {
                "time": "01:30–02:10",
                "label": "情境示範 1：月底工作日誌 Excel 報表",
                "narration": (
                    "情境一：月底要交工作日誌報表。"
                    "點工作日誌頁籤，選「整月」模式，選上個月，點查詢，"
                    "等資料載入完成，點「匯出 Excel」，完成。"
                    "整個流程不到一分鐘。"
                ),
                "actions": ["示範整個流程"],
                "hints": [],
            },
            {
                "time": "02:10–02:40",
                "label": "情境示範 2：找出工時最重的人員",
                "narration": (
                    "情境二：我想知道這個月誰的工作負擔最重。"
                    "回到集團工務概覽，展開「人員負荷與效率分析」，"
                    "點「HR」這欄的標題排序，工時最高的人員就排到最上面了，"
                    "右邊的判斷欄如果是紅色，表示均工時╱件偏高，需要了解原因。"
                ),
                "actions": [
                    "展開人員負荷面板",
                    "點 HR 欄排序",
                    "指向紅色「需關注」Tag",
                ],
                "hints": [],
            },
            {
                "time": "02:40–03:10",
                "label": "情境示範 3：查特定人員某週記錄",
                "narration": (
                    "情境三：主管想了解某位師傅上週的工作狀況。"
                    "工作日誌 → 人員模式 → 選上週起迄日期 → 選人員 → 查詢。"
                    "每天的工作會依日期展開，可以直接截圖給相關人員確認。"
                ),
                "actions": ["示範人員模式查詢"],
                "hints": [],
            },
            {
                "time": "03:10–03:40",
                "label": "結尾",
                "narration": (
                    "好，集團決策 Dashboard 的教學系列到這邊全部結束了。"
                    "\n\n"
                    "快速回顧一下我們學了什麼："
                    "第一集看 KPI 卡片和顏色警示，"
                    "第二集看報修摘要和趨勢圖，"
                    "第三集深入六個分析面板，"
                    "第四集學工作日誌四種查詢模式，"
                    "第五集學從異常提醒出發的處理流程。"
                    "\n\n"
                    "如果對任何功能有疑問，可以點「統計基準說明」頁籤查看詳細的計算口徑說明，"
                    "或是聯繫系統管理員。謝謝收看！"
                ),
                "actions": ["回到集團工務概覽，畫面停在 KPI 卡片做結尾"],
                "hints": [],
            },
        ],
    },
]


# ── Word 文件生成 ─────────────────────────────────────────────────────────────

def build_doc(scripts: list, output_path: Path):
    doc = Document()

    # A4 頁面設定
    sec = doc.sections[0]
    sec.page_height = Inches(11.69)
    sec.page_width  = Inches(8.27)
    sec.top_margin  = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin   = Inches(1.1)
    sec.right_margin  = Inches(1.1)

    # ── 封面 ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_heading("集團決策 Dashboard", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY
    t.runs[0].font.size = Pt(26)

    t2 = doc.add_heading("影音教學旁白腳本", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT
    t2.runs[0].font.size = Pt(20)

    doc.add_paragraph()
    sub = doc.add_paragraph(f"共五集 · {datetime.now().strftime('%Y 年 %m 月 %d 日')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    sub.runs[0].font.size = Pt(12)

    doc.add_paragraph()

    # ── 使用說明 ──────────────────────────────────────────────────────────────
    doc.add_heading("腳本使用說明", 1).runs[0].font.color.rgb = PRIMARY
    tips = [
        "🎬 旁白：錄製時直接念，可依個人習慣調整說話節奏，不必逐字照念。",
        "🖱 操作：每個段落附有「操作」提示框，說明錄製時同步要做的滑鼠動作。",
        "⏸ 提示：黃色「提示」框是錄製前的準備建議，不需要說出來。",
        "[00:00–00:30]：時間軸僅供參考，依實際錄製節奏調整。",
        "旁白中的 \\n\\n 代表說話時的自然停頓（約半秒），不是真的換行。",
    ]
    for t in tips:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(t).font.size = Pt(10.5)

    doc.add_page_break()

    # ── 各集腳本 ──────────────────────────────────────────────────────────────
    for script in scripts:
        # 集標題
        h = doc.add_heading(f"{script['ep']}  {script['title']}", 1)
        h.runs[0].font.color.rgb = PRIMARY
        h.runs[0].font.size = Pt(16)

        # 集資訊列
        info = doc.add_paragraph()
        info.paragraph_format.space_after = Pt(4)
        r1 = info.add_run(f"⏱ {script['duration']}　　")
        r1.font.size = Pt(10)
        r1.font.color.rgb = WARNING
        r1.bold = True

        # 學習目標
        goal_p = doc.add_paragraph()
        goal_p.paragraph_format.space_after = Pt(8)
        goal_p.add_run("學習目標：").bold = True
        goal_p.runs[0].font.color.rgb = ACCENT
        goal_p.add_run(script["goal"])

        section_divider(doc)

        for seg in script["segments"]:
            # 時間軸
            timing_row(doc, seg["time"], seg["label"])

            # 旁白內容（按 \n\n 拆段）
            parts = seg["narration"].split("\n\n")
            for part in parts:
                if part.strip():
                    narration_block(doc, part.strip())

            # 操作提示
            if seg["actions"]:
                for act in seg["actions"]:
                    action_block(doc, act)

            # 錄製前提示
            if seg["hints"]:
                for hint in seg["hints"]:
                    pause_block(doc, hint)

            doc.add_paragraph()

        doc.add_page_break()

    # ── 附錄：術語速查 ────────────────────────────────────────────────────────
    add_heading1(doc, "附錄：術語速查表")
    doc.add_paragraph("錄製時若口誤術語，可參照此表確認正確說法。").runs[0].font.size = Pt(10.5)
    doc.add_paragraph()

    terms = [
        ("完成率",       "Complete Rate / 結案率",        "已結案件數 ÷ 報修總件數 × 100%"),
        ("均工時/件",    "Avg. Hours per Case",           "總工時 ÷ 有工時記錄的件數"),
        ("上期未結",     "Carry-over Cases",              "上月遺留至本月未結案件"),
        ("待辦驗數",     "Pending Verification",          "已修復等待客戶驗收的件數"),
        ("上級交辦",     "Supervisor Task",               "主管指派的工務任務"),
        ("工項類別",     "Work Category",                 "現場報修/上級交辦/緊急事件/例行維護/每日巡檢"),
        ("單位矩陣",     "Category × Unit Matrix",       "工項類別 × 飯店╱商場的交叉統計表"),
        ("Collapse",    "折疊面板",                       "可展開/收合的內容區塊"),
        ("Drawer",      "側拉明細面板",                   "點選列後從右側滑入的詳細資料面板"),
        ("Ragic",       "拉及",                          "後端資料來源系統（發音同英文 ragic）"),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h_text in enumerate(["術語", "英文 / 對應說法", "說明"]):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "1B3A5C")

    for term, en, desc in terms:
        row = table.add_row()
        row.cells[0].text = term
        row.cells[1].text = en
        row.cells[2].text = desc
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 腳本手冊已儲存：{output_path}")


if __name__ == "__main__":
    build_doc(SCRIPTS, OUTPUT_PATH)
    print(f"執行完成！請開啟：{OUTPUT_PATH}")
