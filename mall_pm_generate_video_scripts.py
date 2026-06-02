#!/usr/bin/env python3
"""
商場週期保養表（/mall/periodic-maintenance）影音教學旁白腳本生成器
=================================================================
使用方式：
  python mall_pm_generate_video_scripts.py

輸出：商場週期保養表_影音教學腳本.docx

與飯店版差異：第二集加入「每日巡檢表」Tab 的詳細說明
"""

import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx"); sys.exit(1)

OUTPUT_PATH = Path(__file__).parent / "商場週期保養表_影音教學腳本.docx"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
ORANGE  = RGBColor(0xFA, 0xAD, 0x14)
GRAY    = RGBColor(0x59, 0x59, 0x59)


# ── 輔助函式 ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def narration_block(doc, text: str):
    for part in text.split("\n\n"):
        part = part.strip()
        if not part:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(part); r.font.size = Pt(11.5)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "12"); left.set(qn("w:color"), "4BA8E8")
        pBdr.append(left); pPr.append(pBdr)


def label_box(doc, label: str, text: str, fill_hex: str):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    lc = table.rows[0].cells[0]
    set_cell_bg(lc, fill_hex)
    lp = lc.paragraphs[0]
    lr = lp.add_run(label); lr.bold = True; lr.font.size = Pt(10)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = table.rows[0].cells[1]
    set_cell_bg(rc, "F8F9FA")
    rc.paragraphs[0].add_run(text).font.size = Pt(10.5)
    doc.add_paragraph()


def action_block(doc, text: str):  label_box(doc, "🖱 操作", text, "52C41A")
def hint_block(doc, text: str):    label_box(doc, "⏸ 提示", text, "FA8C16")
def diff_block(doc, text: str):    label_box(doc, "⭐ 商場特有", text, "1565C0")


def timing_row(doc, t_range: str, label: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"[{t_range}]  ")
    r1.bold = True; r1.font.color.rgb = ORANGE; r1.font.size = Pt(10)
    r2 = p.add_run(label)
    r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = PRIMARY


def section_divider(doc):
    p = doc.add_paragraph("─" * 64)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── 腳本資料 ──────────────────────────────────────────────────────────────────
SCRIPTS = [

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第一集",
        "title": "功能簡介 & Dashboard KPI 解讀",
        "duration": "建議時長：4 分鐘",
        "goal": "讓觀眾了解商場週期保養表與飯店版的差異（多一個每日巡檢表 Tab），並掌握 Dashboard 的五個 KPI 卡片。",
        "segments": [
            {
                "time": "00:00–00:40",
                "label": "功能簡介（含與飯店版差異說明）",
                "narration": (
                    "大家好，歡迎收看商場週期保養表教學系列。"
                    "這個模組管理商場所有週期性保養工作的排程、執行與統計追蹤，"
                    "架構和飯店週期保養表高度相同，"
                    "但商場版有一個飯店沒有的功能——「每日巡檢表」Tab。"
                    "\n\n"
                    "如果你已經看過飯店週期保養表的教學，大部分操作方式都一樣，"
                    "重點放在第二集的每日巡檢表說明就好。"
                    "如果是第一次接觸，這系列會完整介紹所有功能。"
                ),
                "actions": ["進入頁面，指向八個頁籤，特別指出「每日巡檢表」是商場獨有"],
                "hints": [],
            },
            {
                "time": "00:40–01:10",
                "label": "八個頁籤快速概覽",
                "narration": (
                    "商場版共有八個頁籤。"
                    "\n\n"
                    "第一個 Dashboard 是主管儀表板，這集的重點。"
                    "第二個「每日巡檢表」是商場獨有的，下一集詳細介紹。"
                    "第三到五是每月、每季、每年的維護統計，"
                    "第六是排程管理，第七是年度計劃表，第八是批次清單。"
                    "\n\n"
                    "第二個以外的所有功能，和飯店版完全相同，"
                    "操作方式可以直接參考飯店版的教學。"
                ),
                "actions": ["依序指向八個頁籤，特別強調第二個"],
                "hints": [],
            },
            {
                "time": "01:10–01:40",
                "label": "年月篩選與本月批次快速入口",
                "narration": (
                    "Dashboard 左上角有年份和月份的篩選選單，"
                    "切換後點「重新整理」更新資料。"
                    "\n\n"
                    "如果本月已有保養批次，頁面底部會顯示批次快速入口，"
                    "點「進入本月批次」直接跳轉到詳細作業頁面。"
                ),
                "actions": ["指向年月篩選，示範切換月份"],
                "hints": ["建議選有資料的近期月份"],
            },
            {
                "time": "01:40–03:20",
                "label": "KPI 五卡解讀",
                "narration": (
                    "Dashboard 頂部是五張 KPI 卡片，都可以點擊展開案件清單。"
                    "\n\n"
                    "「有效項目」深藍色——當月應執行的保養工項總數。"
                    "\n\n"
                    "「已完成」綠色——完成件數，括弧裡是完成率百分比。"
                    "\n\n"
                    "「逾期件數」深紅色——過了排定日期還沒完成的件數，這個最需要關注。"
                    "\n\n"
                    "「異常待追蹤」紫色——有異常記錄需要追蹤的件數。"
                    "\n\n"
                    "「預估工時」藍色——根據保養計劃加總的預計工時（小時）。"
                    "\n\n"
                    "下方有一條完成率進度條，格式是「百分比（完成/總件數）」，藍到綠的漸層。"
                ),
                "actions": [
                    "放大顯示五張卡片，逐一指向",
                    "指向完成率進度條",
                ],
                "hints": [],
            },
            {
                "time": "03:20–04:00",
                "label": "圖表與預警區",
                "narration": (
                    "進度條下方是兩個圖表：左邊是各類別完成率的水平堆疊長條圖，"
                    "右邊是狀態分布的環狀圖。"
                    "\n\n"
                    "再往下是預警區：左邊「逾期項目 Top 10」和右邊「待執行項目」，"
                    "點清單右邊的箭頭按鈕可以跳轉到批次詳細頁。"
                    "\n\n"
                    "好，這集 Dashboard 介紹完了。"
                    "下一集介紹商場獨有的「每日巡檢表」Tab。"
                ),
                "actions": [
                    "指向兩個圖表",
                    "指向預警區",
                ],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第二集",
        "title": "每日巡檢表（商場獨有功能）",
        "duration": "建議時長：4 分鐘",
        "goal": "讓觀眾完整掌握商場獨有的「每日巡檢表」Tab：月份/單日查詢模式、欄位解讀、異常識別。",
        "segments": [
            {
                "time": "00:00–00:25",
                "label": "開場：商場獨有功能",
                "narration": (
                    "這集介紹商場週期保養表最特別的功能——「每日巡檢表」Tab。"
                    "這個 Tab 在飯店版是沒有的，是商場版獨有的功能。"
                    "\n\n"
                    "它整合了商場工務每日巡檢的所有記錄，"
                    "讓管理者可以查看每個樓層、每個設備、每天的巡檢結果。"
                ),
                "actions": ["切換到「每日巡檢表」Tab"],
                "hints": [],
            },
            {
                "time": "00:25–01:00",
                "label": "查詢模式說明",
                "narration": (
                    "篩選列有三個控制項：年份、月份，以及一個可選的日期選擇器。"
                    "\n\n"
                    "預設是「月份模式」——只選年份和月份，"
                    "系統載入整個月的巡檢資料，多個批次合併顯示。"
                    "適合查看本月整體的巡檢狀況。"
                    "\n\n"
                    "如果還額外選了日期，就切換到「單日模式」，"
                    "只顯示那一天的巡檢記錄。"
                    "如果那天沒有資料，頁面會顯示橙色提示橫幅。"
                    "\n\n"
                    "注意：切換年份或月份時，日期選擇會自動清除，"
                    "避免日期跨月的問題。"
                ),
                "actions": [
                    "指向年月篩選說明月份模式",
                    "示範額外選一個日期切換到單日模式",
                ],
                "hints": ["建議先展示月份模式（有完整資料），再切換到有資料的單日示範"],
            },
            {
                "time": "01:00–02:30",
                "label": "表格欄位逐一解讀",
                "narration": (
                    "巡檢記錄表格有七個欄位，我們從左到右說明。"
                    "\n\n"
                    "「樓層」——巡檢的樓層，同一樓層的多筆記錄合併成一格，"
                    "包含 4F、3F、1F 到 3F、1F、B1F 到 B4F 等區域。"
                    "\n\n"
                    "「項目」——巡檢的設備或區域名稱，同一項目的多筆記錄也合併。"
                    "\n\n"
                    "「檢查內容」——具體的檢查項目描述，保留完整換行。"
                    "\n\n"
                    "「實際巡檢人員」——執行這次巡檢的人員姓名。"
                    "\n\n"
                    "「運轉狀況（結果）」——這是最關鍵的欄位，以彩色 Tag 顯示四種狀態："
                    "綠色「正常」、紅色「異常」、黃色「待處理」、灰色「未巡檢」。"
                    "Tag 下方有補充的文字說明。"
                    "\n\n"
                    "「異常說明」——如果是異常狀態，這裡顯示深紅色的詳細說明文字。"
                    "\n\n"
                    "「時間（分）」——這次巡檢耗費的分鐘數。"
                    "有實際記錄時顯示實際值，沒有時顯示標準模板時間。"
                ),
                "actions": [
                    "指向各欄位依序說明",
                    "特別指向「運轉狀況」欄的四種 Tag",
                    "指向異常說明欄的深紅色文字",
                ],
                "hints": [],
            },
            {
                "time": "02:30–03:15",
                "label": "列底色警示與總時間",
                "narration": (
                    "表格的列底色也有意義——"
                    "淺紅底色代表這個巡檢項目有異常記錄，需要特別關注。"
                    "淺灰底色代表這個項目未完成巡檢。"
                    "白色底色是正常完成的項目。"
                    "\n\n"
                    "管理者掃視表格，淺紅色的列就是當天需要追蹤的問題點。"
                    "\n\n"
                    "表格最底部有「總巡檢時間：N 分鐘」的統計。"
                    "如果有實際巡檢記錄就加總實際時間；"
                    "如果沒有實際記錄就用標準模板時間加總，讓管理者知道大概需要多少時間。"
                ),
                "actions": [
                    "指向淺紅底色的列",
                    "指向表格底部總時間",
                ],
                "hints": [],
            },
            {
                "time": "03:15–03:45",
                "label": "常用情境示範",
                "narration": (
                    "快速示範兩個最常用的情境。"
                    "\n\n"
                    "情境一：確認今天商場工務巡檢有沒有異常。"
                    "日期選今天，按查詢，掃視表格找淺紅底色的列，"
                    "點開「運轉狀況」欄看異常說明。"
                    "\n\n"
                    "情境二：查看本月哪些項目有異常記錄。"
                    "不選日期（月份模式），整月資料整合顯示，"
                    "淺紅色的列就是本月曾出現異常的巡檢項目。"
                    "\n\n"
                    "好，商場獨有的每日巡檢表介紹完了。"
                    "下一集從第三個 Tab 開始，介紹三個維護統計 Tab。"
                ),
                "actions": [
                    "示範選今天日期查詢",
                    "示範月份模式指向淺紅列",
                ],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第三集",
        "title": "每月 / 每季 / 每年維護統計（同飯店版）",
        "duration": "建議時長：4 分鐘",
        "goal": "讓觀眾掌握三個維護統計 Tab 的年度矩陣與鑽取功能（與飯店版相同）。",
        "segments": [
            {
                "time": "00:00–00:25",
                "label": "開場接續",
                "narration": (
                    "這集介紹第三到第五個 Tab——每月維護、每季維護、每年維護。"
                    "這三個 Tab 的功能和飯店週期保養表完全相同，"
                    "如果已看過飯店版教學，可以跳過這集。"
                    "第一次看的觀眾請繼續。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:25–01:20",
                "label": "年度矩陣（三個 Tab 共用結構）",
                "narration": (
                    "三個 Tab 都有一張「年度矩陣總表」，"
                    "以 12 個月為橫軸，統計指標為縱軸。"
                    "\n\n"
                    "點擊表格中任一格，下方的 KPI 卡片就會更新到那個月的詳細數據，"
                    "稱為「鑽取」功能。"
                    "\n\n"
                    "左上角有年份選單可以切換查看不同年度，"
                    "方便跨年比較保養完成率的趨勢。"
                ),
                "actions": [
                    "切換到「每月維護」Tab，展示年度矩陣",
                    "點一個月份格子，展示下方 KPI 更新",
                ],
                "hints": [],
            },
            {
                "time": "01:20–02:20",
                "label": "每月維護：KPI 鑽取與未完成說明",
                "narration": (
                    "在「每月維護」Tab，選定月份後，矩陣下方顯示兩區塊。"
                    "\n\n"
                    "左邊「上月累計」：上月遺留未完成數、本月已補做的上月件數、累計完成率。"
                    "\n\n"
                    "右邊「本月統計」：本月項目數、完成數、完成率，搭配進度條。"
                    "\n\n"
                    "底部是「未完成事項說明」表格，列出有填寫備注的未完成項目。"
                ),
                "actions": [
                    "點矩陣一個月份，展示 KPI 鑽取",
                    "指向上月累計和本月統計兩個區塊",
                ],
                "hints": [],
            },
            {
                "time": "02:20–02:55",
                "label": "每季維護",
                "narration": (
                    "「每季維護」Tab，矩陣下方顯示 Q1 到 Q4 的選擇卡片。"
                    "每個卡片顯示該季三個月的加總：工項數、完成數、完成率。"
                    "點選後顯示季度 KPI 卡片和月份分布明細。"
                    "適合做季報。"
                ),
                "actions": [
                    "切換到「每季維護」Tab",
                    "點選一個季度卡片",
                ],
                "hints": [],
            },
            {
                "time": "02:55–03:30",
                "label": "每年維護",
                "narration": (
                    "「每年維護」Tab 顯示全年 KPI 和 Q1 到 Q4 的季度分布，"
                    "適合年底做年度保養成果報告。"
                    "\n\n"
                    "三個維護統計 Tab 介紹完了。"
                    "下一集看排程管理和年度計劃表。"
                ),
                "actions": ["切換到「每年維護」Tab，指向各區塊"],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第四集",
        "title": "排程管理 & 年度計劃表（同飯店版）",
        "duration": "建議時長：4 分鐘",
        "goal": "讓觀眾掌握排程管理的產生排程功能和年度計劃表的矩陣操作（與飯店版相同）。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場接續",
                "narration": (
                    "這集介紹排程管理和年度計劃表，功能和飯店版相同。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:10",
                "label": "排程管理：篩選與產生排程",
                "narration": (
                    "切換到「排程管理」Tab。"
                    "\n\n"
                    "頂部篩選工具列提供年月、保養類別、執行人員、狀態四個條件。"
                    "\n\n"
                    "右側有「產生排程」按鈕——"
                    "每個月初執行一次，系統自動為當月應執行的保養項目建立排程記錄。"
                    "已有排程的項目不會重複產生，可以放心點擊。"
                ),
                "actions": [
                    "切換到「排程管理」Tab",
                    "指向篩選工具列和產生排程按鈕",
                ],
                "hints": [],
            },
            {
                "time": "01:10–01:55",
                "label": "排程表格與狀態",
                "narration": (
                    "排程表格的每列是一筆保養排程，"
                    "包含項目名稱、類別、排定日期、執行人員、狀態、備注。"
                    "\n\n"
                    "狀態有五種顏色：黃色已排定、藍色進行中、綠色已完成、"
                    "深紅色逾期、灰色非本月。"
                    "\n\n"
                    "點編輯圖示可以手動調整排定日期或執行人員。"
                ),
                "actions": [
                    "指向各狀態 Tag",
                    "點一列的編輯圖示展示調整介面",
                ],
                "hints": [],
            },
            {
                "time": "01:55–03:10",
                "label": "年度計劃表",
                "narration": (
                    "切換到「年度計劃表」Tab。"
                    "\n\n"
                    "這是最全面的追蹤視圖——"
                    "橫軸是 12 個月，縱軸是所有保養項目，"
                    "每個格子代表該項目在那個月的執行狀態，以顏色標示。"
                    "\n\n"
                    "點擊格子，右側滑出 Drawer 明細，"
                    "顯示當月的執行人員、完成時間、備注和 Ragic 原始連結。"
                    "\n\n"
                    "如果格子顯示警示訊息「應執行但無排程」，"
                    "要去「排程管理」Tab 執行產生排程。"
                    "\n\n"
                    "年度計劃表截圖很適合放到年報或向主管簡報。"
                    "\n\n"
                    "好，這集排程管理和年度計劃表介紹完了。"
                    "下一集是最後一集，看批次清單和整合情境。"
                ),
                "actions": [
                    "切換到「年度計劃表」Tab",
                    "點一個格子展示 Drawer",
                    "指向有警示的格子（若有）",
                ],
                "hints": [],
            },
        ],
    },

    # ══════════════════════════════════════════════════════════════════
    {
        "ep": "第五集",
        "title": "批次清單 & 整合操作情境",
        "duration": "建議時長：3–4 分鐘",
        "goal": "讓觀眾掌握批次清單的查詢方式，並示範商場版四個完整操作情境（含每日巡檢表的應用）。",
        "segments": [
            {
                "time": "00:00–00:20",
                "label": "開場",
                "narration": (
                    "最後一集，介紹批次清單和整合操作情境，"
                    "特別加入商場版獨有的每日巡檢表應用場景。"
                ),
                "actions": [],
                "hints": [],
            },
            {
                "time": "00:20–01:10",
                "label": "批次清單",
                "narration": (
                    "切換到「批次清單」Tab。"
                    "\n\n"
                    "這裡列出所有歷史保養批次，每個批次對應一個月份的保養作業。"
                    "選年份後顯示該年各月的批次記錄。"
                    "\n\n"
                    "表格包含：保養單號（可點）、保養月份、批次狀態 Tag、"
                    "完成率進度條、逾期件數（紅色徽章）、異常件數（紫色徽章）。"
                    "\n\n"
                    "點保養單號進入批次詳細頁，可以逐項查看保養記錄和填寫結果。"
                    "完成所有項目後，在批次層級做結案動作。"
                ),
                "actions": [
                    "切換到「批次清單」Tab",
                    "點一個保養單號，示範跳轉，再回來",
                ],
                "hints": [],
            },
            {
                "time": "01:10–01:45",
                "label": "情境 1：月初建立本月排程",
                "narration": (
                    "情境一：月初標準流程。"
                    "排程管理 Tab → 選本月 → 點「產生排程」→ 回到 Dashboard 確認 KPI 更新。"
                    "同時確認每日巡檢表有載入資料，若無資料表示巡檢批次尚未建立。"
                ),
                "actions": ["示範月初產生排程流程"],
                "hints": [],
            },
            {
                "time": "01:45–02:20",
                "label": "情境 2：追蹤今日異常巡檢",
                "narration": (
                    "情境二：商場工務主管想知道今天巡檢有沒有問題。"
                    "\n\n"
                    "切換到「每日巡檢表」Tab，日期選今天，點查詢。"
                    "掃視表格，淺紅底色的列就是有異常的巡檢項目，"
                    "查看「異常說明」欄確認問題內容，"
                    "記下樓層和設備名稱後聯絡相關人員處理。"
                ),
                "actions": ["示範切換到每日巡檢表，選今天，找異常列"],
                "hints": [],
            },
            {
                "time": "02:20–02:55",
                "label": "情境 3：月底製作保養完成率報告",
                "narration": (
                    "情境三：月底製作保養報告。"
                    "\n\n"
                    "Dashboard 截圖 KPI 卡片和進度條，"
                    "切換到「每月維護」Tab 截圖年度矩陣，"
                    "切換到「每日巡檢表」Tab（月份模式）截圖整月巡檢摘要。"
                    "\n\n"
                    "這三張截圖配合起來，就是商場保養工作的完整月報素材。"
                ),
                "actions": ["依序截圖三個頁面示範"],
                "hints": [],
            },
            {
                "time": "02:55–03:30",
                "label": "結尾",
                "narration": (
                    "商場週期保養表的五集教學全部結束了。"
                    "\n\n"
                    "快速回顧：第一集學 Dashboard KPI，"
                    "第二集學商場獨有的每日巡檢表，"
                    "第三集學三個維護統計 Tab，"
                    "第四集學排程管理和年度計劃表，"
                    "第五集學批次清單和整合情境。"
                    "\n\n"
                    "整體架構和飯店週期保養表相同，"
                    "最大的差異就是這個「每日巡檢表」Tab，"
                    "讓商場的日常巡檢記錄也能在同一個模組裡查詢管理。謝謝收看！"
                ),
                "actions": ["畫面停在 Dashboard 做結尾"],
                "hints": [],
            },
        ],
    },
]


# ── Word 文件生成 ─────────────────────────────────────────────────────────────
def build_doc(scripts: list, output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height   = Inches(11.69); sec.page_width    = Inches(8.27)
    sec.top_margin    = Inches(0.85);  sec.bottom_margin = Inches(0.85)
    sec.left_margin   = Inches(1.1);   sec.right_margin  = Inches(1.1)

    # 封面
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_heading("商場週期保養表", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY; t.runs[0].font.size = Pt(26)
    t2 = doc.add_heading("影音教學旁白腳本", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT; t2.runs[0].font.size = Pt(20)
    doc.add_paragraph()
    sub = doc.add_paragraph(
        f"共五集 · 路由：/mall/periodic-maintenance\n{datetime.now().strftime('%Y 年 %m 月 %d 日')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    doc.add_paragraph()

    # 使用說明
    doc.add_heading("腳本使用說明", 1).runs[0].font.color.rgb = PRIMARY
    for tip in [
        "🎬 旁白：錄製時直接念，可依個人習慣調整，不必逐字照念。",
        "🖱 操作：錄製時同步要做的滑鼠動作。",
        "⏸ 提示：錄製前的準備建議，不需要說出來。",
        "⭐ 商場特有：藍色標示框代表商場版與飯店版的差異重點。",
        "若觀眾已看過飯店週期保養表教學，第一集開頭說明差異後，可以直接跳到第二集的每日巡檢表。",
        "建議選有巡檢資料且有異常記錄（淺紅底色）的日期錄製第二集，視覺效果最佳。",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(tip).font.size = Pt(10.5)

    doc.add_page_break()

    for script in scripts:
        h = doc.add_heading(f"{script['ep']}  {script['title']}", 1)
        h.runs[0].font.color.rgb = PRIMARY; h.runs[0].font.size = Pt(16)

        info = doc.add_paragraph()
        r1 = info.add_run(f"⏱ {script['duration']}　　")
        r1.font.size = Pt(10); r1.font.color.rgb = ORANGE; r1.bold = True

        gp = doc.add_paragraph()
        gp.paragraph_format.space_after = Pt(8)
        gp.add_run("學習目標：").bold = True
        gp.runs[0].font.color.rgb = ACCENT
        gp.add_run(script["goal"])

        section_divider(doc)

        for seg in script["segments"]:
            timing_row(doc, seg["time"], seg["label"])
            narration_block(doc, seg["narration"])
            for act in seg.get("actions", []):
                action_block(doc, act)
            for hint in seg.get("hints", []):
                hint_block(doc, hint)
            doc.add_paragraph()

        doc.add_page_break()

    # 附錄
    doc.add_heading("附錄：狀態顏色速查", 1).runs[0].font.color.rgb = PRIMARY

    doc.add_heading("保養排程狀態", 2).runs[0].font.color.rgb = ACCENT
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["狀態", "顏色", "說明"]):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1B3A5C")
    for status, color, desc in [
        ("已完成",  "🟢 綠色", "保養完成"),
        ("進行中",  "🔵 藍色", "執行中"),
        ("已排定",  "🟡 黃色", "已排程尚未開始"),
        ("未排定",  "🔴 紅色", "應執行但無排程"),
        ("逾期",    "🔴 深紅", "過期未完成"),
        ("非本月",  "⬜ 灰色", "非本月週期"),
    ]:
        row = table.add_row()
        row.cells[0].text = status; row.cells[1].text = color; row.cells[2].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    doc.add_heading("每日巡檢表（⭐ 商場獨有）", 2).runs[0].font.color.rgb = ACCENT
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Table Grid"
    for i, h in enumerate(["結果狀態", "Tag 顏色", "說明", "列底色"]):
        c = table2.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1B3A5C")
    for status, color, desc, bg in [
        ("正常",   "🟢 綠色 #52C41A", "巡檢正常完成",  "白色"),
        ("異常",   "🔴 紅色 #FF4D4F", "發現異常",      "淺紅底色"),
        ("待處理", "🟡 黃色 #FAAD14", "需要後續處理",  "白色"),
        ("未巡檢", "⬜ 灰色 #d9d9d9", "未完成巡檢",    "淺灰底色"),
    ]:
        row = table2.add_row()
        row.cells[0].text = status; row.cells[1].text = color
        row.cells[2].text = desc;   row.cells[3].text = bg
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 腳本手冊已儲存：{output_path}")


if __name__ == "__main__":
    build_doc(SCRIPTS, OUTPUT_PATH)
    print(f"執行完成！請開啟：{OUTPUT_PATH}")
