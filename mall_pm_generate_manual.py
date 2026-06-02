#!/usr/bin/env python3
"""
商場週期保養表（/mall/periodic-maintenance）教學手冊自動生成腳本
=================================================================
使用方式：
  pip install playwright python-docx
  playwright install chromium
  python mall_pm_generate_manual.py

輸出：
  manual_screenshots_mall_pm/
  商場週期保養表教學手冊.docx

與飯店版差異：多一個「每日巡檢表」Tab（商場工務每日巡檢記錄）
"""

import asyncio
import sys
from pathlib import Path
from datetime import datetime

try:
    from playwright.async_api import async_playwright
except ImportError:
    print("❌ 請先執行：pip install playwright && playwright install chromium"); sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx"); sys.exit(1)

# ── 設定 ──────────────────────────────────────────────────────────────────────
BASE_URL    = "http://localhost:5173"
TARGET_URL  = f"{BASE_URL}/mall/periodic-maintenance"
OUTPUT_DIR  = Path(__file__).parent / "manual_screenshots_mall_pm"
MANUAL_PATH = Path(__file__).parent / "商場週期保養表教學手冊.docx"
USERNAME    = "admin"
PASSWORD    = "Admin@2026"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
GREEN   = RGBColor(0x52, 0xC4, 0x1A)
ORANGE  = RGBColor(0xFA, 0xAD, 0x14)
RED     = RGBColor(0xCF, 0x13, 0x22)
GRAY    = RGBColor(0x59, 0x59, 0x59)

CONTENT_X  = 185
CONTENT_W  = 1390
VIEWPORT_H = 900


# ── 截圖邏輯 ──────────────────────────────────────────────────────────────────
async def capture_screenshots(page, output_dir: Path) -> dict:
    shots = {}

    async def snap(name: str, scroll_y: int = 0, clip=None, wait_ms: int = 800):
        try:
            await page.evaluate(f"window.scrollTo(0, {scroll_y})")
            await page.wait_for_timeout(wait_ms)
            path = output_dir / f"{name}.png"
            await page.screenshot(path=str(path), full_page=False,
                                  **({"clip": clip} if clip else {}))
            shots[name] = path
            print(f"   ✓ {name}")
        except Exception as e:
            print(f"   ✗ {name}：{e}")

    async def click_tab(tab_text: str, wait_ms: int = 2000):
        try:
            tab = page.locator("div.ant-tabs-tab").filter(has_text=tab_text).first
            await tab.click()
            await page.wait_for_timeout(wait_ms)
        except Exception as e:
            print(f"   ⚠ Tab [{tab_text}] 失敗：{e}")

    # ── 登入 ──────────────────────────────────────────────────────────────────
    print("🔐 登入中...")
    await page.goto(f"{BASE_URL}/login")
    await page.wait_for_load_state("domcontentloaded")
    await page.wait_for_timeout(800)
    await page.locator("input[placeholder*='帳號'], input[type='text']").first.fill(USERNAME)
    await page.locator("input[type='password']").first.fill(PASSWORD)
    await page.locator("button[type='submit'], button:has-text('登入')").first.click()
    try:
        await page.wait_for_url(f"{BASE_URL}/dashboard", timeout=15000)
        print("   ✓ 登入成功")
    except Exception:
        print(f"   ✓ 已跳轉至 {page.url}")

    # ── 導航 ──────────────────────────────────────────────────────────────────
    await page.goto(TARGET_URL)
    await page.wait_for_load_state("networkidle")
    try:
        await page.locator(".ant-statistic-content-value").first.wait_for(timeout=12000)
        await page.wait_for_timeout(2000)
        print("   ✓ 頁面載入完成")
    except Exception:
        await page.wait_for_timeout(5000)

    # ════════════════════════════════════════════════════════════
    # Tab 1：Dashboard
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：Dashboard Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await page.wait_for_timeout(600)

    # 1. KPI 五卡 + 年月篩選
    await snap("01_dashboard_kpi",
               clip={"x": CONTENT_X, "y": 55, "width": CONTENT_W, "height": 370})

    # 2. 完成率進度條 + 圖表（Bar + Donut）
    await snap("02_dashboard_charts",
               clip={"x": CONTENT_X, "y": 415, "width": CONTENT_W, "height": 380})

    # 3. 預警區（逾期 Top 10 + 待執行）
    await snap("03_dashboard_alerts",
               scroll_y=790,
               clip={"x": CONTENT_X, "y": 790, "width": CONTENT_W, "height": 360})

    # ════════════════════════════════════════════════════════════
    # Tab 2：每日巡檢表（商場獨有）
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：每日巡檢表 Tab（商場獨有）")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("每日巡檢表", wait_ms=2500)

    # 篩選列
    await snap("04_daily_insp_toolbar",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 110})

    # 巡檢記錄表格（月份模式）
    await snap("05_daily_insp_table",
               clip={"x": CONTENT_X, "y": 260, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab 3：每月維護
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：每月維護 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("每月維護", wait_ms=2500)
    await snap("06_monthly_matrix",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 380})
    await snap("07_monthly_kpi",
               scroll_y=500,
               clip={"x": CONTENT_X, "y": 500, "width": CONTENT_W, "height": 480})

    # ════════════════════════════════════════════════════════════
    # Tab 4：每季維護
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：每季維護 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("每季維護", wait_ms=2500)
    await snap("08_quarterly",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 620})

    # ════════════════════════════════════════════════════════════
    # Tab 5：每年維護
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：每年維護 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("每年維護", wait_ms=2500)
    await snap("09_yearly",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab 6：排程管理
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：排程管理 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("排程管理", wait_ms=2500)
    await snap("10_schedule_toolbar",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 240})
    await snap("11_schedule_table",
               scroll_y=280,
               clip={"x": CONTENT_X, "y": 280, "width": CONTENT_W, "height": 500})

    # ════════════════════════════════════════════════════════════
    # Tab 7：年度計劃表
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：年度計劃表 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("年度計劃表", wait_ms=2500)
    await snap("12_annual_matrix",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 600})

    # ════════════════════════════════════════════════════════════
    # Tab 8：批次清單
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：批次清單 Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("批次清單", wait_ms=2000)
    await snap("13_batch_list",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 550})

    return shots


# ── Word 手冊輔助函式 ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def add_image_safe(doc, path, width_inches: float = 5.8):
    if isinstance(path, Path) and path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph("[ 截圖未能取得，請參考系統實際畫面 ]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = GRAY; p.runs[0].font.italic = True


def add_tip(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run("💡 提示："); r.bold = True; r.font.color.rgb = PRIMARY
    p.add_run(text)


def highlight_box(doc, text: str, bg="E8F5E9"):
    """醒目提示框（綠底）"""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    c = table.rows[0].cells[0]
    set_cell_bg(c, bg)
    c.paragraphs[0].add_run(text).font.size = Pt(10.5)
    doc.add_paragraph()


def hdr_table(doc, headers, bg="1B3A5C"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, bg)
    return table


def h1(doc, t):
    h = doc.add_heading(t, 1); h.runs[0].font.color.rgb = PRIMARY; return h

def h2(doc, t):
    h = doc.add_heading(t, 2); h.runs[0].font.color.rgb = ACCENT; return h


# ── Word 手冊建構 ─────────────────────────────────────────────────────────────
def build_word_manual(shots: dict, output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11.69); sec.page_width  = Inches(8.27)
    sec.top_margin  = Inches(0.9);   sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1.1);   sec.right_margin  = Inches(1.1)

    # ── 封面 ──────────────────────────────────────────────────────────────────
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_heading("商場週期保養表", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY; t.runs[0].font.size = Pt(28)
    t2 = doc.add_heading("教學操作手冊", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT; t2.runs[0].font.size = Pt(22)
    doc.add_paragraph()
    sub = doc.add_paragraph(
        f"路由：/mall/periodic-maintenance\n{datetime.now().strftime('%Y 年 %m 月 %d 日')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    doc.add_page_break()

    # ── 一、功能簡介 ──────────────────────────────────────────────────────────
    h1(doc, "一、功能簡介")
    p = doc.add_paragraph()
    p.add_run("商場週期保養表").bold = True
    p.add_run(
        "（路徑：商場管理 ▶ 週期保養表，路由：/mall/periodic-maintenance）"
        "管理商場所有週期性保養工作的排程、執行與統計追蹤，"
        "涵蓋月維護、季維護、年維護三種頻率。"
        "架構與飯店週期保養表相同，但商場版額外提供「每日巡檢表」頁籤，"
        "用於查詢商場工務每日巡檢記錄。"
    )
    doc.add_paragraph()

    highlight_box(doc,
        "⭐ 與飯店週期保養表的主要差異：\n"
        "商場版多一個「每日巡檢表」Tab（Tab 2），"
        "用於查看商場工務每日巡檢的逐筆記錄，這是飯店版沒有的功能。",
        bg="E3F2FD")

    tab_info = [
        ("Dashboard",   "KPI 五卡、完成率圖表、逾期/待執行預警（同飯店）"),
        ("每日巡檢表",  "⭐ 商場獨有！商場工務每日巡檢記錄（月份/單日模式）"),
        ("每月維護",    "月統計 + 年度矩陣 + 未完成說明（同飯店）"),
        ("每季維護",    "Q1-Q4 季度統計（同飯店）"),
        ("每年維護",    "年度統計 + Q1-Q4 分布（同飯店）"),
        ("排程管理",    "排程管理 + 產生排程功能（同飯店）"),
        ("年度計劃表",  "12個月矩陣視圖（同飯店）"),
        ("批次清單",    "歷史批次列表（同飯店）"),
    ]
    table = hdr_table(doc, ["頁籤", "說明"])
    for name, desc in tab_info:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ── 二、Dashboard（同飯店，簡要說明）─────────────────────────────────────
    h1(doc, "二、Dashboard")

    h2(doc, "2.1  KPI 五卡")
    add_image_safe(doc, shots.get("01_dashboard_kpi"))
    doc.add_paragraph()
    doc.add_paragraph("年月篩選列 + KPI 五卡（同飯店週期保養表，詳細說明請參考飯店版手冊）：")
    kpi5 = [
        ("有效項目",   "當月符合條件的保養工項數",           "深藍"),
        ("已完成",     "已完成件數（括弧內顯示完成率%）",    "綠色"),
        ("逾期件數",   "已過排定日期但未完成的件數",          "深紅"),
        ("異常待追蹤", "有異常記錄需追蹤的件數",              "紫色"),
        ("預估工時",   "計劃工時加總（小時）",                "藍色"),
    ]
    table = hdr_table(doc, ["KPI", "說明", "顏色"])
    for name, desc, color in kpi5:
        row = table.add_row()
        for i, v in enumerate([name, desc, color]):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    h2(doc, "2.2  圖表 & 預警區")
    add_image_safe(doc, shots.get("02_dashboard_charts"))
    doc.add_paragraph("各類別完成率水平 Bar Chart + 狀態分布 Donut 圖（同飯店版）。")
    doc.add_paragraph()
    add_image_safe(doc, shots.get("03_dashboard_alerts"))
    doc.add_paragraph("逾期項目 Top 10（深紅）+ 待執行項目（黃）+ 本月批次快速入口（同飯店版）。")
    doc.add_page_break()

    # ── 三、每日巡檢表（商場獨有）────────────────────────────────────────────
    h1(doc, "三、每日巡檢表（商場獨有）")
    highlight_box(doc,
        "⭐ 此頁籤為商場版獨有，飯店週期保養表沒有此功能。",
        bg="FFF9C4")

    add_image_safe(doc, shots.get("04_daily_insp_toolbar"))
    doc.add_paragraph()
    add_image_safe(doc, shots.get("05_daily_insp_table"))
    doc.add_paragraph()
    doc.add_paragraph(
        "「每日巡檢表」Tab 查詢商場工務每日巡檢的詳細記錄，"
        "整合多張巡檢表（4F / 3F / 1F~3F / 1F / B1F~B4F）。"
    )
    doc.add_paragraph()

    h2(doc, "3.1  查詢模式")
    modes = [
        ("月份模式（預設）", "選年份 + 月份，系統自動載入整月彙整資料，多筆批次合併顯示"),
        ("單日模式",         "額外選擇特定日期（DatePicker），只查看當天的巡檢記錄\n若該日無資料，頁面顯示橙色提示橫幅"),
    ]
    for name, desc in modes:
        p = doc.add_paragraph()
        p.add_run(f"• {name}：").bold = True
        p.add_run(desc)
    doc.add_paragraph()
    add_tip(doc, "切換年份或月份時，日期選擇自動清除（避免日期跨月問題）。")
    doc.add_paragraph()

    h2(doc, "3.2  表格欄位說明")
    col_rows = [
        ("樓層",          "巡檢樓層（跨列合併顯示，同樓層的列合為一格）"),
        ("項目",          "巡檢設備/區域名稱（同項目的列合併）"),
        ("檢查內容",      "具體的檢查項目描述（保留換行）"),
        ("實際巡檢人員",  "執行巡檢的人員姓名"),
        ("運轉狀況（結果）","Tag 顯示：正常（綠）/ 異常（紅）/ 待處理（黃）/ 未巡檢（灰）\n下方顯示補充文字說明"),
        ("異常說明",      "異常時的詳細說明（深紅色字體）"),
        ("時間（分）",    "巡檢耗時（分鐘）；有實際記錄時顯示實際值，否則顯示標準模板值"),
    ]
    table = hdr_table(doc, ["欄位", "說明"])
    for name, desc in col_rows:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph("列底色：")
    doc.add_paragraph("• 淺紅底色：該項目有異常記錄")
    doc.add_paragraph("• 淺灰底色：該項目未完成巡檢")
    doc.add_paragraph("• 白色（正常）：正常完成巡檢")
    doc.add_paragraph()
    doc.add_paragraph("表格底部顯示「總巡檢時間：N 分鐘」，")
    doc.add_paragraph("優先取實際巡檢記錄的時間；若無實際資料則加總標準模板時間。")
    doc.add_page_break()

    # ── 四、每月/每季/每年維護（與飯店版相同）────────────────────────────────
    h1(doc, "四、每月 / 每季 / 每年維護統計（同飯店版）")
    doc.add_paragraph(
        "Tab 3、4、5 的結構與飯店週期保養表完全相同，以下為快速說明："
    )
    doc.add_paragraph()

    h2(doc, "4.1  每月維護")
    add_image_safe(doc, shots.get("06_monthly_matrix"))
    doc.add_paragraph()
    add_image_safe(doc, shots.get("07_monthly_kpi"))
    doc.add_paragraph("年度矩陣 + 單月鑽取（上月累計 / 本月統計 / 未完成說明），詳見飯店版手冊。")
    doc.add_paragraph()

    h2(doc, "4.2  每季維護")
    add_image_safe(doc, shots.get("08_quarterly"))
    doc.add_paragraph("年度矩陣 + Q1-Q4 選擇卡片 + 季度 KPI，詳見飯店版手冊。")
    doc.add_paragraph()

    h2(doc, "4.3  每年維護")
    add_image_safe(doc, shots.get("09_yearly"))
    doc.add_paragraph("年度矩陣 + 全年 KPI + 季度分布，詳見飯店版手冊。")
    doc.add_page_break()

    # ── 五、排程管理（同飯店版）─────────────────────────────────────────────
    h1(doc, "五、排程管理（同飯店版）")
    add_image_safe(doc, shots.get("10_schedule_toolbar"))
    doc.add_paragraph()
    add_image_safe(doc, shots.get("11_schedule_table"))
    doc.add_paragraph()
    doc.add_paragraph(
        "「排程管理」Tab 功能與飯店版完全相同：\n"
        "篩選工具列（年月、類別、人員、狀態）+ 產生排程按鈕 + 排程表格（可手動調整）。"
        "排程狀態顏色說明請參考飯店版手冊的附錄。"
    )
    add_tip(doc, "每月初由管理人員點擊「產生排程」，系統自動為當月保養項目建立排程記錄。")
    doc.add_page_break()

    # ── 六、年度計劃表（同飯店版）──────────────────────────────────────────
    h1(doc, "六、年度計劃表（同飯店版）")
    add_image_safe(doc, shots.get("12_annual_matrix"))
    doc.add_paragraph()
    doc.add_paragraph(
        "12 個月 × 保養項目的全年矩陣。點格子可展開右側 Drawer 查看當月明細。"
        "若格子顯示警示，表示應執行的月份尚未建立排程。"
    )
    doc.add_page_break()

    # ── 七、批次清單（同飯店版）─────────────────────────────────────────────
    h1(doc, "七、批次清單（同飯店版）")
    add_image_safe(doc, shots.get("13_batch_list"))
    doc.add_paragraph()
    doc.add_paragraph(
        "歷史保養批次列表：保養單號（可點擊）、保養月份、批次狀態 Tag、完成率進度條、"
        "逾期件數（紅色徽章）、異常件數（紫色徽章）。"
        "點保養單號進入批次詳細頁逐項查看。"
    )
    doc.add_page_break()

    # ── 八、商場版特有操作情境 ───────────────────────────────────────────────
    h1(doc, "八、商場版特有操作情境")

    h2(doc, "情境 1：查看今日商場工務巡檢結果")
    doc.add_paragraph("步驟：切換到「每日巡檢表」Tab → 日期選今日 → 點「查詢」")
    doc.add_paragraph("說明：若今日已有巡檢批次，顯示當日各樓層的逐筆巡檢記錄；若無資料顯示橙色提示橫幅。")
    doc.add_paragraph()

    h2(doc, "情境 2：確認本月哪幾天有異常巡檢記錄")
    doc.add_paragraph("步驟：切換到「每日巡檢表」Tab → 年月選本月（不選日期）→ 查看淺紅底色的列")
    doc.add_paragraph("說明：月份模式整合當月所有批次，淺紅底色代表有異常記錄的巡檢項目。")
    doc.add_paragraph()

    h2(doc, "情境 3：製作本月保養完成率報告")
    doc.add_paragraph("步驟：Dashboard Tab → 查看 KPI 卡片 → 截圖進度條 → 切換到「每月維護」Tab 截圖矩陣")
    doc.add_paragraph("說明：和飯店版操作相同，加上每日巡檢表的截圖作為補充。")

    # ── 附錄：狀態顏色對照表 ─────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "附錄：狀態顏色對照表")

    h2(doc, "保養排程狀態（Dashboard / 排程管理 共用）")
    table = hdr_table(doc, ["狀態", "顏色", "說明"])
    sched_rows = [
        ("已完成",      "🟢 綠色 #52C41A",  "保養工項已完成"),
        ("進行中",      "🔵 藍色 #4BA8E8",  "執行中"),
        ("已排定",      "🟡 黃色 #FAAD14",  "已排程尚未開始"),
        ("未排定",      "🔴 紅色 #FF4D4F",  "應執行但無排程"),
        ("逾期",        "🔴 深紅 #C0392B",  "過期未完成"),
        ("非本月",      "⬜ 灰色 #999999",  "非本月頻率週期"),
    ]
    for s, c, d in sched_rows:
        row = table.add_row()
        row.cells[0].text = s; row.cells[1].text = c; row.cells[2].text = d
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    h2(doc, "每日巡檢表（商場獨有）")
    table = hdr_table(doc, ["結果狀態", "Tag 顏色", "說明", "列底色"])
    insp_rows = [
        ("正常",   "🟢 綠色",  "巡檢正常完成", "白色"),
        ("異常",   "🔴 紅色",  "發現異常",     "淺紅底色"),
        ("待處理", "🟡 黃色",  "需要後續處理", "白色"),
        ("未巡檢", "⬜ 灰色",  "未完成巡檢",   "淺灰底色"),
    ]
    for s, c, d, bg in insp_rows:
        row = table.add_row()
        row.cells[0].text = s; row.cells[1].text = c
        row.cells[2].text = d; row.cells[3].text = bg
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 手冊已儲存：{output_path}")


# ── 主程式 ────────────────────────────────────────────────────────────────────
async def main():
    OUTPUT_DIR.mkdir(exist_ok=True)
    print("=" * 55)
    print("  商場週期保養表 教學手冊生成器")
    print("=" * 55)
    print(f"目標：{TARGET_URL}")
    print(f"截圖：{OUTPUT_DIR}")
    print(f"手冊：{MANUAL_PATH}\n")

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=["--disable-blink-features=AutomationControlled"],
        )
        context = await browser.new_context(
            viewport={"width": 1600, "height": VIEWPORT_H},
            device_scale_factor=1.5,
            locale="zh-TW",
            timezone_id="Asia/Taipei",
        )
        page = await context.new_page()
        shots = await capture_screenshots(page, OUTPUT_DIR)
        await browser.close()

    print(f"\n截圖完成，共 {len(shots)} 張")
    print("\n📄 生成 Word 手冊...")
    build_word_manual(shots, MANUAL_PATH)
    print(f"\n✅ 執行完成！請開啟：{MANUAL_PATH}")


if __name__ == "__main__":
    asyncio.run(main())
