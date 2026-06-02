#!/usr/bin/env python3
"""
商場工務報修（/luqun-repair/dashboard）教學手冊自動生成腳本
===========================================================
使用方式：
  pip install playwright python-docx
  playwright install chromium
  python luqun_generate_manual.py

輸出：
  manual_screenshots_luqun/
  商場工務報修教學手冊.docx

架構與飯店工務報修（dazhi-repair）完全相同，
主要差異：商場資料來源、路由、標題、最後一 Tab 名稱（報修清單總表）
"""

import asyncio
import sys
from pathlib import Path
from datetime import datetime

try:
    from playwright.async_api import async_playwright
except ImportError:
    print("❌ 請先執行：pip install playwright && playwright install chromium"); sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 請先執行：pip install python-docx"); sys.exit(1)

# ── 設定 ──────────────────────────────────────────────────────────────────────
BASE_URL    = "http://localhost:5173"
TARGET_URL  = f"{BASE_URL}/luqun-repair/dashboard"
OUTPUT_DIR  = Path(__file__).parent / "manual_screenshots_luqun"
MANUAL_PATH = Path(__file__).parent / "商場工務報修教學手冊.docx"
USERNAME    = "admin"
PASSWORD    = "Admin@2026"

PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT  = RGBColor(0x4B, 0xA8, 0xE8)
GREEN   = RGBColor(0x52, 0xC4, 0x1A)
ORANGE  = RGBColor(0xFA, 0xAD, 0x14)
RED     = RGBColor(0xCF, 0x13, 0x22)
PURPLE  = RGBColor(0x72, 0x2E, 0xD1)
GRAY    = RGBColor(0x59, 0x59, 0x59)

CONTENT_X  = 185
CONTENT_W  = 1390
VIEWPORT_H = 900


# ── 截圖邏輯 ──────────────────────────────────────────────────────────────────
async def capture_screenshots(page, output_dir: Path) -> dict:
    shots = {}

    async def snap(name: str, scroll_y: int = 0, clip=None, wait_ms: int = 800):
        try:
            await page.evaluate(f"window.scrollTo(0, {scroll_y})")
            await page.wait_for_timeout(wait_ms)
            path = output_dir / f"{name}.png"
            await page.screenshot(path=str(path), full_page=False,
                                  **({"clip": clip} if clip else {}))
            shots[name] = path
            print(f"   ✓ {name}")
        except Exception as e:
            print(f"   ✗ {name}：{e}")

    async def click_tab(tab_text: str, wait_ms: int = 2000):
        try:
            tab = page.locator("div.ant-tabs-tab").filter(has_text=tab_text).first
            await tab.click()
            await page.wait_for_timeout(wait_ms)
        except Exception as e:
            print(f"   ⚠ Tab [{tab_text}] 失敗：{e}")

    # ── 登入 ──────────────────────────────────────────────────────────────────
    print("🔐 登入中...")
    await page.goto(f"{BASE_URL}/login")
    await page.wait_for_load_state("domcontentloaded")
    await page.wait_for_timeout(800)
    await page.locator("input[placeholder*='帳號'], input[type='text']").first.fill(USERNAME)
    await page.locator("input[type='password']").first.fill(PASSWORD)
    await page.locator("button[type='submit'], button:has-text('登入')").first.click()
    try:
        await page.wait_for_url(f"{BASE_URL}/dashboard", timeout=15000)
        print("   ✓ 登入成功")
    except Exception:
        print(f"   ✓ 已跳轉至 {page.url}")

    # ── 導航 ──────────────────────────────────────────────────────────────────
    await page.goto(TARGET_URL)
    await page.wait_for_load_state("networkidle")
    try:
        await page.locator(".ant-statistic-content-value").first.wait_for(timeout=12000)
        await page.wait_for_timeout(2500)
        print("   ✓ 頁面載入完成")
    except Exception:
        await page.wait_for_timeout(5000)

    # ════════════════════════════════════════════════════════════
    # Tab: Dashboard
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：Dashboard Tab")
    await page.evaluate("window.scrollTo(0, 0)")
    await page.wait_for_timeout(600)

    # 1. 頁頭 + 查詢列 + KPI 7 卡
    await snap("01_header_kpi",
               clip={"x": CONTENT_X, "y": 55, "width": CONTENT_W, "height": 220})

    # 2. 費用 KPI 3 張
    await snap("02_fee_kpi",
               clip={"x": CONTENT_X, "y": 320, "width": CONTENT_W, "height": 160})

    # 3. 趨勢圖 + 圓餅圖
    await snap("03_charts_top",
               clip={"x": CONTENT_X, "y": 490, "width": CONTENT_W, "height": 330})

    # 4. 樓層分布 + 狀況分布
    await snap("04_charts_bottom",
               scroll_y=720,
               clip={"x": CONTENT_X, "y": 720, "width": CONTENT_W, "height": 280})

    # 5. Top 10 三欄
    await snap("05_top10",
               scroll_y=1010,
               clip={"x": CONTENT_X, "y": 1010, "width": CONTENT_W, "height": 480})

    # 6. 點未完成件數 KPI 展開 Modal
    print("📸 截圖：KPI Modal")
    try:
        await page.evaluate("window.scrollTo(0, 0)")
        await page.wait_for_timeout(400)
        uncompleted = page.locator(".ant-card").filter(has_text="未完成件數").first
        await uncompleted.click()
        await page.wait_for_timeout(1500)
        await snap("06_kpi_modal",
                   clip={"x": 100, "y": 60, "width": 1350, "height": 640})
        close = page.locator(".ant-modal-close").first
        await close.click()
        await page.wait_for_timeout(500)
    except Exception as e:
        print(f"   ✗ KPI Modal：{e}")

    # 7. 點 Top 10 展開案件 Drawer
    print("📸 截圖：案件詳情 Drawer")
    try:
        await page.evaluate("window.scrollTo(0, 1010)")
        await page.wait_for_timeout(500)
        first_case = page.locator(".ant-card").filter(has_text="未完成案件").first \
                          .locator("div[style*='cursor: pointer']").first
        await first_case.click()
        await page.wait_for_timeout(1500)
        await snap("07_case_drawer",
                   clip={"x": 900, "y": 55, "width": 680, "height": 700})
        close = page.locator(".ant-drawer-close").first
        await close.click()
        await page.wait_for_timeout(500)
    except Exception as e:
        print(f"   ✗ Drawer：{e}")

    # ════════════════════════════════════════════════════════════
    # Tab: 3.1 報修統計
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：3.1 報修統計")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("3.1 報修", wait_ms=2500)
    await snap("08_repair_stats",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab: 3.2 結案時間
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：3.2 結案時間")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("3.2 結案時間", wait_ms=2500)
    await snap("09_closing_time",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab: 3.3 報修類型
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：3.3 報修類型")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("3.3 報修類型", wait_ms=2500)
    await snap("10_repair_type",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab: 3.4 本月客房報修表
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：3.4 客房報修表")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("3.4 本月客房報修表", wait_ms=2500)
    await snap("11_room_repair",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab: 金額統計
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：金額統計")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("金額統計", wait_ms=2500)
    await snap("12_fee_stats",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 580})

    # ════════════════════════════════════════════════════════════
    # Tab: 報修清單總表
    # ════════════════════════════════════════════════════════════
    print("📸 截圖：報修清單總表")
    await page.evaluate("window.scrollTo(0, 0)")
    await click_tab("報修清單總表", wait_ms=2500)
    await snap("13_detail_toolbar",
               clip={"x": CONTENT_X, "y": 150, "width": CONTENT_W, "height": 130})
    await snap("14_detail_list",
               clip={"x": CONTENT_X, "y": 280, "width": CONTENT_W, "height": 520})

    return shots


# ── Word 手冊輔助函式 ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)


def add_image_safe(doc, path, width_inches: float = 5.8):
    if isinstance(path, Path) and path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph("[ 截圖未能取得，請參考系統實際畫面 ]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = GRAY; p.runs[0].font.italic = True


def add_tip(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run("💡 提示："); r.bold = True; r.font.color.rgb = PRIMARY
    p.add_run(text)


def hdr_table(doc, headers, bg="1B3A5C"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, bg)
    return table


def h1(doc, t):
    h = doc.add_heading(t, 1); h.runs[0].font.color.rgb = PRIMARY; return h

def h2(doc, t):
    h = doc.add_heading(t, 2); h.runs[0].font.color.rgb = ACCENT; return h


# ── Word 手冊建構 ─────────────────────────────────────────────────────────────
def build_word_manual(shots: dict, output_path: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11.69); sec.page_width  = Inches(8.27)
    sec.top_margin  = Inches(0.9);   sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1.1);   sec.right_margin  = Inches(1.1)

    # ── 封面 ──────────────────────────────────────────────────────────────────
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_heading("商場工務報修", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = PRIMARY; t.runs[0].font.size = Pt(28)
    t2 = doc.add_heading("教學操作手冊", 0)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = ACCENT; t2.runs[0].font.size = Pt(22)
    doc.add_paragraph()
    sub = doc.add_paragraph(
        f"春大直 · 路由：/luqun-repair/dashboard\n{datetime.now().strftime('%Y 年 %m 月 %d 日')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    doc.add_page_break()

    # ── 一、功能簡介 ──────────────────────────────────────────────────────────
    h1(doc, "一、功能簡介")
    p = doc.add_paragraph()
    p.add_run("商場工務報修").bold = True
    p.add_run(
        "（路徑：商場管理 ▶ 商場工務報修，路由：/luqun-repair/dashboard）"
        "管理商場（陸群大樓）工務部的所有報修案件，"
        "提供 KPI 總覽、費用追蹤、結案分析、工時統計與 PowerPoint 匯出。"
        "\n\n"
        "功能架構與飯店工務報修（dazhi-repair）完全對應，"
        "如已熟悉飯店版，只需注意資料來源為商場端。"
    )
    doc.add_paragraph()

    tab_info = [
        ("Dashboard",     "KPI 7 卡 + 費用 3 卡 + 圖表 + Top 10 快速摘要（預設 Tab）"),
        ("3.1 報修",      "月份報修走勢折線圖 + 可點月份查看案件"),
        ("3.2 結案時間",  "結案天數分布分析"),
        ("3.3 報修類型",  "各報修類型件數與占比"),
        ("3.4 本月客房報修表", "依客房房號列出本月商場報修記錄"),
        ("金額統計",      "年度費用彙整（委外/維修/扣款）"),
        ("報修清單總表",  "完整報修案件清單（多條件篩選 + 匯出 Excel）"),
    ]
    table = hdr_table(doc, ["頁籤", "說明"])
    for name, desc in tab_info:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ── 二、查詢列 & 頁頭工具 ─────────────────────────────────────────────────
    h1(doc, "二、查詢列 & 頁頭工具")
    add_image_safe(doc, shots.get("01_header_kpi"))
    doc.add_paragraph()
    doc.add_paragraph("頁面頂部查詢列控制所有 Tab 的資料範圍（各 Tab 共享同一查詢條件）：")
    for name, desc in [
        ("年度選單",       "選擇查詢年份（有資料的年份）"),
        ("月份選單",       "選「全年」或特定月份"),
        ("查詢按鈕",       "確認條件後點擊，各 Tab 同步更新"),
        ("重設按鈕",       "回到預設年月"),
        ("匯出 PowerPoint","紫色漸層按鈕，匯出 .pptx 簡報"),
        ("連線測試",       "測試後端與 Ragic 的連線狀態"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"• {name}：").bold = True
        p.add_run(desc)
    doc.add_page_break()

    # ── 三、Dashboard — KPI 7 卡 ──────────────────────────────────────────────
    h1(doc, "三、Dashboard — KPI 7 卡")
    add_image_safe(doc, shots.get("01_header_kpi"))
    doc.add_paragraph()
    doc.add_paragraph(
        "7 張 KPI 卡片全部可點擊，展開符合條件的案件清單 Modal（含 ℹ 圖示說明計算口徑）："
    )
    doc.add_paragraph()

    kpi_rows = [
        ("本月相關案件", "深藍", "上期未結 + 本期報修件數（跨月累計）"),
        ("已完成件數",   "綠色", "累計已完成 + 本期已完成（以已驗收為標準）"),
        ("待辦驗件數",   "黃色", "處理狀況 = 待辦驗的商場報修案件"),
        ("未完成件數",   "紅色", "累計未完成 + 本期未完成"),
        ("平均結案天數", "藍色", "已完成案件的平均完工天數"),
        ("本月工時統計", "青色", "有工時記錄的商場報修案件總工時（hr），附換算工作天參考"),
        ("客房報修件數", "橙色", "本期涉及商場客房的報修件數"),
    ]
    table = hdr_table(doc, ["KPI 名稱", "顏色", "統計口徑說明"])
    for name, color, desc in kpi_rows:
        row = table.add_row()
        for i, v in enumerate([name, color, desc]):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    add_tip(doc, "點擊任一 KPI 卡片，展開 Modal 查看符合條件的每筆案件明細。ℹ 圖示停留可查看計算口徑說明。")
    doc.add_paragraph()
    add_image_safe(doc, shots.get("06_kpi_modal"), width_inches=5.5)
    doc.add_paragraph("▲ 點擊 KPI 卡片後展開的案件清單 Modal（以未完成件數為例）")
    doc.add_page_break()

    # ── 四、費用 KPI 3 張 ─────────────────────────────────────────────────────
    h1(doc, "四、Dashboard — 費用 KPI 3 張")
    add_image_safe(doc, shots.get("02_fee_kpi"))
    doc.add_paragraph()
    fee_rows = [
        ("委外+維修費用", "紫色", "YTD 累計至選定月，點擊展開費用明細 Modal（含委外/維修各別金額）"),
        ("扣款費用",      "紅色", "YTD 累計扣款費用，點擊展開扣款事項明細"),
        ("當月金額",      "青色", "當月委外+維修 / 扣款費用 / 扣款專櫃 / 小計（不可點擊）"),
    ]
    table = hdr_table(doc, ["卡片", "顏色", "說明"])
    for name, color, desc in fee_rows:
        row = table.add_row()
        for i, v in enumerate([name, color, desc]):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ── 五、圖表區 ────────────────────────────────────────────────────────────
    h1(doc, "五、Dashboard — 圖表區")
    add_image_safe(doc, shots.get("03_charts_top"))
    doc.add_paragraph()
    for title, desc in [
        ("近 12 個月報修趨勢（折線）", "深藍線=報修件數，綠線=完成件數，反映商場工務量走勢"),
        ("報修類型分布（圓餅）",       "各報修類型占比，停留顯示精確數字"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"▸ {title}：").bold = True
        p.add_run(desc)
    doc.add_paragraph()
    add_image_safe(doc, shots.get("04_charts_bottom"))
    doc.add_paragraph()
    for title, desc in [
        ("發生樓層分布（水平 Bar）", "各樓層報修件數，找出商場報修熱點樓層"),
        ("處理狀況分布（水平 Bar）", "各狀態件數，顏色對應狀態類型"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"▸ {title}：").bold = True
        p.add_run(desc)
    doc.add_page_break()

    # ── 六、Top 10 快速摘要 & 案件 Drawer ─────────────────────────────────────
    h1(doc, "六、Dashboard — Top 10 快速摘要 & 案件詳情 Drawer")
    add_image_safe(doc, shots.get("05_top10"))
    doc.add_paragraph()
    top10_rows = [
        ("未完成案件 Top 10", "🔴 紅色", "依等待天數（pending_days）降序；紅色≥30天，橙色7–29天"),
        ("高費用案件 Top 10", "🟣 紫色", "依費用金額降序；顯示金額與是否已結案"),
        ("高工時案件 Top 10", "🔵 藍色", "依工時降序；顯示工時（hr）與是否已結案"),
    ]
    table = hdr_table(doc, ["清單", "顏色", "說明"])
    for name, color, desc in top10_rows:
        row = table.add_row()
        for i, v in enumerate([name, color, desc]):
            row.cells[i].text = v
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()
    add_image_safe(doc, shots.get("07_case_drawer"), width_inches=4.0)
    doc.add_paragraph()
    doc.add_paragraph(
        "點擊任一 Top 10 清單的案件，右側滑出案件詳情 Drawer（寬 520px），"
        "包含：報修編號、標題（粗體）、報修類型、發生樓層、發生時間、"
        "負責單位、工時、處理狀況 Tag、委外/維修/總費用、扣款、驗收/結案資訊、"
        "維修圖片（Lightbox）、「在 Ragic 查看」連結。"
    )
    add_tip(doc, "維修圖片由系統自動從資料庫載入，點縮圖可放大預覽並左右切換。")
    doc.add_page_break()

    # ── 七、統計分析 Tab ──────────────────────────────────────────────────────
    h1(doc, "七、統計分析 Tab（3.1 – 3.4）& 金額統計")

    h2(doc, "7.1  3.1 報修統計")
    add_image_safe(doc, shots.get("08_repair_stats"))
    doc.add_paragraph("月份折線圖：可點圖表月份查看當月案件明細。觀察商場工務量的季節性規律。")
    doc.add_paragraph()

    h2(doc, "7.2  3.2 結案時間統計")
    add_image_safe(doc, shots.get("09_closing_time"))
    doc.add_paragraph("結案天數分布（0–7天、8–30天等區間），分析商場工務結案效率。")
    doc.add_paragraph()

    h2(doc, "7.3  3.3 報修類型統計")
    add_image_safe(doc, shots.get("10_repair_type"))
    doc.add_paragraph("各報修類型（衛浴、建築、設備等）的件數占比，作為資源配置依據。")
    doc.add_paragraph()

    h2(doc, "7.4  3.4 本月客房報修表")
    add_image_safe(doc, shots.get("11_room_repair"))
    doc.add_paragraph("依客房房號列出本月商場報修記錄，找出常態問題的客房單元。")
    doc.add_paragraph()

    h2(doc, "7.5  金額統計")
    add_image_safe(doc, shots.get("12_fee_stats"))
    doc.add_paragraph("年度費用彙整，顯示委外/維修/扣款費用的月份分布與全年合計。")
    doc.add_page_break()

    # ── 八、報修清單總表 ──────────────────────────────────────────────────────
    h1(doc, "八、報修清單總表（春大直）")
    add_image_safe(doc, shots.get("13_detail_toolbar"))
    doc.add_paragraph()
    add_image_safe(doc, shots.get("14_detail_list"))
    doc.add_paragraph()
    doc.add_paragraph(
        "「報修清單總表」是商場版的完整案件清單，標籤名稱為「春大直-報修清單總表」。"
        "功能與飯店版「飯店工務報修」Tab 完全相同——"
        "多條件篩選（年月/類型/樓層/狀態/關鍵字）+ 點列展開 Drawer + 匯出 Excel。"
    )
    doc.add_paragraph()

    col_rows = [
        ("報修編號",   "點擊跳轉 Ragic 商場報修 Sheet"),
        ("標題",       "報修摘要"),
        ("樓層",       "商場樓層"),
        ("報修時間",   "報修申請時間"),
        ("狀態",       "彩色 Tag"),
        ("完工時間",   "驗收完成時間"),
        ("結案天數",   "報修到結案天數"),
        ("工時(hr)",   "維修工時"),
        ("費用",       "委外+維修費用合計"),
    ]
    table = hdr_table(doc, ["欄位", "說明"])
    for name, desc in col_rows:
        row = table.add_row()
        row.cells[0].text = name; row.cells[1].text = desc
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)
    add_tip(doc, "右上角「匯出 Excel」按鈕下載目前篩選結果，檔名含年月條件。")
    doc.add_page_break()

    # ── 九、常見操作情境 ──────────────────────────────────────────────────────
    h1(doc, "九、常見操作情境")
    scenarios = [
        ("月底確認商場工務未完成件數",
         ["查詢列選本月 → Dashboard",
          "查看「未完成件數」KPI，點擊展開 Modal",
          "查看「未完成案件 Top 10」找等待最久的案件"]),
        ("確認高費用商場報修案件",
         ["Dashboard → 點「委外+維修費用」卡展開 Modal",
          "或查看「高費用案件 Top 10」點入 Drawer 確認明細"]),
        ("分析商場哪個類型報修最多",
         ["切換到「3.3 報修類型」Tab，查看圓餅圖或長條圖占比",
          "結果作為商場保養計劃資源配置依據"]),
        ("查詢特定商場報修案件完整記錄",
         ["報修清單總表 Tab → 關鍵字輸入報修編號 → 點列展開 Drawer",
          "確認完整欄位，點「在 Ragic 查看」確認原始記錄"]),
        ("匯出月報 PowerPoint",
         ["確認查詢列選定特定月份",
          "點頁頭「匯出 PowerPoint」按鈕，等待下載"]),
    ]
    for i, (title, steps) in enumerate(scenarios, 1):
        p = doc.add_paragraph()
        p.add_run(f"情境 {i}：{title}").bold = True
        p.runs[0].font.color.rgb = PRIMARY
        for step in steps:
            doc.add_paragraph(f"　→ {step}")
        doc.add_paragraph()

    # ── 附錄 ──────────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "附錄：處理狀況顏色對照表")
    table = hdr_table(doc, ["處理狀況", "Tag 顏色", "意義"])
    status_rows = [
        ("已驗收 / 已結案 / 完修", "🟢 success 綠色",   "案件已完成"),
        ("處理中",                  "🔵 processing 藍色","執行中"),
        ("待維修 / 待驗收",         "🟡 warning 黃色",   "等待下一步"),
        ("待協調",                  "🟠 orange 橙色",   "需協調"),
        ("待排除",                  "🔴 error 紅色",    "有問題需即刻處理"),
    ]
    for s, c, d in status_rows:
        row = table.add_row()
        row.cells[0].text = s; row.cells[1].text = c; row.cells[2].text = d
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    h1(doc, "附錄：與飯店工務報修（dazhi-repair）對照")
    table = hdr_table(doc, ["項目", "飯店工務報修", "商場工務報修"])
    diff_rows = [
        ("路由",          "/hotel/dazhi-repair/dashboard",      "/luqun-repair/dashboard"),
        ("頁面標題",      "飯店工務報修",                        "商場工務報修"),
        ("副標題",        "（無）",                              "春大直 - 報修清單總表"),
        ("最後 Tab",      "飯店工務報修",                        "報修清單總表"),
        ("Ragic Sheet",   "lequn-public-works/8",               "luqun-public-works-repair-reporting-system/6"),
        ("Tab 數量",      "7 個",                               "7 個"),
        ("KPI 結構",      "7 卡（同）",                         "7 卡（同）"),
        ("功能",          "完全相同",                            "完全相同"),
    ]
    for item, dazhi, luqun in diff_rows:
        row = table.add_row()
        row.cells[0].text = item; row.cells[1].text = dazhi; row.cells[2].text = luqun
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(str(output_path))
    print(f"\n✅ 手冊已儲存：{output_path}")


# ── 主程式 ────────────────────────────────────────────────────────────────────
async def main():
    OUTPUT_DIR.mkdir(exist_ok=True)
    print("=" * 55)
    print("  商場工務報修 教學手冊生成器")
    print("=" * 55)
    print(f"目標：{TARGET_URL}")
    print(f"截圖：{OUTPUT_DIR}")
    print(f"手冊：{MANUAL_PATH}\n")

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=["--disable-blink-features=AutomationControlled"],
        )
        context = await browser.new_context(
            viewport={"width": 1600, "height": VIEWPORT_H},
            device_scale_factor=1.5,
            locale="zh-TW",
            timezone_id="Asia/Taipei",
        )
        page = await context.new_page()
        shots = await capture_screenshots(page, OUTPUT_DIR)
        await browser.close()

    print(f"\n截圖完成，共 {len(shots)} 張")
    print("\n📄 生成 Word 手冊...")
    build_word_manual(shots, MANUAL_PATH)
    print(f"\n✅ 執行完成！請開啟：{MANUAL_PATH}")


if __name__ == "__main__":
    asyncio.run(main())
